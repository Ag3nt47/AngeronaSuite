from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parent
ANALYSIS = ROOT / "analysis"
VERSION = "1.10.1"
DATE = "2026-08-24"


def copy_document(source: str, target: str) -> tuple[Path, Document]:
    src = ANALYSIS / source
    dst = ANALYSIS / target
    shutil.copy2(src, dst)
    return dst, Document(dst)


def set_properties(doc: Document, *, title: str, subject: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.version = VERSION
    props.last_modified_by = "Angerona documentation loop"


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def insert_bullet_before(paragraph, text: str):
    return paragraph.insert_paragraph_before(text, style="List Bullet")


def update_security() -> Path:
    path, doc = copy_document(
        "Angerona_Security_Assessment_v2.3_2026-07-27.docx",
        f"Angerona_Security_Assessment_v{VERSION}_{DATE}.docx",
    )
    doc.paragraphs[2].text = (
        f"Assessment version {VERSION} · {DATE} · Host Adaption improvement-loop "
        "assessment, remediation, and final validation."
    )
    set_properties(
        doc,
        title=f"Angerona Security Assessment v{VERSION}",
        subject="Host Adaption improvement-loop security assessment",
    )

    doc.add_paragraph(f"Loop remediation — {DATE}", style="Heading 1")
    doc.add_paragraph(
        "Scope: the new Adaption / Adapt to Host service, dashboard workbench, "
        "Windows Firewall mutation boundary, context automation, exception and "
        "feedback logic, rollback, and performance path. This was one combined "
        "bug, adversary, remediation, performance, and visionary loop; no external "
        "penetration test or real elevated firewall mutation was performed."
    )
    doc.add_paragraph("Findings found and fixed", style="Heading 2")
    add_bullet(doc, "R8-RT-01 (Medium) — stale automation authority: fixed with monotonic state revisions, compare-and-swap updates, a single-flight transaction, fresh context checks, and final pre-execution authorization.")
    add_bullet(doc, "R8-RT-02 (Medium) — command success without effective-state proof: fixed with explicit Firewall ActiveStore collection, exact apply postconditions, automatic verified recovery, and restore postconditions.")
    add_bullet(doc, "R8-RT-03 (Medium) — over-broad exceptions and feedback: fixed with exact before/after finding fingerprints, one-shot labels, three distinct reviewed findings before tuning, and a 0.75 weight floor.")
    add_bullet(doc, "R8-RT-04 (Medium) — weaker context precedence: fixed with strongest-posture ordering, conservative Public-network detection across all active adapters, and refusal of automatic relaxation.")
    add_bullet(doc, "R8-RT-05 (Medium) — incomplete collector authority: fixed with explicit quality metadata, privacy-minimized service/listener identity, effective firewall profiles, and a bounded rule inventory; incomplete categories are not scored as healthy.")
    add_bullet(doc, "R8-BT-01 through R8-BT-03 — context re-entry, simultaneous Public-network classification, and batch selfcheck exit-code propagation were found and fixed.")
    doc.add_paragraph("Mitigations and safety boundary", style="Heading 2")
    add_bullet(doc, "Balanced, Public Network, and Emergency Lockdown remain a closed firewall-only catalog. Plans are short-lived, host/precondition-bound, previewed exactly, and require exact approval; the sandbox is no-write.")
    add_bullet(doc, "Production apply requires a verified Windows Firewall export. Trusted absolute Windows tools and a sanitized child environment are used. Receipt failure or postcondition mismatch enters automatic recovery.")
    add_bullet(doc, "Automation defaults to proposal-only. Separately armed auto-apply is revision-bound, serialized, context-revalidated, non-relaxing, and rate-limited by a persistent circuit breaker.")
    doc.add_paragraph("Deferred and external", style="Heading 2")
    add_bullet(doc, "Deferred implementation depth: per-rule program/service/address/port firewall filter joins and executable signer/content-hash attestation for service binaries.")
    add_bullet(doc, "External acceptance: real elevated apply/rollback with connectivity-loss recovery and a physical simultaneous Public/Private/VPN topology.")
    doc.add_paragraph("Verification", style="Heading 2")
    add_bullet(doc, "Focused host-adaptation/UI/performance set: 20/20 passed.")
    add_bullet(doc, "Full repository suite: 1077 passed, 3 intentional platform skips, 0 failed in 75.71 seconds.")
    add_bullet(doc, "Headless selfcheck: 26/26 passed.")

    doc.save(path)
    return path


def update_capability() -> Path:
    path, doc = copy_document(
        "Angerona_Capability_Doc_v1.9.4.docx",
        f"Angerona_Capability_Doc_v{VERSION}.docx",
    )
    doc.paragraphs[2].text = f"Document version {VERSION} · {DATE}"
    doc.paragraphs[3].text = (
        "66 auto-discovered security modules; v1.10.1 adds the Host Adaption "
        "workbench and its audited, reversible Windows Firewall profile workflow."
    )
    set_properties(
        doc,
        title=f"Angerona Capability Doc v{VERSION}",
        subject="Host Adaption capability update",
    )

    anchor = doc.paragraphs[4]
    heading = anchor.insert_paragraph_before(
        f"What's New in v{VERSION} — Host Adaption", style="Heading 1"
    )
    # Insert content before the prior historical What's New block, preserving order.
    items = [
        "The dashboard's top-left ADAPTION button opens a menu-complete seven-tab workbench: Overview, Audit & Drift, Exceptions & Feedback, Profiles & Rollback, Sandbox, Automation, and Activity.",
        "Deep local audits capture hardware, privacy-minimized services and listeners, network context, effective Windows Firewall ActiveStore profiles, and a bounded rule inventory. Collector completeness and truncation remain visible; golden baselines, exact exceptions, drift scoring, and JSON/CSV exports are built in.",
        "Balanced, Public Network, and Emergency Lockdown use a closed typed catalog. Exact short-lived dry-run plans show their command stack; the sandbox cannot write. Apply is approval-, precondition-, snapshot-, postcondition-, and rollback-gated.",
        "SSID, VPN, and Public-network triggers begin proposal-only. Auto-apply is separately armed, revision-bound, single-flight, re-authorized immediately before mutation, strongest-posture-first, non-relaxing, and circuit-broken.",
        "Bug/adversary convergence fixed five Medium safety findings and three QA defects. Unchanged 500-row activity refreshes measured 14.6x faster. Final evidence: focused 20/20, pytest 1077 passed with 3 intentional platform skips, and selfcheck 26/26.",
        "Deferred: deep firewall program/service/address/port joins, service executable signer/content-hash attestation, event-driven network wakeups, versioned baseline promotion, crash-independent trial leases, shadow feedback, and signed non-executable posture packs.",
    ]
    # insert_paragraph_before always inserts immediately before anchor; iterate forward.
    for text in items:
        anchor.insert_paragraph_before(text, style="List Bullet")

    doc.add_paragraph(f"v{VERSION} Host Adaption verification", style="Heading 1")
    add_bullet(doc, "66 modules remain auto-discovered; Host Adaption is a core service/workbench, not a new BaseModule.")
    add_bullet(doc, "Final focused host-adaptation/UI/performance tests: 20/20 passed; full pytest: 1077 passed / 3 intentional platform skips; selfcheck: 26/26.")
    doc.save(path)
    return path


def update_master() -> Path:
    path, doc = copy_document(
        "Angerona_Master_Manual_v1.10.0_Cycle22.docx",
        f"Angerona_Master_Manual_v{VERSION}.docx",
    )
    doc.paragraphs[4].text = (
        f"Host Adaption improvement-loop edition  |  24 August 2026  |  v{VERSION}"
    )
    set_properties(
        doc,
        title="Angerona Consolidated Master Manual",
        subject=f"Version {VERSION} — Host Adaption verified addendum",
    )
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                if "v1.10 Local SOC consolidated 2026-08-22" in run.text:
                    run.text = run.text.replace(
                        "v1.10 Local SOC consolidated 2026-08-22",
                        f"v{VERSION} Host Adaption verified {DATE}",
                    )

    # Contents entry immediately before Appendix A.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Appendix A. Consolidation Sources":
            paragraph.insert_paragraph_before(
                "23. Host Adaption — Adapt to Host", style="List Bullet"
            )
            break

    # Exact current companion-document cross-reference near the front matter.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("This manual consolidates"):
            paragraph.insert_paragraph_before(
                "Current companion documents: "
                f"Angerona_Capability_Doc_v{VERSION}.docx; "
                f"Angerona_Security_Assessment_v{VERSION}_{DATE}.docx; and "
                f"Angerona_Vulnerabilities_Assessment_Remediation_v{VERSION}.docx.",
                style="Manual Note",
            )
            break
    for paragraph in doc.paragraphs:
        if "full module-by-module table" in paragraph.text and "companion Capability Doc" in paragraph.text:
            paragraph.text = (
                "66 BaseModule subclasses are auto-discovered at startup. The full "
                "module-by-module table is maintained in the companion "
                f"Angerona_Capability_Doc_v{VERSION}.docx; Host Adaption is a core "
                "service/workbench and does not change the module count."
            )
            break

    # Condensed version-history row.
    history = doc.tables[6]
    row = history.add_row().cells
    row[0].text = f"v{VERSION} / Round 8"
    row[1].text = (
        "Host Adaption workbench; quality-aware audit/baseline/drift; exact "
        "exceptions and bounded feedback; typed preview/sandbox/snapshot/apply/"
        "verified rollback; conservative revision-bound automation and breaker; "
        "five Medium adversary findings and three QA defects fixed; 1077 tests "
        "passed with 3 intentional platform skips."
    )

    doc.add_paragraph("23. Host Adaption — Adapt to Host", style="Heading 1")
    doc.add_paragraph(
        "Verified addendum — 24 August 2026. Host Adaption is a local, "
        "operator-controlled Windows Firewall posture workflow. It separates "
        "observation, planning, simulation, approval, mutation, verification, "
        "rollback, automation, and feedback."
    )
    doc.add_paragraph("23.1 Operator surface", style="Heading 2")
    add_bullet(doc, "Open the top-left ADAPTION button. File, Audit, Profiles, Safety, and Help menus complement Overview, Audit & Drift, Exceptions & Feedback, Profiles & Rollback, Sandbox, Automation, and Activity tabs.")
    add_bullet(doc, "Status cards show baseline, drift risk, active profile, automation, and breaker state. Incomplete collector coverage is prominent in the audit, baseline confirmation, and footer status.")
    doc.add_paragraph("23.2 Audit, baseline, drift, and exports", style="Heading 2")
    add_bullet(doc, "The snapshot includes hardware; privacy-minimized service command/account identity; protocol/family/scope-preserving listeners; SSID/VPN/network category; effective ActiveStore firewall profiles; and a bounded rule inventory.")
    add_bullet(doc, "Each bounded collector reports completeness, availability, truncation, skipped rows, and a sanitized reason. An incomplete baseline/current pair is not scored as healthy drift.")
    add_bullet(doc, "The operator may save a golden baseline, run drift checks, pin an exact finding fingerprint as a known-good exception, dismiss distinct false positives into bounded feedback, and export formula-safe JSON or CSV.")
    doc.add_paragraph("23.3 Preview, sandbox, apply, and recovery", style="Heading 2")
    add_bullet(doc, "Balanced, Public Network, and Emergency Lockdown are a closed firewall-only catalog. A short-lived immutable plan binds host, exact profile actions, command preview, and the effective firewall precondition digest.")
    add_bullet(doc, "Sandbox projects the plan against captured state and cannot write. Real apply requires exact plan confirmation, unchanged preconditions, a verified Windows Firewall export, trusted absolute tools, a sanitized environment, and effective ActiveStore postconditions.")
    add_bullet(doc, "Apply failure or receipt-state failure enters automatic recovery. One-click rollback requires explicit approval, verifies the signed manifest and firewall artifact digest, imports only the bound snapshot, and verifies restored effective profiles and bounded rules.")
    doc.add_paragraph("23.4 Context automation, breaker, and feedback", style="Heading 2")
    add_bullet(doc, "SSID, VPN-active, and Public-network triggers are opt-in and proposal-only by default. Auto-apply is separately armed and re-authorized at the final mutation boundary.")
    add_bullet(doc, "All matching rules are ordered by profile strength. Public evidence cannot be hidden by adapter enumeration or weakened by a familiar SSID, and automation never relaxes the observed posture.")
    add_bullet(doc, "Revision compare-and-swap checks and a single-flight transaction stop stale or concurrent workers. The persistent breaker enforces cooldown and repeated/drastic-change budgets. Feedback requires three distinct exact reviewed findings before its bounded effect begins.")
    doc.add_paragraph("23.5 Loop results and remaining gates", style="Heading 2")
    add_bullet(doc, "Adversary remediation closed five Medium findings; QA closed context re-entry, simultaneous Public-category, and batch selfcheck exit-code defects.")
    add_bullet(doc, "Qt timer state I/O moved off the interface thread, signed-state reads were coalesced, and unchanged table rebuilds were skipped. A 500-row activity refresh measured 0.869 ms unchanged versus 12.671 ms forced (14.6x).")
    add_bullet(doc, "Final verification: focused 20/20; repository pytest 1077 passed / 3 intentional platform skips / 0 failed; selfcheck 26/26.")
    add_bullet(doc, "External/deferred: real elevated connectivity acceptance, physical mixed-network topology, deep firewall program/service/address/port joins, service executable signer/content-hash attestation, event-driven wakeups, versioned baseline lifecycle, trial leases, shadow feedback, and signed posture packs.")
    doc.save(path)
    return path


def update_vulnerabilities() -> Path:
    path, doc = copy_document(
        "Angerona_Vulnerabilities_Assessment_Remediation_v1.9.4.docx",
        f"Angerona_Vulnerabilities_Assessment_Remediation_v{VERSION}.docx",
    )
    doc.paragraphs[2].text = (
        f"Document version {VERSION} · {DATE} · Host Adaption improvement loop: "
        "the new capability is recorded separately from eight implementation "
        "defects, all fixed; deeper collection and physical-host acceptance remain explicit."
    )
    set_properties(
        doc,
        title=f"Angerona Vulnerabilities Assessment and Remediation v{VERSION}",
        subject="Host Adaption feature-versus-defect remediation record",
    )
    doc.add_paragraph(f"Host Adaption feature-versus-defect record — v{VERSION}", style="Heading 1")
    doc.add_paragraph("Feature delivered", style="Heading 2")
    doc.add_paragraph(
        "The ADAPTION dashboard entry and Host Adaption workbench are a new "
        "defensive capability, not a defect closure. They add quality-aware "
        "audit/baselining/drift, exact exceptions, JSON/CSV export, typed firewall "
        "profiles, dry-run preview, a no-write sandbox, snapshot/apply/verified "
        "rollback, context triggers, a circuit breaker, bounded feedback, and a "
        "complete operator menu/tab surface."
    )
    doc.add_paragraph("Defects found and fixed", style="Heading 2")
    add_bullet(doc, "R8-RT-01 (Medium): stale auto-cycle authority — fixed with revision/CAS checks and final authorization.")
    add_bullet(doc, "R8-RT-02 (Medium): unverified command/restore outcome — fixed with effective ActiveStore apply and rollback postconditions plus automatic recovery.")
    add_bullet(doc, "R8-RT-03 (Medium): broad exception/feedback scope — fixed with exact fingerprints, distinct one-shot review gating, and a 0.75 floor.")
    add_bullet(doc, "R8-RT-04 (Medium): weaker SSID context could outrank Public — fixed with all-network collection, strongest-profile ordering, and no automatic relaxation.")
    add_bullet(doc, "R8-RT-05 (Medium): incomplete service/listener/firewall coverage could look authoritative — fixed with quality metadata, stronger privacy-minimized identities, explicit ActiveStore profiles, and bounded rules.")
    add_bullet(doc, "R8-BT-01: a matched context returning after no-match was incorrectly deduplicated — fixed.")
    add_bullet(doc, "R8-BT-02: simultaneous Public network could be hidden by first-adapter selection — fixed with conservative all-active-category precedence.")
    add_bullet(doc, "R8-BT-03: run-selfcheck.bat could return success after Python failure — fixed by preserving the Python exit code.")
    doc.add_paragraph("Deferred enhancement, not unresolved defect", style="Heading 2")
    add_bullet(doc, "Deep per-rule firewall filter joins for program, service, local/remote address, and local/remote port are deferred.")
    add_bullet(doc, "Service executable signer and content-hash attestation are deferred; the current audit uses a privacy-minimized command digest, executable name, and account type/identifier.")
    add_bullet(doc, "Crash-independent trial leases, event-driven wakeups, versioned baseline promotion, poisoning-resistant shadow feedback, and signed posture packs remain proposed enhancements.")
    doc.add_paragraph("Verification and acceptance boundary", style="Heading 2")
    add_bullet(doc, "Focused host-adaptation/UI/performance set: 20/20 passed; full pytest: 1077 passed, 3 intentional platform skips, 0 failed; selfcheck: 26/26.")
    add_bullet(doc, "A real elevated firewall apply/rollback/connectivity-loss exercise and physical mixed-network topology remain controlled operator acceptance tests, not claimed automated evidence.")
    doc.save(path)
    return path


def main() -> None:
    outputs = [
        update_security(),
        update_capability(),
        update_master(),
        update_vulnerabilities(),
    ]
    for path in outputs:
        reopened = Document(path)
        joined = "\n".join(p.text for p in reopened.paragraphs)
        assert VERSION in joined
        assert "Host Adaption" in joined
        print(path)


if __name__ == "__main__":
    main()
