"""Build the canonical Angerona Master Manual from the Analysis DOCX set.

This is intentionally a content-aware rebuild rather than a raw OOXML append:
all source documents are retained, exact repeated paragraphs/tables are included
once, styles are normalized, and the current UI/enterprise direction is added.
"""
from __future__ import annotations

import argparse
import hashlib
import re
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Iterable, Iterator

from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
ANALYSIS_DIR = REPO_ROOT / "analysis"

BLUE = "1F9CFF"
DARK_BLUE = "164E75"
INK = "0B1726"
MUTED = "64748B"
LIGHT_BLUE = "E8F4FF"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAFBF2"
PALE_GOLD = "FFF8E8"
BORDER = "CBD5E1"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_V = 80
CELL_MARGIN_H = 120


@dataclass(frozen=True)
class Source:
    filename: str
    chapter: str
    purpose: str


SOURCES = (
    Source(
        "Angerona_Capabilities_Bragsheet_v1.9.4.docx",
        "Executive Capability Summary",
        "Concise capability and release proof summary.",
    ),
    Source(
        "Angerona_Master_Manual_v1.9.4.docx",
        "Core Operator Manual",
        "Installation, operation, security model, resilience, data, and history.",
    ),
    Source(
        "Angerona_Capability_Doc_v1.9.4.docx",
        "Detailed Capability Catalog",
        "Layer, module, console, UI, and security capability reference.",
    ),
    Source(
        "Angerona_System_Flow_v1.9.4.docx",
        "System Architecture and Data Flow",
        "Architecture nodes, trust boundaries, failure behavior, and proof flow.",
    ),
    Source(
        "Angerona_Security_Assessment_v2.3_2026-07-27.docx",
        "Security Assessment",
        "Assessment scope, findings, controls, remediation, and residual risk.",
    ),
    Source(
        "Angerona_Vulnerabilities_Assessment_Remediation_v1.9.4.docx",
        "Vulnerability and Remediation Record",
        "Validated defects, remediation history, verification, and open decisions.",
    ),
)


def _font(run, name: str = "Calibri", size: float | None = None,
          color: str | None = None, bold: bool | None = None,
          italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", CELL_MARGIN_V),
        ("bottom", CELL_MARGIN_V),
        ("start", CELL_MARGIN_H),
        ("end", CELL_MARGIN_H),
    ):
        el = tc_mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _set_table_geometry(table: Table, widths: list[int]) -> None:
    """Apply fixed DXA geometry: tblW/tblInd/tblGrid/tcW all agree."""
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    tbl_pr = table._tbl.tblPr
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440.0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)


def _smart_widths(rows: list[list[str]]) -> list[int]:
    columns = max((len(row) for row in rows), default=1)
    if columns == 1:
        return [CONTENT_WIDTH_DXA]
    scores: list[float] = []
    for col in range(columns):
        texts = [row[col] if col < len(row) else "" for row in rows]
        max_len = max((min(90, len(t)) for t in texts), default=1)
        avg_len = sum(min(90, len(t)) for t in texts) / max(1, len(texts))
        scores.append(max(8.0, max_len * 0.55 + avg_len * 0.45))
    # Short ID/status/number columns stay compact.
    for col in range(columns):
        values = [row[col].strip() for row in rows[1:] if col < len(row)]
        if values and all(len(v) <= 18 for v in values):
            scores[col] = min(scores[col], 18.0)
    minimum = 900 if columns >= 4 else 1200
    remaining = CONTENT_WIDTH_DXA - minimum * columns
    if remaining < 0:
        minimum = CONTENT_WIDTH_DXA // columns
        remaining = CONTENT_WIDTH_DXA - minimum * columns
    total = sum(scores) or 1.0
    widths = [minimum + int(remaining * score / total) for score in scores]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _normalize(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip().casefold()
    return text


def _canonicalize_stale_instruction(text: str) -> str:
    """Replace historical instructions that are unsafe in a canonical manual."""
    stripped = text.strip()
    stripped = re.sub(
        r"secrets confined to a git-ignored \.env",
        (
            "credentials stored in the current-user Windows DPAPI-protected store "
            "(legacy .env is migration input only)"
        ),
        stripped,
        flags=re.IGNORECASE,
    )
    stripped = re.sub(
        (
            r"credentials stored in the current-user Windows DPAPI-protected store "
            r"\(legacy \.env is migration input only\)\s*"
            r"\(\.env, \.env\.\* ignored, \.env\.example kept\)"
        ),
        (
            "Credentials stored in the current-user Windows DPAPI-protected store "
            "(legacy .env files remain git-ignored migration input only)"
        ),
        stripped,
        flags=re.IGNORECASE,
    )
    if stripped.startswith("%LOCALAPPDATA%\\Angerona\\ (or ANGERONA_DATA override)"):
        return (
            "<install-folder>\\runtime-data (or the protected packaged-data root) — "
            "flight-recorder DB, settings, logs, diagnostics, scanner evidence, "
            "reports, and forensic case folders. %LOCALAPPDATA%\\Angerona is legacy "
            "migration input only, never a production write target."
        )
    if stripped.startswith("Secrets only in git-ignored .env"):
        return (
            "Credentials are stored in the current-user Windows DPAPI-protected store. "
            "A legacy .env file is migration input only and must not remain the active "
            "credential source."
        )
    if stripped.startswith(
        "kill-all-angerona.bat — elevated, terminates every pythonw.exe/python.exe"
    ):
        return (
            "kill-all-angerona.bat — elevated; terminates only verified Angerona-owned "
            "Python entry points and unloads resident Angerona llama3 models"
        )
    return stripped


def _table_key(rows: list[list[str]]) -> str:
    joined = "\n".join("\t".join(_normalize(cell) for cell in row) for row in rows)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()


def _iter_blocks(parent: _Document | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        parent_obj = parent
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
        parent_obj = parent
    else:
        raise TypeError(type(parent))
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent_obj)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent_obj)


def _looks_like_heading(text: str, style_name: str) -> int | None:
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except Exception:
            return 2
    if style_name in {"Title", "Subtitle"}:
        return 1
    stripped = text.strip()
    if re.match(r"^\d+(?:\.\d+)*\.?\s+\S", stripped) and len(stripped) <= 125:
        depth = min(3, stripped.split()[0].count(".") + 1)
        return depth
    if re.match(r"^[A-Z]-\d+\s+[-\u2014]", stripped) and len(stripped) <= 150:
        return 3
    if re.match(r"^v\d+\.\d+(?:\.\d+)?\b", stripped, flags=re.I) and len(stripped) <= 150:
        return 2
    if stripped.rstrip(":") in {
        "Verification",
        "Remediation steps",
        "What's new",
        "Operator summary",
        "Architecture changes",
        "Release gate",
        "Evidence reviewed",
        "Implemented corrections",
        "Validation",
        "Residual assurance",
        "Residual assurance work",
        "Residual decisions",
        "Executive conclusion",
        "Public-release guidance",
    }:
        return 3
    return None


def _configure_styles(doc: _Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for name, color, fill in (
        ("Manual Callout", DARK_BLUE, LIGHT_BLUE),
        ("Manual Note", INK, LIGHT_GRAY),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.base_style = normal
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.left_indent = Inches(0.14)
        style.paragraph_format.right_indent = Inches(0.14)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15


def _shade_paragraph(paragraph: Paragraph, fill: str, border: str = BORDER) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        node = p_bdr.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            p_bdr.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), border)


def _add_callout(doc: _Document, label: str, text: str,
                 style: str = "Manual Callout", fill: str = LIGHT_BLUE) -> Paragraph:
    paragraph = doc.add_paragraph(style=style)
    first = paragraph.add_run(f"{label}: ")
    _font(first, size=10.5, color=DARK_BLUE, bold=True)
    body = paragraph.add_run(text)
    _font(body, size=10.5, color=INK)
    _shade_paragraph(paragraph, fill)
    return paragraph


def _add_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def _configure_sections(doc: _Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run("ANGERONA  /  CONSOLIDATED MASTER MANUAL")
        _font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        run = fp.add_run("Angerona v1.9.4  |  Consolidated 2026-07-27  |  Page ")
        _font(run, size=8.5, color=MUTED)
        _add_page_field(fp)


def _add_cover(doc: _Document) -> None:
    logo = REPO_ROOT / "assets" / "icons" / "angerona_icon_256.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    if logo.exists():
        run = p.add_run()
        shape = run.add_picture(str(logo), width=Inches(1.05))
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", "Angerona shield logo")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("DEFENSIVE SECURITY REFERENCE")
    _font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("ANGERONA")
    _font(run, size=31, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Consolidated Master Manual")
    _font(run, size=18, color=DARK_BLUE, bold=False)

    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.paragraph_format.space_after = Pt(50)
    run = edition.add_run("Version 1.9.4 consolidated edition  |  27 July 2026")
    _font(run, size=10.5, color=MUTED, italic=True)

    _add_callout(
        doc,
        "Safety boundary",
        "Angerona is a defensive, local-first security suite. Its red-team "
        "simulations use benign reversible markers; they do not deploy real "
        "exploits, credential theft, persistence, ransomware, or destructive payloads.",
        fill=PALE_GREEN,
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(22)
    run = note.add_run(
        "This manual consolidates the former Bragsheet, Capability Document, "
        "Master Manual, System Flow, Security Assessment, and Vulnerability "
        "Remediation documents. The source files remain preserved as historical evidence."
    )
    _font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def _add_contents(doc: _Document) -> None:
    doc.add_heading("Contents", level=1)
    chapters = (
        "1. Current-State Guide",
        "2. Executive Capability Summary",
        "3. Core Operator Manual",
        "4. Detailed Capability Catalog",
        "5. System Architecture and Data Flow",
        "6. Security Assessment",
        "7. Vulnerability and Remediation Record",
        "8. Enterprise Upgrade Direction",
        "Appendix A. Consolidation Sources",
        "Appendix B. Authoritative Engineering References",
    )
    for item in chapters:
        p = doc.add_paragraph(style="List Number")
        # The displayed chapter numbers are part of the actual title; suppress
        # fake numbering by using a bullet-style visual list instead.
        p.style = doc.styles["List Bullet"]
        run = p.add_run(item)
        _font(run, size=11, color=INK)
    _add_callout(
        doc,
        "Reading note",
        "Historical release notes and assessment findings retain their original "
        "dates and evidence counts. The Current-State Guide and latest repository "
        "tests take precedence over stale historical instructions.",
        style="Manual Note",
        fill=PALE_GOLD,
    )
    doc.add_page_break()


def _add_current_state(doc: _Document) -> None:
    doc.add_heading("1. Current-State Guide", level=1)
    intro = doc.add_paragraph(
        "Angerona is currently strongest as a Windows-first, single-host defensive "
        "suite with a modular sensor/response architecture, local ARIA assistance, "
        "non-destructive validation drills, signed proof artifacts, and independent "
        "resilience helpers. It has enterprise-grade foundations, but it is not yet "
        "a centrally managed enterprise fleet product."
    )
    intro.paragraph_format.keep_with_next = True

    doc.add_heading("Canonical current corrections", level=2)
    corrections = (
        "Credentials are stored in a current-user DPAPI-protected store. Legacy "
        "plaintext .env import requires an explicit migration action.",
        "Runtime data resolves through the canonical configured data root; the "
        "project must not silently spill scanner, drill, report, or database data to C:.",
        "External Python capabilities require a detached manifest, exact source "
        "digest, declared permissions/privacy budget, trusted Ed25519 publisher, "
        "and signature, except explicit hash-pinned development mode.",
        "Remediation closure requires deterministic action evidence and postcondition "
        "proof. Local AI may explain or recommend; it is not an execution authority.",
        "Release checksums, SBOMs, and provenance improve traceability but do not "
        "replace publisher authentication. Authenticode/MSIX signing remains open.",
        "Fleet enrollment, organization RBAC/audit, centrally signed policy, "
        "cross-endpoint search/storage, HA, and case management remain future work.",
    )
    for text in corrections:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    doc.add_heading("Dashboard and operator experience", level=2)
    ui_items = (
        "Every top-row action now has a distinct Qt-rendered vector icon, avoiding "
        "platform-dependent emoji rendering.",
        "Hovering or keyboard-focusing a top-row action presents its name and plain-"
        "language definition, with accessible name/description metadata.",
        "Top-row clicks use a 280 ms vertical-line-to-panel reveal and ignore duplicate "
        "clicks during the transition. Windows/app reduced-motion settings disable it.",
        "The System Pulse card beside ARIA shows CPU, RAM, available memory, Wi-Fi "
        "signal, and aggregate receive/send throughput.",
        "System Pulse sampling runs in a single-flight background worker; slower "
        "Windows/Wi-Fi queries do not run on the Qt UI thread.",
        "Eco Mode continues to pause heavy background scanners and wakes them one at "
        "a time, waiting for first-cycle boundaries to avoid a resource stampede.",
    )
    for text in ui_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    doc.add_heading("Current documentation model", level=2)
    _add_callout(
        doc,
        "Canonical references",
        "Use this Master Manual for human-facing architecture, operations, "
        "capabilities, security history, and roadmap context. Use llms.txt for the "
        "dense AI/tooling contract. Use ENTERPRISE_UPGRADE_TODO.txt as the actionable "
        "engineering backlog with stable task IDs and release gates.",
    )


def _copy_runs(source: Paragraph, destination: Paragraph) -> None:
    run_text = "".join(run.text for run in source.runs)
    if run_text != source.text:
        run = destination.add_run(source.text)
        _font(run, size=11, color=INK)
        return
    for source_run in source.runs:
        run = destination.add_run(source_run.text)
        is_code = (
            source_run.font.name in {"Consolas", "Courier New", "Cascadia Mono"}
            or "`" in source_run.text
        )
        _font(
            run,
            name="Consolas" if is_code else "Calibri",
            size=9.5 if is_code else 11,
            color=INK,
            bold=source_run.bold,
            italic=source_run.italic,
        )
        run.underline = source_run.underline


def _add_picture_from_paragraph(doc: _Document, paragraph: Paragraph) -> bool:
    rel_ids = paragraph._p.xpath(".//a:blip/@r:embed")
    if not rel_ids:
        return False
    added = False
    for rel_id in rel_ids:
        part = paragraph.part.related_parts.get(rel_id)
        if part is None:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        shape = run.add_picture(BytesIO(part.blob), width=Inches(6.25))
        shape._inline.docPr.set(
            "descr",
            "Angerona system architecture and control-flow diagram",
        )
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        r = caption.add_run("Figure: Angerona system architecture and control flow")
        _font(r, size=9, color=MUTED, italic=True)
        added = True
    return added


def _add_source_table(doc: _Document, source: Table) -> list[list[str]]:
    rows = [
        [
            _canonicalize_stale_instruction(
                re.sub(r"\s+", " ", cell.text).strip()
            )
            for cell in row.cells
        ]
        for row in source.rows
    ]
    if not rows:
        return rows
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = _smart_widths(rows)
    _set_table_geometry(table, widths)
    for ri, row in enumerate(rows):
        for ci in range(columns):
            cell = table.cell(ri, ci)
            cell.text = row[ci] if ci < len(row) else ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            if ri == 0:
                _set_cell_shading(cell, LIGHT_BLUE)
            elif ri % 2 == 0:
                _set_cell_shading(cell, "F8FAFC")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.08
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if ci == 0 and columns >= 4
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in p.runs:
                    _font(
                        run,
                        size=9.2,
                        color=INK,
                        bold=(ri == 0),
                    )
    _set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return rows


def _add_source_chapter(
    doc: _Document,
    chapter_number: int,
    source_spec: Source,
    seen_paragraphs: set[str],
    seen_tables: set[str],
) -> tuple[int, int, int]:
    source_path = ANALYSIS_DIR / source_spec.filename
    source = Document(source_path)
    doc.add_page_break()
    doc.add_heading(f"{chapter_number}. {source_spec.chapter}", level=1)
    source_line = doc.add_paragraph()
    source_line.paragraph_format.space_after = Pt(10)
    r = source_line.add_run(
        f"Consolidated from {source_spec.filename}. {source_spec.purpose}"
    )
    _font(r, size=9.5, color=MUTED, italic=True)

    included_p = skipped_p = included_t = 0
    skip_companion_section = False
    current_heading_level = 1
    for block in _iter_blocks(source):
        if isinstance(block, Paragraph):
            original_text = re.sub(r"\s+", " ", block.text).strip()
            text = _canonicalize_stale_instruction(original_text)
            style_name = block.style.name if block.style is not None else "Normal"
            if _add_picture_from_paragraph(doc, block):
                continue
            if not text:
                continue
            normalized = _normalize(text)
            if (
                normalized.startswith("angerona -")
                or normalized.startswith("angerona \u2014")
                or normalized.startswith("angeronasuite (canonical gui app)")
                or normalized.startswith("document version ")
            ) and included_p < 3:
                continue
            if normalized.startswith("12. companion documents"):
                skip_companion_section = True
                continue
            heading_level = _looks_like_heading(text, style_name)
            if skip_companion_section and heading_level is None:
                continue
            if skip_companion_section and heading_level is not None:
                skip_companion_section = False

            if heading_level is None and normalized in seen_paragraphs:
                skipped_p += 1
                continue

            if heading_level is not None:
                # Source headings become subheadings within the consolidated chapter.
                requested_level = 2 if heading_level <= 1 else 3
                level = min(requested_level, current_heading_level + 1)
                destination = doc.add_heading(text, level=level)
                current_heading_level = level
            else:
                if style_name.startswith("List Bullet"):
                    destination = doc.add_paragraph(style="List Bullet")
                elif style_name.startswith("List Number"):
                    destination = doc.add_paragraph(style="List Number")
                else:
                    destination = doc.add_paragraph()
                if text != original_text:
                    destination.add_run(text)
                else:
                    _copy_runs(block, destination)
                seen_paragraphs.add(normalized)
            included_p += 1
        elif isinstance(block, Table):
            rows = [
                [
                    _canonicalize_stale_instruction(
                        re.sub(r"\s+", " ", cell.text).strip()
                    )
                    for cell in row.cells
                ]
                for row in block.rows
            ]
            key = _table_key(rows)
            if key in seen_tables:
                continue
            seen_tables.add(key)
            _add_source_table(doc, block)
            included_t += 1

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    r = note.add_run(
        f"Consolidation record: {included_p} paragraph blocks and {included_t} "
        f"tables retained; {skipped_p} exact repeated paragraph blocks were omitted."
    )
    _font(r, size=8.5, color=MUTED, italic=True)
    return included_p, skipped_p, included_t


def _add_enterprise_direction(doc: _Document) -> None:
    doc.add_page_break()
    doc.add_heading("8. Enterprise Upgrade Direction", level=1)
    doc.add_paragraph(
        "The enterprise roadmap deliberately prioritizes trust and operability over "
        "feature count. Angerona should remain a strong standalone product while an "
        "optional fleet layer is built around unique device identity, typed remote "
        "jobs, signed content, tenant isolation, durable ingestion, search, cases, "
        "high availability, and verifiable updates."
    )

    doc.add_heading("Build sequence", level=2)
    rows = [
        ["Phase", "Primary outcome", "Release meaning"],
        ["0 - Public release", "Signed, privacy-clean, reproducible, soak-tested standalone build", "Trustworthy standalone"],
        ["1 - Fleet foundation", "Enrollment, mTLS, RBAC, audit, ingestion, signed policy", "Fleet preview"],
        ["2 - Operations", "Search, hunts, cases, remote updates, backup and DR", "Enterprise beta"],
        ["3 - Advanced defense", "Isolation, cross-host analytics, broader platform depth", "GA differentiation"],
        ["4 - Research", "Federated learning, attestation, counterfactual defense", "Lab only until proven"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    widths = [1500, 4860, 3000]
    _set_table_geometry(table, widths)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = text
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                _set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    _font(run, size=9.4, color=INK, bold=(ri == 0))
    _set_repeat_table_header(table.rows[0])

    doc.add_heading("Highest-priority engineering outcomes", level=2)
    priorities = (
        "Gate A: public-release secrets/privacy sweep, publisher-authenticated "
        "installer, SBOM/provenance, update rollback, and 24-hour soak.",
        "Deterministic drill remediation: map benign markers to typed cleanup, "
        "verify postconditions, and score only proven closures.",
        "Fleet identity: unique locally generated keys, one-time enrollment, mTLS, "
        "rotation, revocation, and quarantine.",
        "Tenant and admin boundary: OIDC, least-privilege RBAC, separation of "
        "duties, and immutable admin audit.",
        "Signed policy and content: effective-policy explanation, canary rollout, "
        "health gates, last-known-good, and rollback.",
        "Typed remote operations only: no generic remote shell and no AI-generated "
        "endpoint execution.",
        "Durable bounded data plane: versioned envelope, offline queue, backpressure, "
        "deduplication, sequence-gap visibility, retention, and privacy classification.",
    )
    for text in priorities:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Innovation guardrails", level=2)
    _add_callout(
        doc,
        "Research rule",
        "Kernel drivers, autonomous investigation, federated learning, deception "
        "fabric, and self-healing remain isolated research until their exact need, "
        "threat model, privacy impact, resource budget, rollback, and independent "
        "review are proven.",
        fill=PALE_GOLD,
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "The complete task-level roadmap, stable IDs, acceptance gates, source "
        "references, and 30/90/180/365-day sequence are maintained in "
        "ENTERPRISE_UPGRADE_TODO.txt at the repository root."
    )
    _font(r, name="Consolas", size=9.5, color=DARK_BLUE)


def _add_source_appendix(doc: _Document, records: list[tuple[Source, tuple[int, int, int]]]) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A. Consolidation Sources", level=1)
    doc.add_paragraph(
        "The versioned source documents remain in the Analysis folder as historical "
        "evidence. This manual is the canonical human-facing consolidation."
    )
    rows = [["Source", "Manual chapter", "Retained", "Exact repeats omitted"]]
    for source, (included, skipped, tables) in records:
        rows.append(
            [
                source.filename,
                source.chapter,
                f"{included} paragraphs; {tables} tables",
                str(skipped),
            ]
        )
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    widths = [3300, 2640, 2160, 1260]
    _set_table_geometry(table, widths)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = text
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                _set_cell_shading(cell, LIGHT_BLUE)
            elif ri % 2 == 0:
                _set_cell_shading(cell, "F8FAFC")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    _font(run, size=8.8, color=INK, bold=(ri == 0))
    _set_repeat_table_header(table.rows[0])


def _add_reference_appendix(doc: _Document) -> None:
    doc.add_heading("Appendix B. Authoritative Engineering References", level=1)
    references = (
        ("NIST SP 800-207 Zero Trust Architecture", "https://csrc.nist.gov/pubs/sp/800/207/final"),
        ("NIST SP 800-218 Secure Software Development Framework", "https://csrc.nist.gov/pubs/sp/800/218/final"),
        ("NIST AI Risk Management Framework", "https://www.nist.gov/itl/ai-risk-management-framework"),
        ("OCSF", "https://ocsf.io/"),
        ("OASIS TAXII 2.1", "https://docs.oasis-open.org/cti/taxii/v2.1/os/taxii-v2.1-os.html"),
        ("MITRE ATT&CK Detection Strategies", "https://attack.mitre.org/detectionstrategies/"),
        ("MITRE D3FEND", "https://d3fend.mitre.org/about/"),
        ("SPIFFE Workload API", "https://spiffe.io/docs/latest/spiffe-specs/spiffe_workload_api/"),
        ("SLSA specification", "https://slsa.dev/spec/v1.2/"),
        ("Sigstore signing overview", "https://docs.sigstore.dev/cosign/signing/overview/"),
        ("The Update Framework", "https://theupdateframework.github.io/specification/"),
        ("OpenTelemetry", "https://opentelemetry.io/docs/"),
    )
    for name, url in references:
        p = doc.add_paragraph(style="List Bullet")
        name_run = p.add_run(f"{name}: ")
        _font(name_run, size=10, color=INK, bold=True)
        url_run = p.add_run(url)
        _font(url_run, name="Consolas", size=8.8, color=DARK_BLUE)


def build(output: Path) -> None:
    missing = [str(ANALYSIS_DIR / source.filename) for source in SOURCES
               if not (ANALYSIS_DIR / source.filename).exists()]
    if missing:
        raise FileNotFoundError("Missing source documents:\n" + "\n".join(missing))

    doc = Document()
    _configure_styles(doc)
    _configure_sections(doc)
    _add_cover(doc)
    _add_contents(doc)
    _add_current_state(doc)

    seen_paragraphs: set[str] = set()
    seen_tables: set[str] = set()
    records: list[tuple[Source, tuple[int, int, int]]] = []
    for number, source in enumerate(SOURCES, start=2):
        record = _add_source_chapter(
            doc,
            number,
            source,
            seen_paragraphs,
            seen_tables,
        )
        records.append((source, record))

    _add_enterprise_direction(doc)
    _add_source_appendix(doc, records)
    _add_reference_appendix(doc)

    # Make layout hints stable and keep headings with their first body paragraph.
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if paragraph.text.strip().startswith(("http://", "https://")):
            for run in paragraph.runs:
                _font(run, name="Consolas", size=8.8, color=DARK_BLUE)

    core = doc.core_properties
    core.title = "Angerona Consolidated Master Manual"
    core.subject = "Architecture, operations, capabilities, security, remediation, and roadmap"
    core.keywords = "Angerona, defensive security, endpoint, ARIA, manual"
    core.comments = "Consolidated from the versioned Angerona Analysis document set."
    core.created = datetime(2026, 7, 27)
    core.modified = datetime(2026, 7, 27)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=REPO_ROOT / "analysis" / "manual_build" / "Angerona_Master_Manual_unscrubbed.docx",
    )
    args = parser.parse_args()
    build(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
