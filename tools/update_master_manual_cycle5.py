"""Apply the verified Cycle 5 delta to the consolidated Master Manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def _find(doc: Document, text: str, style: str | None = None, last: bool = False):
    matches = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip() == text
        and (style is None or paragraph.style.name == style)
    ]
    if not matches:
        raise RuntimeError(f"paragraph not found: {style!r} {text!r}")
    return matches[-1] if last else matches[0]


def _replace_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _insert(target, text: str = "", style: str | None = None):
    paragraph = target.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def update(source: Path, destination: Path) -> None:
    doc = Document(source)

    _replace_text(
        _find(
            doc,
            "Version 1.9.4 consolidated edition  |  27 July 2026",
            "Normal",
        ),
        "Cycle 5 consolidated edition  |  28 July 2026",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    if footer.runs:
        footer.runs[0].text = (
            "Angerona  |  Cycle 5 consolidated 2026-07-28  |  Page "
        )

    contents_appendix = _find(
        doc, "Appendix A. Consolidation Sources", "List Bullet"
    )
    _insert(contents_appendix, "9. Cycle 5 Integrated Upgrade", "List Bullet")

    dashboard_docs = _find(doc, "Current documentation model", "Heading 2")
    _insert(
        dashboard_docs,
        "The Modules Running, Alerts, Critical, and Threat Level cards; Modules, "
        "Live Alerts, SOAR Queue, ARIA, Console, System Pulse; module-history "
        "alerts; and both bottom status rows now open live detail views.",
        "List Bullet",
    )
    _insert(
        dashboard_docs,
        "The accent line expands into the real destination window, where a "
        "bounded scanning header continues the motion. Reduced-motion settings "
        "remain authoritative; Self-Test and Eco remain immediate actions.",
        "List Bullet",
    )
    _insert(
        dashboard_docs,
        "Expanded views reuse existing bounded dashboard snapshots. They do not "
        "start additional host scans or response actions.",
        "List Bullet",
    )

    old_reveal = _find(
        doc,
        "Top-row clicks use a 280 ms vertical-line-to-panel reveal and ignore "
        "duplicate clicks during the transition. Windows/app reduced-motion "
        "settings disable it.",
        "List Bullet",
    )
    _replace_text(
        old_reveal,
        "Window-opening controls use a 360 ms vertical-line-to-real-window reveal "
        "and ignore duplicate clicks during the transition. Windows/app "
        "reduced-motion settings disable it.",
    )

    appendix = _find(doc, "Appendix A. Consolidation Sources", "Heading 1")
    page_break = _insert(appendix)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    _insert(appendix, "9. Cycle 5 Integrated Upgrade", "Heading 1")
    _insert(
        appendix,
        "This section records the 28 July 2026 integrated dashboard, security, "
        "performance, and enterprise-visionary pass. It describes verified "
        "current behavior; historical release evidence elsewhere in this manual "
        "remains unchanged.",
        "Normal",
    )

    _insert(appendix, "Operator-facing dashboard changes", "Heading 2")
    for item in (
        "Every primary dashboard surface now has a live drill-down view: "
        "Modules Running, Alerts, Critical, Threat Level, Modules, Live Alerts, "
        "SOAR Queue, ARIA, Console, System Pulse, module-history alerts, and the "
        "two bottom module/status rows.",
        "The animation now reveals the actual destination window and continues "
        "inside it with a restrained scanning header and live metric tiles.",
        "System Pulse retains at most 90 samples. Expanded console text and "
        "module-resource activity are bounded, and detail timers stop when their "
        "windows close.",
        "The Red Team console is resizable and screen-aware. Its configuration "
        "scrolls on smaller displays, the kill-chain wraps into two rows, and "
        "Launch/Stop remain reachable in a sticky footer at 700 by 520 pixels.",
        "Self-Test and Eco remain immediate controls. Reduced-motion and Windows "
        "animation preferences are hard overrides.",
    ):
        _insert(appendix, item, "List Bullet")

    _insert(appendix, "Security and privacy closures", "Heading 2")
    for item in (
        "External capabilities execute the exact bytes that passed manifest "
        "verification, closing the verification-to-import race.",
        "Causal-graph receipt links are labeled as unverified references until "
        "the receipt and underlying action record are checked.",
        "Receipt authenticity validation verifies the signed/hash-chained action "
        "record and its outcome instead of accepting a stored label.",
        "Incident-response redaction now covers IPv6 addresses, UNC paths, URLs, "
        "and hostnames.",
        "Secure-store temporary files use exclusive randomized creation, and "
        "malformed security booleans fail closed.",
        "ARIA webhook destinations are stored in the Windows current-user DPAPI "
        "secret store rather than plaintext settings.",
        "Shark and Red Team requests pass a bounded fail-closed safety preflight "
        "before a worker or marker starts. Run ground truth is versioned, "
        "HMAC-attested, SHA-256-chained, bounded, and atomically written.",
        "After-Action Reports and Evolution rule synthesis refuse unsigned, "
        "tampered, legacy, or over-budget drill histories.",
    ):
        _insert(appendix, item, "List Bullet")

    _insert(appendix, "Long-session performance changes", "Heading 2")
    for item in (
        "Provenance graph nodes and edges are bounded, with constant-time "
        "duplicate rejection.",
        "Remediation-ledger catch-up is incremental and attack-feed lookup scans "
        "the newest blocks first.",
        "Compatible modules share cached network-connection snapshots rather "
        "than repeating the same host walk.",
        "Mobile alert digests are bounded, MCP startup is idempotent, and MCP "
        "queue shutdown reliably wakes its worker.",
        "FlightCache batches commits. In the focused 20,000-put benchmark, "
        "runtime improved from 1.494 seconds to 0.796 seconds and commits fell "
        "from 20,000 to 157.",
    ):
        _insert(appendix, item, "List Bullet")

    _insert(appendix, "Enterprise capability drift auditor", "Heading 2")
    _insert(
        appendix,
        "The new read-only Capability Drift Auditor inspects extension source "
        "with Python's abstract syntax tree. It never imports or executes the "
        "inspected extension. It compares observed behavior with declared "
        "permissions, checks digest and entrypoint drift, flags unsafe "
        "constructs, and privacy-minimizes paths. It is an audit aid, not an "
        "automatic blocking or remediation authority.",
        "Normal",
    )

    _insert(appendix, "Integrated verification", "Heading 2")
    for item in (
        "Repository test suite: 173 passed, 1 platform skip, 0 failed.",
        "Focused fresh UI, security, performance, visionary, and drill gate: "
        "40 passed.",
        "Enterprise drill-contract regression gate: 9 passed.",
        "Headless application, dialog, drill, and module integration harness: "
        "26 passed, 0 failed.",
        "ARIA self-tests: 13 passed, 0 failed.",
        "Compile gate: 220 Python files scanned, 0 failed.",
        "Module auto-discovery: 63 modules.",
    ):
        _insert(appendix, item, "List Bullet")
    _insert(
        appendix,
        "Verification note: expected stopped, idle, optional-driver, and "
        "local-model-unavailable module outcomes printed inside the headless "
        "drill are classified by the harness and did not fail the 26-part "
        "integration gate.",
        "Manual Callout",
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle5.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
