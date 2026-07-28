"""Add the macOS Observe architecture foundation to the consolidated manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def _find(doc: Document, text: str, style: str | None = None, last: bool = False):
    matches = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip() == text
        and (style is None or paragraph.style.name == style)
    ]
    if not matches:
        raise RuntimeError(f"paragraph not found: {style!r} {text!r}")
    return matches[-1] if last else matches[0]


def _replace(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _insert(target, text: str = "", style: str | None = None):
    paragraph = target.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def _bullet(target, text: str) -> None:
    _insert(target, text, "List Bullet")


def update(source: Path, destination: Path) -> None:
    doc = Document(source)

    # Safe re-run for a previously generated edition: refresh verification
    # evidence without duplicating the chapter or Contents entry.
    if any(
        paragraph.text.strip()
        == "Cross-platform foundation edition  |  28 July 2026"
        for paragraph in doc.paragraphs
    ):
        old_repository = [
            paragraph for paragraph in doc.paragraphs
            if paragraph.text.strip()
            == "Repository pytest: 192 passed, 1 platform skip, 0 failed."
        ]
        old_focused = [
            paragraph for paragraph in doc.paragraphs
            if paragraph.text.strip()
            == "Focused macOS/platform contract: 7 passed, 0 failed."
        ]
        for paragraph in old_repository:
            _replace(
                paragraph,
                "Repository pytest: 193 passed, 1 platform skip, 0 failed.",
            )
        for paragraph in old_focused:
            _replace(
                paragraph,
                "Focused macOS/platform contract: 8 passed, 0 failed.",
            )
        old_native = (
            "native/macos contains source scaffolding for a signed host, "
            "SMAppService background lifecycle, OSSystemExtensionRequest "
            "activation, and an Endpoint Security client subscribed only to "
            "NOTIFY events."
        )
        new_native = (
            "native/macos contains source scaffolding for a signed host, "
            "SMAppService background lifecycle, FSEvents file observation that "
            "preserves rescan/overflow flags, OSSystemExtensionRequest "
            "activation, and an Endpoint Security client subscribed only to "
            "NOTIFY events."
        )
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == old_native:
                _replace(paragraph, new_native)
        destination.parent.mkdir(parents=True, exist_ok=True)
        doc.save(destination)
        return

    _replace(
        _find(doc, "Cycle 5 consolidated edition  |  28 July 2026", "Normal"),
        "Cross-platform foundation edition  |  28 July 2026",
    )
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        if footer.runs:
            footer.runs[0].text = (
                "Angerona  |  Cross-platform foundation 2026-07-28  |  Page "
            )

    appendix_contents = _find(
        doc, "Appendix A. Consolidation Sources", "List Bullet"
    )
    _insert(
        appendix_contents,
        "10. macOS Observe Architecture Foundation",
        "List Bullet",
    )

    _replace(
        _find(
            doc,
            "Angerona is currently strongest as a Windows-first, single-host "
            "defensive suite with a modular sensor/response architecture, local "
            "ARIA assistance, non-destructive validation drills, signed proof "
            "artifacts, and independent resilience helpers. It has "
            "enterprise-grade foundations, but it is not yet a centrally "
            "managed enterprise fleet product.",
            "Normal",
        ),
        "Angerona remains strongest as a Windows Protect, single-host defensive "
        "suite with modular sensors and response, local ARIA assistance, "
        "non-destructive validation drills, signed proof artifacts, and "
        "independent resilience helpers. The shared core now also supports a "
        "source-available macOS Observe preview and an explicit Linux headless "
        "sensor path. macOS native enforcement and centralized enterprise fleet "
        "management are not yet shipped.",
    )
    _replace(
        _find(
            doc,
            "Credentials are stored in a current-user DPAPI-protected store. "
            "Legacy plaintext .env import requires an explicit migration action.",
            "List Bullet",
        ),
        "Credentials use the current-user DPAPI-protected store on Windows and "
        "the current user's Keychain on macOS. Unsupported platforms fail "
        "closed. Legacy plaintext .env import requires an explicit migration "
        "action and a verified protected write.",
    )

    dashboard_heading = _find(
        doc, "Dashboard and operator experience", "Heading 2"
    )
    for item in (
        "The production Windows edition is the Protect release. The macOS "
        "edition is an Observe developer preview; the Linux edition is an "
        "explicit headless sensor path.",
        "Every module declares supported platforms and a capability mode. "
        "Unavailable sensors never start and cannot inflate protection posture.",
        "Platform collectors normalize bounded process, file, network, "
        "authentication, system, and security observations before EventBus "
        "publication.",
    ):
        _insert(dashboard_heading, item, "List Bullet")

    appendix = _find(doc, "Appendix A. Consolidation Sources", "Heading 1")
    page_break = _insert(appendix)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    _insert(appendix, "10. macOS Observe Architecture Foundation", "Heading 1")
    _insert(
        appendix,
        "This chapter records the first implemented cross-platform architecture "
        "slice. It distinguishes reusable core services from operating-system "
        "sensor and response layers and deliberately avoids representing a "
        "developer preview as production protection.",
        "Normal",
    )

    _insert(appendix, "Edition and capability boundary", "Heading 2")
    for item in (
        "Windows Protect remains the supported packaged edition, with the full "
        "ETW, WMI/CIM, AMSI, WFP, resilience, detection, and response path.",
        "macOS Observe is source-only. It shares the GUI, storage, EventBus, "
        "local AI triage, Remote Bridge, and resource governance where each "
        "capability explicitly declares macOS support.",
        "Linux remains a headless sensor path. Its BCC/eBPF module now declares "
        "Linux support explicitly instead of appearing healthy and inert on "
        "other operating systems.",
        "A capability reports one of four modes: observe, detect, protect, or "
        "respond. Unsupported capabilities remain visible as unavailable where "
        "appropriate, never start, and do not count as enabled protection.",
    ):
        _bullet(appendix, item)
    _insert(
        appendix,
        "Release truth rule: do not label the macOS preview EDR protection until "
        "an entitled native sensor is installed, healthy, signed, notarized, "
        "upgradeable, and independently validated.",
        "Manual Callout",
    )

    _insert(appendix, "Shared platform and event contracts", "Heading 2")
    for item in (
        "core/platforms.py defines one canonical Windows, macOS, and Linux "
        "vocabulary. Existing modules without an explicit declaration fail "
        "closed as Windows-only.",
        "On non-Windows hosts, ModuleManager parses the literal platform "
        "declaration from the source AST before import. Incompatible legacy "
        "files are not imported, avoiding top-level ETW, AMSI, Registry, or "
        "Windows API failures.",
        "Capability inventory rows include the current platform, supported "
        "platforms, availability, availability reason, requirements, and "
        "capability mode.",
        "core/sensor_events.py defines normalized-event schema version 1 for "
        "process, file, network, authentication, system, and security "
        "observations. Text, fields, containers, and total serialized size are "
        "bounded before an event reaches consumers.",
    ):
        _bullet(appendix, item)

    _insert(appendix, "macOS Observe sensor", "Heading 2")
    for item in (
        "The Python observer takes a baseline, then reports newly started "
        "processes and newly established network flows at a conservative cadence.",
        "Command lines and usernames are excluded by default because they often "
        "contain credentials, tokens, personal paths, or private content.",
        "Process and connection identity state is hard-bounded for long-running "
        "systems. Network enumeration runs less frequently than process "
        "enumeration to reduce steady CPU and permission pressure.",
        "The module reports observe-only health and explicitly states that "
        "Endpoint Security enforcement is not installed. Snapshot errors degrade "
        "that module without crashing the core.",
    ):
        _bullet(appendix, item)

    _insert(appendix, "Native extension trust boundary", "Heading 2")
    for item in (
        "A future native host must emit only the normalized event envelope. The "
        "Python bridge rejects oversized, malformed, stale, future-dated, "
        "replayed, unsigned, wrongly signed, non-macOS, or schema-invalid frames.",
        "The bridge signature is HMAC-SHA-256 over the exact schema version, "
        "timestamp, nonce, and event body. A bounded nonce cache prevents replay.",
        "native/macos contains source scaffolding for a signed host, SMAppService "
        "background lifecycle, FSEvents file observation that preserves "
        "rescan/overflow flags, OSSystemExtensionRequest activation, and an "
        "Endpoint Security client subscribed only to NOTIFY events.",
        "No authorization event subscription, content filter, block action, "
        "quarantine, delete, isolation, or autonomous remediation is present in "
        "the macOS preview.",
        "The later Network Extension containment phase must keep Python and local "
        "AI outside the synchronous flow-decision path and provide deterministic "
        "fallback, allowlisting, audit, rollback, and uninstall behavior.",
    ):
        _bullet(appendix, item)

    _insert(appendix, "Secret storage and privacy", "Heading 2")
    for item in (
        "Windows retains the current-user DPAPI blob and private ACL. macOS uses "
        "Security.framework generic-password operations in the current user's "
        "Keychain.",
        "The Keychain adapter does not pass secret values on a command line, "
        "write them to a temporary plaintext file, or silently fall back to a "
        "project .env file.",
        "Protected writes are read back and compared before legacy plaintext "
        "migration is allowed to delete a source file.",
        "The Observe event contract records privacy classes and excludes raw "
        "file content, usernames, and full command lines by default.",
    ):
        _bullet(appendix, item)

    _insert(appendix, "Packaging and production gates", "Heading 2")
    for item in (
        "Obtain Apple's Endpoint Security entitlement for the production "
        "Developer ID team and provisioning profile.",
        "Create the containing app and system-extension targets in Xcode with "
        "stable bundle identifiers and Team ID; placeholder identifiers in the "
        "source scaffold are not release identities.",
        "Implement authenticated native host-to-extension XPC, Keychain bridge "
        "key provisioning, activation/deactivation, upgrade, rollback, and "
        "complete uninstall.",
        "Build with the hardened runtime, sign nested code in the correct order, "
        "notarize the containing app, and staple the notarization ticket.",
        "Complete field-level privacy review and validate on supported macOS "
        "versions in VMs and on physical Apple Silicon. Measure event loss, "
        "latency, sleep/wake, resource use, upgrade, and uninstall.",
        "Only after those gates pass should distribution add a signed/notarized "
        ".app or installer and consider authorization or containment features.",
    ):
        _bullet(appendix, item)

    _insert(appendix, "Implemented verification", "Heading 2")
    for item in (
        "Repository pytest: 193 passed, 1 platform skip, 0 failed.",
        "Focused macOS/platform contract: 8 passed, 0 failed.",
        "Compile gate: 230 Python files parsed, 0 failed.",
        "Windows discovery: 64 declared modules, 62 available, 0 discovery "
        "errors. The Linux eBPF and macOS Observe sensors are correctly marked "
        "unavailable on Windows.",
        "Simulated macOS discovery: 4 explicitly compatible capabilities "
        "imported, 0 discovery errors.",
    ):
        _bullet(appendix, item)
    _insert(
        appendix,
        "Verification limit: native Xcode compilation, Apple entitlement "
        "approval, signing, notarization, and physical macOS execution were not "
        "possible on this Windows workstation and remain mandatory external "
        "release gates.",
        "Manual Note",
    )

    reference_tail = _find(
        doc, "OpenTelemetry: https://opentelemetry.io/docs/", "List Bullet"
    )
    apple_references = (
        "Apple Endpoint Security: "
        "https://developer.apple.com/documentation/endpointsecurity",
        "Apple Endpoint Security entitlement: "
        "https://developer.apple.com/documentation/BundleResources/Entitlements/"
        "com.apple.developer.endpoint-security.client",
        "Apple Network Extension content filters: "
        "https://developer.apple.com/documentation/networkextension/"
        "content-filter-providers",
        "Apple File System Events: "
        "https://developer.apple.com/documentation/coreservices/"
        "file_system_events",
        "Apple Keychain Services: "
        "https://developer.apple.com/documentation/security/keychain-services/",
        "Apple SMAppService: "
        "https://developer.apple.com/documentation/servicemanagement/smappservice",
        "Apple notarization: "
        "https://developer.apple.com/documentation/security/"
        "notarizing-macos-software-before-distribution",
    )
    cursor = reference_tail
    for item in apple_references:
        # Appending after the final reference keeps this inside Appendix B.
        cursor = doc.add_paragraph(item, style="List Bullet")

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print(
            "usage: update_master_manual_macos.py "
            "SOURCE.docx DESTINATION.docx"
        )
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
