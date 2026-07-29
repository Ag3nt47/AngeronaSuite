"""Append the verified 2026-07-29 Cycle 6 record to the Master Manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def _find(doc, text: str, style: str | None = None, last: bool = False):
    matches = [
        p for p in doc.paragraphs
        if p.text.strip() == text and (style is None or p.style.name == style)
    ]
    if not matches:
        raise RuntimeError(f"paragraph not found: {style!r} {text!r}")
    return matches[-1] if last else matches[0]


def _replace(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _insert(target, text: str = "", style: str | None = None):
    paragraph = target.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def _bullets(target, items) -> None:
    for item in items:
        _insert(target, item, "List Bullet")


def update(source: Path, destination: Path) -> None:
    doc = Document(source)

    _replace(
        _find(doc, "Cross-platform foundation edition  |  28 July 2026"),
        "Cycle 6 current-development edition  |  29 July 2026",
    )
    footer = doc.sections[0].footer.paragraphs[0]
    if footer.runs:
        footer.runs[0].text = "Angerona  |  Cycle 6 consolidated 2026-07-29  |  Page "

    contents_appendix = _find(doc, "Appendix A. Consolidation Sources", "List Bullet")
    _insert(contents_appendix, "11. Cycle 6 Enterprise Hardening Sweep", "List Bullet")

    appendix = _find(doc, "Appendix A. Consolidation Sources", "Heading 1")
    page_break = _insert(appendix)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    _insert(appendix, "11. Cycle 6 Enterprise Hardening Sweep", "Heading 1")
    _insert(
        appendix,
        "Cycle 6 ran three connected loops: enterprise design and adversarial "
        "discovery; gated security, network, performance, and UI implementation; "
        "then independent red-team, bug, performance, and integration regression. "
        "The findings below describe the verified current-development tree, not "
        "an enterprise certification or a released signed build.",
        "Normal",
    )

    _insert(appendix, "Security findings and closures", "Heading 2")
    _bullets(appendix, (
        "The Teams development authentication bypass is no longer persisted or "
        "configurable through Settings. It requires an ephemeral process opt-in, "
        "a direct loopback peer, and no forwarding headers; tunneled or forwarded "
        "requests fail closed.",
        "Event-signing and stand-down authority now use separate keys. Elevated "
        "startup establishes a protected runtime parent before key access, and "
        "unsafe or attacker-precreated key material is quarantined rather than "
        "accepted and re-labeled as trusted.",
        "The optimized telemetry cursor now verifies each persisted event HMAC. "
        "Forged or altered rows do not reach the analyst display; a canonical "
        "Ledger Integrity event reports the rejection without repeating attacker "
        "controlled content.",
        "WFP containment treats submitted plans as untrusted at the apply boundary. "
        "Targets, timestamps, 30-second-to-24-hour TTL, recovery exclusions, plan "
        "identity, and approval binding are reconstructed and verified before a "
        "privileged executor can be called.",
        "The kernel-boundary posture ledger HMAC-authenticates every record, "
        "retains a signed pruning anchor, refuses malformed or missing initialized "
        "history, and never silently overwrites a failed chain.",
    ))

    _insert(appendix, "Kernel and network roadblocks", "Heading 2")
    _insert(
        appendix,
        "The new Kernel-Boundary Posture Ledger is a read-only Windows user-mode "
        "sensor. Every five minutes it records Secure Boot, VBS/HVCI, boot debug, "
        "test-signing and integrity-check flags, Code Integrity channel health, "
        "and a bounded kernel-driver-service inventory digest. Disabled controls "
        "raise risk; unreadable controls remain unknown/degraded. Driver-set drift "
        "is evidence, not an automatic malware verdict.",
        "Normal",
    )
    _insert(
        appendix,
        "Transactional containment adds typed IP, CIDR, port, and executable "
        "targets; mandatory loopback, DNS, and DHCP recovery exclusions; expiry; "
        "dry-run previews; explicit approval; rollback receipts; and an independent "
        "verification hook. The current repository supplies the contract and proof "
        "layer. A separately privileged, signed WFP broker is still required for "
        "production enforcement; Angerona does not silently mutate the firewall.",
        "Normal",
    )

    _insert(appendix, "Telemetry continuity and long-session performance", "Heading 2")
    _bullets(appendix, (
        "A bounded, thread-safe Telemetry Coverage Accountant is subscribed to the "
        "live EventBus and exported through status.json and status.txt. It reports "
        "sequence gaps, duplicates, regressions, stale or unsequenced sensors, and "
        "sensor-cardinality evictions. No observations means unknown, never healthy.",
        "The telemetry GUI worker retains a thread-owned read-only SQLite connection "
        "and indexed rowid cursor. A 5,000-row benchmark reduced 1,000 idle polls "
        "from 2,338.7 ms to 156.6 ms: 14.9 times faster and 93.3 percent less work.",
        "Resolve Center replaced hundreds of per-alert button widgets with one "
        "shared action bar and 25-row pagination. A simulated 5,000-alert storm "
        "opened in about 0.10 seconds; forced refresh measured about 3.38 ms while "
        "all bounded-window alerts remained reachable.",
        "Kernel posture work runs on a module worker at a 300-second cadence with "
        "6-to-8-second command timeouts, a 2,048-service scan cap, and a 256-record "
        "ledger. It performs no work on the GUI thread.",
    ))

    _insert(appendix, "Verification evidence", "Heading 2")
    _bullets(appendix, (
        "Repository pytest: 223 passed, 2 intentional platform skips, 0 failed.",
        "Python source inventory and compile gate: 233 of 233 files compiled.",
        "Module discovery: 65 modules with 0 discovery errors.",
        "Headless integration self-check: 26 of 26 passed.",
        "Module drill: 51 passed; 15 stopped, optional, local-model, or platform "
        "outcomes were classified as expected environment skips.",
        "Focused kernel-boundary and telemetry-continuity regression: 11 passed.",
    ))
    _insert(
        appendix,
        "A first final self-check exposed a genuine first-run defect: the new "
        "kernel-posture self-test expected a live ledger before the first "
        "observation. It now validates an authenticated temporary ledger without "
        "changing live posture; the final rerun passed 26 of 26.",
        "Manual Callout",
    )

    _insert(appendix, "Known boundaries and remaining enterprise work", "Heading 2")
    _bullets(appendix, (
        "Angerona remains a user-mode Python/Qt application. These controls raise "
        "cost and improve evidence, but cannot defeat an already-compromised "
        "Administrator, SYSTEM principal, kernel, or trusted in-process module.",
        "The editable source launcher is structurally checked and its runtime "
        "custody is hardened, but it cannot equal a signed, administrator-owned "
        "packaged installation. Signed installer, SBOM, provenance, publisher "
        "identity, update rollback, and release soak gates remain open.",
        "The optional AngeronaSensor driver remains laboratory-only. Production "
        "kernel deployment is blocked pending independent review, HVCI/HLK and "
        "Driver Verifier evidence, signing, crash recovery, compatibility testing, "
        "and a justified user-mode visibility gap.",
        "Telemetry sequence accounting is only as complete as producer sequence "
        "metadata. Fleet-wide durable deduplication, signed policy distribution, "
        "device identity, tenant isolation, RBAC, and central control remain roadmap work.",
        "No result in this section establishes enterprise readiness, regulatory "
        "certification, tamper-proofing, or autonomous destructive authority.",
    ))

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle6.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
