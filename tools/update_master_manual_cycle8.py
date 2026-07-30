"""Append the verified Cycle 8 implementation record to the master manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def bullets(doc, values):
    for value in values:
        doc.add_paragraph(value, style="List Bullet")


def update(source: Path, destination: Path) -> None:
    doc = Document(source)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(
        "Cycle 8. Local Fleet Preview and Response Safety", level=1
    )
    doc.add_paragraph(
        "Verification date: 29 July 2026. This section records the next "
        "offline-first enterprise tranche and preserves the distinction between "
        "a locally testable preview and production enterprise deployment."
    )

    doc.add_heading("Operator experience and configuration", level=2)
    bullets(doc, (
        "Settings now has one canonical owner for every configuration area. "
        "The Mobile Integration controls moved out of Advanced Console, which "
        "continues to provide operational diagnostics rather than duplicate state.",
        "The searchable Information tab provides capability definitions, "
        "step-by-step procedures, verification criteria, privacy boundaries, and "
        "Take me there navigation to the owning setting or operational window.",
        "Explicit Close, Cancel, and native X actions use the same "
        "reduced-motion-aware reverse panel transition as opening.",
        "Ordinary startup no longer rewrites the highest-privilege scheduled "
        "task. Registration changes only after an explicit operator setting change.",
        "The elevated source launcher validates its volume through the Windows "
        "DriveInfo API and requires a ready fixed drive. This replaces a blank "
        "PowerShell path-object DriveType property that incorrectly refused a "
        "valid D: checkout. Reparse-point, required-file, protected-key-custody, "
        "and independent Python source/interpreter checks remain fail closed.",
    ))

    doc.add_heading("Authenticated local fleet preview", level=2)
    bullets(doc, (
        "The opt-in service binds only to the loopback interface and remains "
        "disabled by default. It refuses public or Local Area Network binds.",
        "Requests authenticate the complete path and query, payload digest, "
        "timestamp, and bounded nonce with a protected per-install secret. "
        "Stale, replayed, malformed, or tampered requests fail closed.",
        "The local SQLite control plane requires tenant predicates on inventory "
        "and event operations, rejects identity-key conflicts, deduplicates "
        "events, enforces quarantine/revocation state, and signs ingestion receipts.",
        "This preview does not claim production mutual Transport Layer Security "
        "(mTLS), certificate lifecycle, internet exposure, central high "
        "availability, or Single Sign-On (SSO). Those remain deployment gates.",
    ))

    doc.add_heading("Typed enterprise response boundary", level=2)
    bullets(doc, (
        "Response operations must be registered with an exact identifier, risk, "
        "argument validator, rollback description, and bounded result handler. "
        "There is no generic remote shell.",
        "Dry-run proposals validate and describe rollback but never invoke the "
        "handler. Live operations require a distinct non-requester approval; "
        "high and critical operations require two distinct approvers.",
        "Proposal expiry, target and argument binding, idempotency, conflict "
        "rejection, bounded results, and HMAC-authenticated receipts are enforced.",
    ))

    doc.add_heading("Privacy-minimized identity defense", level=2)
    bullets(doc, (
        "Bounded local authentication analytics detect password spray, "
        "distributed account targeting, repeated failures, new privileged "
        "sign-in sources, and interactive service-account use.",
        "Account and source values are HMAC-tokenized before retention. The "
        "analytics contract never accepts passwords, tokens, or credential material.",
    ))

    doc.add_heading("Privacy-minimized network behavior", level=2)
    bullets(doc, (
        "A fixed-memory Network Detection and Response (NDR) layer detects "
        "periodic beacon timing, private lateral fanout, broad external fanout, "
        "and large asymmetric uploads.",
        "Process and destination identities are HMAC-tokenized before retention. "
        "The layer does not accept or store packet payloads.",
    ))

    doc.add_heading("Durable exposure lifecycle", level=2)
    bullets(doc, (
        "Host-applicable vulnerabilities have a durable local lifecycle with "
        "ownership, deadlines, optimistic concurrency, mitigation, bounded risk "
        "acceptance, exception expiry, and due-item queries.",
        "Resolved or closed findings require a SHA-256 closure artifact identity; "
        "state changes produce authenticated audit receipts.",
    ))

    doc.add_heading("Signed plugin lifecycle and interoperability", level=2)
    bullets(doc, (
        "Reviewed capability extensions can be staged from exact verified source "
        "and manifest snapshots, revalidated before activation, versioned, "
        "revoked, and quarantined. The lifecycle manager never imports plugin code.",
        "Trust changes, source or manifest tampering, unsigned content, and "
        "entrypoint collisions fail closed before ModuleManager import.",
        "Privacy-reviewed OCSF, STIX, OTLP, and Angerona envelopes use a durable "
        "signed queue with idempotency, bounded retry backoff, and dead letters. "
        "External egress remains denied by default.",
    ))

    doc.add_heading("Enterprise governance baseline", level=2)
    bullets(doc, (
        "Repository policy now defines supported editions and honest claims, "
        "semantic and protocol compatibility, downgrade/deprecation behavior, "
        "and six accepted architecture decisions.",
        "Detection-content governance defines owner/reviewer roles, telemetry, "
        "fixtures, ATT&CK mapping, false-positive risk, privacy, signing, and "
        "promotion gates. Support operations define ownership, severity, response "
        "targets, diagnostic privacy, escalation, and disclosure.",
    ))

    doc.add_heading("Safe fleet hunts and collections", level=2)
    bullets(doc, (
        "Fleet hunts use a closed artifact catalog, explicit device/group "
        "targets, host/byte/time/expiry budgets, independent approval, and "
        "two-person approval for restricted evidence.",
        "Authenticated atomic state survives restart and fails on tampering. "
        "Terminal results are budget-checked and signed. Arbitrary commands, "
        "scripts, paths, and remote shells are forbidden.",
        "The authenticated hunt workspace stores typed queries, bounded analyst "
        "notes, findings, decisions, and immutable result/evidence references "
        "with optimistic concurrency. It has no executable cells.",
        "Sanitized notebook export redacts identities and paths, omits device "
        "tokens and raw artifacts, and excludes restricted results by default.",
        "Per-host progress and coded failures persist as authenticated records "
        "with non-regressing timestamps and byte counters. Exact host and total "
        "byte budgets apply before insertion; bounded summaries expose the "
        "latest state without retaining raw device identity.",
        "Verified result references promote idempotently into a deterministic "
        "investigating case. Evidence remains reference-only and every promoted "
        "item receives an authenticated custody chain.",
        "Safe live-response sessions bind one target, requester, exact capability "
        "set and a maximum 30-minute expiry. Read-only sessions require an "
        "independent approval; host-changing sessions require two.",
        "Only registered read-only query handlers and separately approval-gated "
        "Response Broker operations are available. Executable fields, target "
        "escape, unregistered capabilities, expiry, and duplicate execution fail "
        "closed. The privacy-minimized transcript forms an authenticated hash chain.",
        "Read-only query handlers execute outside the session-manager lock so a "
        "slow collector cannot block close or expiry. Results arriving after "
        "the session ends are recorded as discarded and are not returned.",
    ))

    doc.add_heading("Reliability and release evidence", level=2)
    bullets(doc, (
        "A short bounded atomic-replacement retry tolerates temporary antivirus "
        "sharing locks on hunt state while persistent or unrelated failures "
        "remain visible.",
        "Registered deterministic recovery drills enforce explicit retryable "
        "errors, attempt limits, delay caps, total time budgets, and "
        "content-addressed outcomes. Current fixtures cover transient database "
        "locks, collector outage, unknown errors, service restart, and dead letters.",
        "The fixed local release gate runs serial tests, bytecode compilation, "
        "Ruff, dependency audit, and documentation drift without accepting "
        "operator-supplied commands or invoking a shell.",
        "Its bounded evidence pack binds the complete source commit and source "
        "epoch to redacted summaries, command/output digests, limitations, and "
        "a canonical manifest digest. Publisher signing and long-duration "
        "physical-host soaks remain separate gates.",
        "Core Manager and peer Watchdog now use separate HMAC-authenticated "
        "atomic recovery ledgers. Dead and suspended components receive a "
        "privacy-minimized authenticated health snapshot before recovery.",
        "Exponential restart backoff, rolling crash budgets, restart deadlines, "
        "diagnostic digests, and safe mode survive supervisor restarts. A forged "
        "ledger fails closed until an authenticated manual restart resets it; "
        "stopped heartbeats and signed stand-down remain stopped.",
    ))

    doc.add_heading("Encrypted backup, restore, and audit export", level=2)
    bullets(doc, (
        "Selected canonical data-root files and consistent SQLite snapshots "
        "stream into an Advanced Encryption Standard Galois/Counter Mode "
        "(AES-GCM) authenticated archive. Filenames, paths, manifests, and "
        "payloads remain inside the encrypted stream.",
        "Backup rejects traversal, symlink or reparse-point input, an in-root "
        "archive destination, missing required data, byte overflow, wrong keys, "
        "tampering, and digest mismatch.",
        "Restore is planned separately and binds the exact archive, target, "
        "manifest, item list, requester, and expiry. It requires Angerona "
        "offline plus two distinct non-requester approvals, stages every item, "
        "and retains replaced files in a rollback scope.",
        "An interrupted replacement restores the previous file. Protected key "
        "recovery, scheduling, retention, and full disaster-recovery exercises "
        "remain operational gates.",
        "Backup policy is disabled by default and side-effect free. It computes "
        "cadence status, exact selections, destination class, a minimum verified "
        "copy floor, count/age retention, and an HMAC-authenticated deletion "
        "proposal; it never schedules work or removes an archive.",
        "Named site-loss, database-corruption, control-plane-compromise, bad "
        "policy/update, and lost-signing-key scenarios define a Recovery Point "
        "Objective (RPO), Recovery Time Objective (RTO), minimum verified copies, "
        "owner, and review date.",
        "Recovery drill evidence measures backup age and recovery duration, "
        "requires archive, manifest, service-health, and rollback verification, "
        "lists every objective violation, and authenticates the final result. "
        "Operational alternate-site and key-recovery exercises remain open.",
        "Audit export filters an exact tenant, scope, inclusive time range and "
        "record limit. Restricted fields are removed, sensitive values and "
        "actors are tokenized, free text is redacted, and truncation is explicit.",
        "Exported audit records form a cryptographic custody chain; an "
        "HMAC-authenticated manifest binds the chain head, record digest, scope, "
        "time range, privacy policy, and requester token.",
    ))

    doc.add_heading("Proof-driven drill remediation", level=2)
    bullets(doc, (
        "The After-Action Report closure counter previously had no producer, "
        "which structurally kept simulated finding closure at zero. Closure now "
        "uses a durable lifecycle rather than a display-only flag.",
        "Each unique actionable technique receives a typed action contract with "
        "exact detector and marker scope, preconditions, operator-reviewed "
        "authorization, safety checks, idempotency, verification expiry, and "
        "an exact rollback. The contract accepts no command, shell, script, "
        "code, or arbitrary executable field.",
        "Installing a Purple Remediation Guard candidate and cleaning inert "
        "markers creates an HMAC-authenticated applied receipt. Applied does "
        "not mean fixed and cannot close the source run.",
        "The VERIFIED_CLOSED state requires a fresh, technique-bound Purple "
        "Guard echo from a different run and after the action was applied. Wrong run, "
        "detector, technique, contract identifier or digest, pre-apply evidence, "
        "and tampered lifecycle state fail closed.",
        "Application, verification, and exact-policy rollback produce standalone "
        "authenticated receipts. Rollback removes only the selected technique "
        "and preserves unrelated detector candidates.",
        "Verification expires after a bounded interval and a later miss reopens "
        "the issue. Repeated occurrences are idempotent and retained within a "
        "fixed bound.",
        "The scorecard separately reports detector coverage, same-run correlated "
        "Security Orchestration, Automation, and Response (SOAR), deterministic "
        "action application, and verified closure. Closure counts unique "
        "actionable techniques; resilience and no-detector-by-design stages "
        "remain outside the denominator.",
    ))

    doc.add_heading("Engineering efficiency and verification", level=2)
    bullets(doc, (
        "Hypothesis 6.163.0 provides deterministic, database-free, synthetic, "
        "offline property fuzzing as a development-only dependency. It derives "
        "120 replayable examples per property for normalized sensor events, "
        "authenticated Inter-Process Communication frames, signed capability "
        "manifests, detection packages, fleet authentication, encrypted-backup "
        "metadata, portable restore paths, signed release envelopes, portable "
        "archive metadata, and bounded Sigma rule documents.",
        "The property contract requires complete normalization or the documented "
        "fail-closed validation error. Its first pass found and corrected "
        "malformed platform and authentication-header exception leaks, mixed-type "
        "diagnostics, silent list truncation, non-finite values, schema smuggling, "
        "Windows alternate-data-stream/device aliases, permissive signed-release "
        "metadata, archive collisions and expansion hazards, an unpinned speech "
        "redirect, and unbounded Sigma aliases and types.",
        "Update packages, verified offline speech models, and inert drill-marker "
        "archives share one pre-read validator. It enforces canonical portable "
        "paths, Unicode normalization, entry/member/expanded-size and compression "
        "ratio budgets, duplicate and case-collision rejection, supported "
        "compression, and regular-file-only content. No encrypted or special "
        "members are accepted.",
        "Signed update envelopes use exact schemas, bounded typed artifact "
        "metadata, strict Ed25519 signature encoding, and exact 32-byte trusted "
        "keys. Speech downloads remain on their approved Hypertext Transfer "
        "Protocol Secure (HTTPS) origin. Sigma YAML is byte, document, node, and "
        "depth bounded; aliases, non-plain types, and non-finite values fail "
        "closed. Current local user-mode document/archive coverage is complete. "
        "A future native Input/Output Control boundary requires a separate gate.",
        "Shared-memory ring version 2 authenticates scanner-to-core records with "
        "a dedicated protected per-install Hash-based Message Authentication "
        "Code using Secure Hash Algorithm 256 key. Schema, sensor identifier, "
        "full 64-bit sequence, and payload are bound before parsing. Wrong-key, "
        "modified, replayed, out-of-position, unknown-schema, oversized, and "
        "malformed records are discarded and counted.",
        "Known process telemetry uses strict UTF-8 JavaScript Object Notation "
        "(JSON) with exact bounded fields and finite values. A valid non-object "
        "JSON payload can no longer terminate the ring-drain thread. Rejection "
        "alerts are rate-limited and never retain rejected raw payloads.",
        "Ring authentication detects corruption, replay, and forgery without "
        "the protected key. It is not process isolation: a fully compromised "
        "scanner that can read the key can forge a valid record. Code integrity, "
        "supervision, and future native isolation remain separate controls.",
        "Long-running Graphical User Interface telemetry now bounds both sides "
        "of its queued signal. Worker and presentation queues retain the newest "
        "2,000 events, signal and render batches are capped at 250, and EventBus "
        "identity retention is limited to 4,096 identity/timestamp pairs. "
        "Observable counters distinguish retained and dropped work.",
        "Stopping the presentation flusher rejects stale completion. System "
        "Pulse ignores post-shutdown samples. Voice resolver/listener startup "
        "and full self-test execution are single-flight, and canary echo "
        "ingestion is non-blocking with a newest-64 bound.",
        "Deterministic lifecycle coverage proves non-overlapping generation "
        "restart, deferred-restart cancellation, shutdown draining, sleep and "
        "resume, closed-dialog result suppression, duplicate-action rejection, "
        "and queue saturation. A synthetic 100,000-event burst retained exactly "
        "2,000 newest records, accounted for 98,000 drops, and enqueued in "
        "216.46 milliseconds.",
        "The current local user-mode security-boundary matrix is complete. Its "
        "74-test selection covers forged endpoint identity, signed policy and "
        "update content, authenticated receipts, After-Action Reports, tenant "
        "isolation, replay and expiry, path traversal and reparse points, "
        "untrusted extensions, and model prompt/tool abuse.",
        "Initial endpoint enrollment now binds the stable device identifier to "
        "the enrolling Ed25519 public key. Identity and access state use a "
        "current-key signature, and rotation is accepted only against the "
        "caller's already-trusted device identifier and old public key. "
        "Malformed, expired, replayed, wrong-key, or modified envelopes fail "
        "closed.",
        "External plugin catalogs are untrusted exact-schema input. Capability "
        "identifiers, portable one-file entrypoints, digests, state, timestamps, "
        "and catalog size are bounded; duplicate fields, traversal, aliases, "
        "symlink or reparse roots, and out-of-root quarantine paths are refused. "
        "Activation rebinds the verified capability identity, entrypoint, and "
        "source digest. Protected launch disables the unsigned-development "
        "override even if it was inherited from the user environment.",
        "The Artificial Intelligence Security Broker now accepts only plain, "
        "finite, exact-schema JavaScript Object Notation with conclusions tied "
        "to known evidence. Typed tool requests receive a private broker "
        "Hash-based Message Authentication Code using Secure Hash Algorithm "
        "256 authorization, expire after 1 to 300 seconds, and are consumed "
        "before handler entry. Forged, modified, cross-broker, expired, or "
        "replayed calls cannot invoke a handler; results are bounded JSON. "
        "Authenticated audit receipts can be minted only for a response "
        "validated by that broker.",
        "Ruff 0.16.0 adds a fast correctness gate and found a latent ARP "
        "Watchdog return-type spelling defect, which was corrected.",
        "pytest-xdist 3.8.0 is available for isolated test groups. The complete "
        "suite remains serial because Windows Access Control List (ACL) tests "
        "intentionally mutate permissions and are not process-parallel-safe.",
        "pip-audit 2.10.1 reports no known vulnerability in installed "
        "dependencies. Public-repository scanning found no committed real "
        "credential, private key, user-profile path, database, cache, or runtime data.",
        "A development-only toolkit pins GitHub release assets and SHA-256 "
        "digests for uv 0.12.0, py-spy 0.4.2, hyperfine 1.20.0, and the official "
        "GitHub CLI 2.96.0. Bandit 1.9.4, Vulture 2.16, and pytest-timeout 2.4.0 "
        "run in a separate ignored environment and are not product runtime "
        "dependencies.",
        "The first Bandit pass exposed a weak SHA-1 alert identifier and "
        "shell-based console clearing. Both were corrected and the "
        "high-severity re-scan is clean.",
        "Telemetry Scanner process-start evidence now includes the resolved "
        "executable location, parent process name and identifier, and a bounded "
        "command line. Common inline credential, token, API-key, authorization, "
        "and bearer values are redacted before authenticated ring persistence. "
        "Only new process identifiers are enriched, preserving the lightweight "
        "recurring table-diff path. Access denied, process exited, and unavailable "
        "states are represented explicitly.",
        "Alert Detail keeps the signed JSON as its source of truth and adds an "
        "Observed evidence panel for Event, Subject, Location, Parent, Command "
        "line, and Source. The strict two-kibibyte IPC slot accepts the enriched "
        "schema and normalizes legacy frames without fabricating evidence.",
        "The destination-window coordinator now covers later top-level Angerona "
        "dialogs globally, including legacy and indirect open paths, and reverses "
        "the same transition on X or Close. ARIA Voice and Microphone settings "
        "are explicitly routed from their source button. Tooltips and pop-up "
        "menus are excluded, and reduced-motion settings remain authoritative.",
        "Final repository evidence: 517 tests passed, 2 intentional "
        "platform-dependent skips, 0 failures; Python compilation, Ruff, "
        "dependency audit, and whitespace checks passed.",
    ))

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | Cycle 8 consolidated 2026-07-29 | Page "
                )
                break
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle8.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
