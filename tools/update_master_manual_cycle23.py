"""Surgically update the canonical Angerona Master Manual for Cycle 23.

The first invocation snapshots the pre-Cycle-23 manual under ``.tmp``.  Later
invocations always rebuild from that snapshot so visual-QA iterations cannot
duplicate sections or progressively damage the canonical source layout.
"""
from __future__ import annotations

import os
from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Angerona_Master_Manual.docx"
BUILD_ROOT = ROOT / ".tmp" / "docx_cycle23"
PRISTINE = BUILD_ROOT / "Angerona_Master_Manual_pre_cycle23.docx"
STAGED = BUILD_ROOT / "Angerona_Master_Manual_cycle23.docx"


def _set_text(paragraph, text: str) -> None:
    """Replace visible text while retaining the paragraph and first-run style."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _find_exact(document: Document, text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) > 1:
        heading_matches = [
            paragraph
            for paragraph in matches
            if paragraph.style.name.startswith("Heading")
        ]
        if len(heading_matches) == 1:
            return heading_matches[0]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _insert_before(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _insert_blocks(
    document: Document,
    *,
    before: str,
    marker: str,
    blocks: tuple[tuple[str, str], ...],
) -> None:
    if any(p.text.strip() == marker for p in document.paragraphs):
        raise ValueError(f"Cycle 23 marker already exists: {marker}")
    anchor = _find_exact(document, before)
    for style, text in blocks:
        _insert_before(document, anchor, text, style)


def _update_document_control(document: Document) -> None:
    if len(document.tables) < 2:
        raise ValueError("manual is missing its document-control tables")
    _set_text(
        document.tables[0].cell(0, 0).paragraphs[0],
        "Current release: v1.10.3 - state-grade defensive hardening for SSH, "
        "audit-log continuity, untrusted physical network paths, Personal "
        "Sentinel path attestation, sanitized live activity, and governed "
        "ARIA Defense Memory; the final gate remains zero-failure.",
    )
    rows = {
        row.cells[0].text.strip(): row.cells[1]
        for row in document.tables[1].rows
        if len(row.cells) >= 2
    }
    required = {"Version", "Release state", "Source of truth"}
    if not required.issubset(rows):
        raise ValueError("manual document-control fields changed unexpectedly")
    _set_text(rows["Version"].paragraphs[0], "1.10.3")
    _set_text(
        rows["Release state"].paragraphs[0],
        "Current local release candidate; three-round Cycle 23 defensive "
        "hardening and full serial validation complete",
    )
    _set_text(
        rows["Source of truth"].paragraphs[0],
        "Current repository code and analysis/loop/cycle23 evidence as of "
        "26 August 2026",
    )

    validation_tables = [
        table
        for table in document.tables
        if table.rows
        and len(table.rows[0].cells) >= 2
        and table.rows[0].cells[0].text.strip() == "Gate"
        and table.rows[0].cells[1].text.strip() == "Final result"
    ]
    if len(validation_tables) != 1:
        raise ValueError("manual validation summary table changed unexpectedly")
    validation_rows = {
        row.cells[0].text.strip(): row.cells[1]
        for row in validation_tables[0].rows[1:]
        if len(row.cells) >= 2
    }
    current_results = {
        "Full pytest": (
            "1,465 collected across 208 files; 1,460 passed; 5 expected "
            "host-capability skips; 0 failed"
        ),
        "Compile": "321/321 product Python files",
        "Static/runtime quality": (
            "Ruff, 73 module-file imports, 71/71 discovery, 58/58 register "
            "hooks, duplicate name/code, and patch-integrity gates passed"
        ),
        "Module self-tests": (
            "50 genuine pass; 0 genuine fail; 21 explicit inactive/platform "
            "skips; EventBus pipeline passed"
        ),
        "Core/Shark": "22/22 passed",
        "Application selfcheck": "26/26 direct and 26/26 batch",
    }
    if not set(current_results).issubset(validation_rows):
        raise ValueError("manual validation summary rows changed unexpectedly")
    for field, value in current_results.items():
        _set_text(validation_rows[field].paragraphs[0], value)

    first_footer = document.sections[0].first_page_footer
    first_footer_dates = [
        paragraph
        for paragraph in first_footer.paragraphs
        if "Angerona Suite" in paragraph.text
    ]
    if len(first_footer_dates) != 1:
        raise ValueError("manual first-page footer changed unexpectedly")
    _set_text(
        first_footer_dates[0],
        "Angerona Suite  |  v1.10.3  |  26 August 2026",
    )


def _add_cycle23_content(document: Document) -> None:
    _insert_blocks(
        document,
        before="4. Installation, launch, data, and secrets",
        marker="3.2 State-grade defense-in-depth and network trust flow",
        blocks=(
            ("Heading 2", "3.2 State-grade defense-in-depth and network trust flow"),
            ("Normal", "Cycle 23 adds actor-neutral, observe-only controls around remote access, telemetry continuity, and hostile network paths. Observation, evidence, decision, authorization, mutation, and verification remain separate trust stages; no new detector can authorize a response by itself."),
            ("Code Block", "Angerona host\n  -> operator-controlled Personal Sentinel gateway/firewall\n  -> upstream or ISP router\n  -> Internet"),
            ("List Bullet", "Every active non-loopback physical Wi-Fi or Ethernet path starts untrusted. Private addressing, an operating-system network category, and a familiar wireless name are context, not identity."),
            ("List Bullet", "The Personal Sentinel component is a vendor-neutral HTTPS attestation client for one explicitly enrolled default-gateway path. It requires normal certificate validation plus a certificate pin, a fresh nonce, a pinned policy digest, direct proxy-free transport, no redirect, and stable route context before and after the exchange."),
            ("List Bullet", "A verified gateway labels only that exact path as gateway-attested. It never trusts applications, users, destination resources, firmware, the upstream router, or the Internet, and it grants no response authority."),
            ("List Bullet", "The client does not discover, configure, credential, or operate a router. The actual intermediate appliance/server, routing policy, firmware measurement, and separate monotonic witness remain deployment projects."),
        ),
    )
    _insert_blocks(
        document,
        before="6. Detection, evidence, and interoperability",
        marker="5.2 Live Defense Activity",
        blocks=(
            ("Heading 2", "5.2 Live Defense Activity"),
            ("Normal", "The dashboard now includes a bounded Live Defense Activity card for observable operational activity. It requests at most 16 recent public events, displays at most five sanitized rows, and refreshes through the existing dashboard cadence. It reads timestamp, severity, module, and public message only; Event.details, local identifiers, credentials, source code, and private model reasoning are outside the card's boundary."),
            ("Normal", "Use this surface to confirm coarse module and EventBus activity, then open Live Alerts, Forensics, or Local SOC for evidence. It is not a debugger, code tracer, evidence viewer, live-code display, or chain-of-thought window."),
        ),
    )
    _insert_blocks(
        document,
        before="7. Adversary Combat - autonomous response",
        marker="6.4 SSH Surface / Key / Tunnel Guard",
        blocks=(
            ("Heading 2", "6.4 SSH Surface / Key / Tunnel Guard"),
            ("Normal", "The cross-platform SSH guard performs bounded read-only inspection of canonical OpenSSH configuration and Include graphs, Match ambiguity, effective authentication and forwarding posture, configured per-user key sources, file and parent-directory custody, public-key fingerprints and restrictions, host-key digests, services, listeners, client forwarding options, connections, and supported authentication/tunnel evidence."),
            ("List Bullet", "Windows OpenSSH evidence uses fixed Admin and Operational channels and fixed event identities, with bounded retry, source reopen, recovery-tail inspection, and an honest retained warning after blind intervals."),
            ("List Bullet", "The guard tokenizes accounts, paths, processes, endpoints, and sources. It never reads private keys, captures credentials, probes a listener, attempts a login, removes a key, terminates a tunnel, rewrites policy, or auto-promotes its first baseline."),
            ("Heading 2", "6.5 Audit Log Integrity Guard"),
            ("Normal", "The Windows Audit Log Integrity Guard watches fixed Security, System, and Sysmon clear, policy, service, capacity, and internal-error evidence. Authenticated channel generations, record anchors, cursors, bounded retained replay, and transition-stability checks expose explicit clearing, regression, refill, retention gaps, record reuse, source blindness, and checkpoint tamper without retaining raw event XML."),
            ("List Bullet", "A gap establishes missing confidence; it cannot identify everything that happened during the missing interval or restore deleted records. An unreadable or unauthenticated checkpoint fails closed rather than creating a clean baseline."),
            ("List Bullet", "Local HMAC custody detects alteration but cannot detect replay of an older matching state pair. The protocol for a separate high-water authority is built and tested, but no independently administered server or policy-bound hardware authority is configured by default; freshness therefore remains local-authenticity-only."),
            ("Heading 2", "6.6 Zero-Trust Network Path Monitor"),
            ("Normal", "The network monitor tokenizes physical interface, wireless, DNS, DHCP, route, gateway-neighbor, profile, and connection-epoch evidence. It reports incomplete collection, concurrent/default-path anomalies, weak or unknown wireless protection, infrastructure drift, and explicit path addition while keeping every physical path untrusted by default."),
            ("List Bullet", "A newly added path is persisted only as a provisional authenticated transition. Promotion requires the exact pending path to remain active and unchanged; absence, concurrent drift, lost freshness, failed persistence, or bounded-history pressure freezes the prior authenticated comparison baseline."),
            ("List Bullet", "The monitor is observe-only: it does not discover devices, change routes, operate a VPN, reconfigure profiles, change firewall policy, or trust endpoints reached through an attested gateway."),
        ),
    )
    _insert_blocks(
        document,
        before="10. Host Adaptation, resilience, and resource governance",
        marker="9.4 Governed ARIA Defense Memory",
        blocks=(
            ("Heading 2", "9.4 Governed ARIA Defense Memory"),
            ("Normal", "ARIA now indexes a built-in defensive reference with 10 entries covering Angerona capabilities and workflow, zero trust, SSH, erased-log response, hostile network paths, Personal Sentinel, live activity, evidence-led response, actor-neutral advanced tradecraft mappings, and assurance limits. The JSON asset is read-only, strict-schema, duplicate-key rejecting, structurally bounded, stable-read checked, resource-root confined, and pinned to a canonical SHA-256 digest before retrieval."),
            ("Normal", "Local retrieval may use the complete bounded synthesized reference. Optional cloud fallback can receive only selected redacted excerpts whose exact source is angerona://defense-memory; live telemetry and arbitrary operator runbooks remain local. The memory is static advisory data, not a self-writing learning store, executable playbook, prompt override, or action authority."),
        ),
    )
    _insert_blocks(
        document,
        before="14. Performance evidence",
        marker="13.2 Cycle 23 final validation",
        blocks=(
            ("Heading 2", "13.2 Cycle 23 final validation"),
            ("Normal", "Three sequential defensive review rounds closed 15 findings: Round 1 fixed nine; Round 2 fixed five and retained one independently administered high-water authority as deferred; Round 3 found and fixed one physical-path restart-continuity defect. The final actor-neutral review reproduced no remaining Cycle 23 Critical or High issue and did not re-label the external dependency as shipped."),
            ("List Bullet", "The authoritative one-process suite collected 1,465 tests across 208 files: 1,460 passed, five expected host-capability skips, and zero failed."),
            ("List Bullet", "All 321 product Python files compiled. Ruff passed; 73 module files imported; 71 of 71 modules discovered; all 58 compatibility register hooks were valid; core and Shark self-tests passed 22 of 22."),
            ("List Bullet", "The module harness produced 50 passes, zero failures, and 21 explicit inactive/platform skips plus a passing EventBus pipeline. Direct and batch selfcheck both passed 26 of 26."),
            ("Normal", "These gates are strong project evidence, not independent certification, state-actor attribution, or proof of complete attack coverage."),
        ),
    )
    _insert_blocks(
        document,
        before="15. Security posture, limits, and external gates",
        marker="14.2 Cycle 23 measured hot-path decisions",
        blocks=(
            ("Heading 2", "14.2 Cycle 23 measured hot-path decisions"),
            ("Normal", "Round 1 avoided redundant unchanged audit-state writes (97.6% in the measured case), limited SSH command-line work to admitted clients (90.9%), and skipped rebuilding an already-untrusted immutable network snapshot (93.3%). Round 2 replaced quadratic per-user SSH token expansion with one bounded pass, improving token-heavy stress cases by 98.8% to 99.4%."),
            ("Normal", "Round 3 retained direct pending-path security logic after maximum-bound evaluation measured 2.817 to 3.113 ms and proposed alternatives were either negligible at the 64-path cap or slower in normal/stress cases. Security clarity, freshness checks, anchors, retries, and pre/post route observations were not traded for benchmark numbers."),
        ),
    )
    _insert_blocks(
        document,
        before="16. Project positioning and resume value",
        marker="15.4 Cycle 23 assurance boundary",
        blocks=(
            ("Heading 2", "15.4 Cycle 23 assurance boundary"),
            ("List Bullet", "Built now: observe-only SSH posture/runtime drift, fixed-source Windows audit-clear and continuity evidence, zero-trust physical-path drift, a pinned Personal Sentinel attestation client, sanitized dashboard activity, and pinned governed ARIA defensive reference memory."),
            ("List Bullet", "Built as contracts but not a deployed independent service: compact witness receipt verification and the IndependentHighWater interface with rollback, fork, clone, outage, migration, and crash tests."),
            ("List Bullet", "Proposed or deferred: the actual gateway/firewall appliance and routing role, a separately administered monotonic witness server or policy-bound hardware authority, firmware/signed-boot measurement, resource-scoped egress leases, and authoritative SSH key-to-session provenance."),
            ("List Bullet", "A compromised Administrator or SYSTEM principal, kernel, hypervisor, firmware, upstream provider, stolen identity, or separately administered witness can still defeat or deny user-mode evidence. Zero trust limits implicit authority; it does not make any layer infallible."),
        ),
    )
    _insert_blocks(
        document,
        before="Appendix A. Command reference",
        marker="17.5 v1.10.3 state-grade defensive hardening (2026-08-26)",
        blocks=(
            ("Heading 2", "17.5 v1.10.3 state-grade defensive hardening (2026-08-26)"),
            ("Normal", "Cycle 23 translates public advanced-operator tradecraft into actor-neutral defensive controls instead of an attribution engine. It adds SSH posture/key/tunnel observation, Windows audit-log clear and continuity evidence, hostile-by-default physical network evaluation, exact-path Personal Sentinel attestation, sanitized live defense activity, and governed ARIA Defense Memory."),
            ("List Bullet", "Fifteen findings were fixed across three review rounds. The final Round 3 path-addition repair authenticates pending membership and requires active unchanged confirmation across restart before promotion."),
            ("List Bullet", "Final automated evidence: 1,460 passed, five expected skips, zero failed from 1,465 tests across 208 files; 321 of 321 product files compiled; 71 of 71 modules discovered; Ruff and both 26 of 26 selfchecks passed."),
            ("List Bullet", "Measured applied wins range from 90.9% to 99.4% in their declared bounded cases. Round 3 intentionally retained clearer direct security transitions when alternatives were negligible or slower."),
            ("Normal", "The separately administered monotonic authority remains deferred. Without it, authenticated local state is explicitly local-authenticity-only and not independently fresh. The Personal Sentinel client attests one path; it is not an appliance, router manager, firmware verifier, endpoint trust grant, or response authority."),
        ),
    )
    _insert_blocks(
        document,
        before="Appendix C. Glossary",
        marker="B.1 Cycle 23 primary defensive sources",
        blocks=(
            ("Heading 2", "B.1 Cycle 23 primary defensive sources"),
            ("Normal", "CISA AA25-239A, network-device and SSH persistence: https://www.cisa.gov/news-events/cybersecurity-advisories/aa25-239a"),
            ("Normal", "CISA AA24-038A, critical-infrastructure persistence and telemetry guidance: https://www.cisa.gov/sites/default/files/2024-03/aa24-038a_csa_prc_state_sponsored_actors_compromise_us_critical_infrastructure_3.pdf"),
            ("Normal", "FBI/NSA 2026 router and DNS/DHCP risk advisory: https://www.ic3.gov/PSA/2026/PSA260407"),
            ("Normal", "NSA Improve Router Hygiene: https://media.defense.gov/2026/Jul/09/2003959498/-1/-1/0/CSA_IMPROVE_ROUTER_HYGIENE.PDF"),
            ("Normal", "NIST SP 800-207, Zero Trust Architecture: https://csrc.nist.gov/pubs/sp/800/207/final"),
            ("Normal", "Microsoft OpenSSH Server configuration: https://learn.microsoft.com/en-us/windows-server/administration/openssh/openssh-server-configuration"),
            ("Normal", "Microsoft OpenSSH verbose logging: https://learn.microsoft.com/en-us/troubleshoot/windows-server/system-management-components/enable-openssh-verbose-logging"),
            ("Normal", "MITRE ATT&CK T1070.001, Clear Windows Event Logs: https://attack.mitre.org/techniques/T1070/001/"),
            ("Normal", "Microsoft BadPilot campaign analysis: https://www.microsoft.com/en-us/security/blog/2025/02/12/the-badpilot-campaign-seashell-blizzard-subgroup-conducts-multiyear-global-access-operation/"),
            ("Normal", "These sources describe observable techniques and defensive priorities. A matching behavior is not proof of any agency, government, campaign, or operator."),
        ),
    )


def _update_existing_text(document: Document) -> None:
    cover_dates = [p for p in document.paragraphs if p.text.strip() == "25 August 2026"]
    if len(cover_dates) != 1:
        raise ValueError(f"expected one cover date, found {len(cover_dates)}")
    _set_text(cover_dates[0], "26 August 2026")

    contents = _find_exact(document, "Contents")
    contents.paragraph_format.page_break_before = True

    glossary = _find_exact(document, "Appendix C. Glossary")
    glossary.paragraph_format.page_break_before = True

    overview = _find_exact(
        document,
        "Angerona unifies Windows endpoint and network visibility, local evidence retention, case work, MITRE ATT&CK mapping, governed containment, recovery, local AI assistance, and non-destructive purple-team validation. Its 68 discovered modules share a typed EventBus and explicit platform contracts.",
    )
    _set_text(
        overview,
        "Angerona unifies Windows endpoint and network visibility, local evidence retention, case work, MITRE ATT&CK mapping, governed containment, recovery, local AI assistance, and non-destructive purple-team validation. Its 71 discovered modules share a typed EventBus and explicit platform contracts.",
    )

    skips = _find_exact(
        document,
        "The three skipped pytest cases are intentional platform gates. Expected inactive module results identify stopped sensors, idle/unarmed SOAR or Combat, and unavailable optional Ollama; they are not hidden failures.",
    )
    _set_text(
        skips,
        "The five skipped pytest cases are expected host-capability gates. Expected inactive module results identify stopped sensors, idle or unarmed SOAR or Combat, unavailable optional Ollama, and unsupported platform features; they are not hidden failures.",
    )

    evidence = _find_exact(
        document,
        "Repository evidence: analysis/loop/LOOP_LOG.md; innovation_ideas.md; round1-round7 findings, remediation, bug-test, and performance summaries; cycle7/round3 validation; final Combat validation artifacts; and the frozen source/tests.",
    )
    _set_text(
        evidence,
        "Repository evidence: analysis/loop/LOOP_LOG.md; analysis/loop/cycle23 research and round1-round3 findings, remediation, bug-test, performance, and visionary summaries; final Combat validation artifacts; and the current source/tests.",
    )


def update() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    BUILD_ROOT.mkdir(parents=True, exist_ok=True)
    if not PRISTINE.exists():
        probe = Document(SOURCE)
        if any(
            p.text.strip() == "17.5 v1.10.3 state-grade defensive hardening (2026-08-26)"
            for p in probe.paragraphs
        ):
            raise ValueError("cannot create a pristine snapshot from an already updated manual")
        shutil.copy2(SOURCE, PRISTINE)

    document = Document(PRISTINE)
    _update_existing_text(document)
    _update_document_control(document)
    _add_cycle23_content(document)

    visible = "\n".join(p.text for p in document.paragraphs)
    required = (
        "3.2 State-grade defense-in-depth and network trust flow",
        "5.2 Live Defense Activity",
        "6.4 SSH Surface / Key / Tunnel Guard",
        "6.5 Audit Log Integrity Guard",
        "6.6 Zero-Trust Network Path Monitor",
        "9.4 Governed ARIA Defense Memory",
        "13.2 Cycle 23 final validation",
        "14.2 Cycle 23 measured hot-path decisions",
        "15.4 Cycle 23 assurance boundary",
        "17.5 v1.10.3 state-grade defensive hardening (2026-08-26)",
        "B.1 Cycle 23 primary defensive sources",
        "1,460 passed",
        "local-authenticity-only",
    )
    missing = [item for item in required if item not in visible]
    if missing:
        raise ValueError(f"updated manual is missing required content: {missing}")

    document.core_properties.title = "Angerona Master Manual"
    document.core_properties.subject = (
        "v1.10.3 operator reference with Cycle 23 state-grade defensive hardening"
    )
    document.core_properties.comments = (
        "Canonical manual updated 26 August 2026; actor-neutral defensive evidence."
    )
    document.save(STAGED)

    reopened = Document(STAGED)
    headings = [
        p.text.strip()
        for p in reopened.paragraphs
        if p.style.name.startswith("Heading")
    ]
    for heading in required[:10]:
        if headings.count(heading) != 1:
            raise ValueError(f"expected one updated heading after reopen: {heading}")
    if reopened.paragraphs[11].text.strip() != "26 August 2026":
        raise ValueError("cover date failed reopen validation")
    os.replace(STAGED, SOURCE)
    print(f"updated {SOURCE}")


if __name__ == "__main__":
    update()
