"""Append verified Cycles 12-13 fleet transport and admission evidence."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def bullets(document: Document, values: tuple[str, ...]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _insert_after(document: Document, paragraph, text: str):
    created = document.add_paragraph(text, style=paragraph.style)
    paragraph._p.addnext(created._p)
    return created


def _remove_duplicate_page_breaks(document: Document) -> None:
    previous_was_break = False
    for paragraph in list(document.paragraphs):
        xml = paragraph._p.xml
        is_break_only = not paragraph.text.strip() and 'w:type="page"' in xml
        if is_break_only and previous_was_break:
            paragraph._element.getparent().remove(paragraph._element)
            continue
        previous_was_break = is_break_only


def update(source: Path, destination: Path) -> None:
    document = Document(source)
    if any(
        paragraph.text.strip().startswith("Cycles 12-13. Compressed Delivery")
        for paragraph in document.paragraphs
    ):
        raise ValueError("Cycles 12-13 are already present in the source manual")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Cycle 11 current-development edition"):
            paragraph.text = "Cycle 13 current-development edition  |  1 August 2026"
        elif text == "16. Cycle 11 Atomic Fleet Delivery":
            _insert_after(
                document,
                paragraph,
                "17. Cycles 12-13 Compressed Delivery and Admission Control",
            )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "Cycles 12-13. Compressed Delivery and Admission Control",
        level=1,
    )
    document.add_paragraph(
        "Verification date: 1 August 2026. These cycles add bounded request "
        "compression, explicit transport negotiation, and local overload "
        "admission to the authenticated loopback Fleet Preview. They reduce "
        "bandwidth and event-storm risk without claiming production transport, "
        "distributed quota enforcement, or high availability."
    )

    document.add_heading("Authenticated bounded compression", level=2)
    bullets(document, (
        "The request Hash-based Message Authentication Code using Secure Hash "
        "Algorithm 256 covers the exact wire representation. A request is not "
        "decoded until freshness, signature, and one-time nonce checks pass.",
        "Identity and gzip are the only supported request encodings. A gzip body "
        "must contain exactly one complete member. Malformed or incomplete "
        "streams, trailing bytes, and concatenated members fail closed.",
        "Wire and decoded bodies are independently capped at 5 MiB. Expansion "
        "beyond the decoded budget fails before JavaScript Object Notation "
        "parsing. The 256 KiB per-event and 4 MiB aggregate normalized batch "
        "limits remain separate post-decoding controls.",
        "The decoder streams into a fixed output budget; it does not use an "
        "unbounded convenience decompressor on authenticated request data.",
    ))

    document.add_heading("Capability negotiation", level=2)
    bullets(document, (
        "The deterministic OpenAPI 3.1 contract is version 1.2.0. It documents "
        "identity/gzip support and the wire, decoded, normalized, batch-item, "
        "compression-threshold, and default admission limits.",
        "An authenticated ingestion-capabilities route returns the same bounded "
        "transport contract so endpoint software can choose compression and "
        "batch size without sending probe failures.",
        "The capability response contains no tenant, endpoint, host, user, local "
        "path, event payload, credential, nonce, signature, or key identity.",
    ))

    document.add_heading("Per-tenant and per-device admission", level=2)
    bullets(document, (
        "Thread-safe token buckets enforce default tenant and endpoint event "
        "rates. Bucket count is hard-bounded and refill uses a monotonic clock.",
        "Admission runs only after tenant authorization and proof that every "
        "named device is enrolled and active, but before event lookup or insert. "
        "Missing, quarantined, revoked, or retired identities cannot allocate "
        "rate-limit state.",
        "A denied batch raises an explicit bounded retry delay and rolls the "
        "database transaction back before evidence, last-seen state, or durable "
        "ingestion counters change. The HTTP boundary returns status 429 and a "
        "fixed numeric Retry-After header.",
        "Admission attempts are process-local operational metrics and reset at "
        "restart. Stored, duplicate, clock-quality, and batch counters remain "
        "durable and semantically separate.",
    ))

    document.add_heading("Verification and remaining gates", level=2)
    bullets(document, (
        "The authoritative serial Windows repository suite passes 586 tests with "
        "2 intentional platform-dependent skips and 0 failures.",
        "Ruff correctness checks, Python bytecode compilation, whitespace "
        "validation, and focused Bandit medium/high scans pass.",
        "Adversarial tests cover valid compressed batches, gzip bombs, malformed "
        "and ambiguous streams, unsupported encoding, capability fidelity, "
        "tenant/device bucket isolation, refill, atomic denial, health metrics, "
        "and HTTP retry semantics.",
        "Dynamic endpoint backpressure, distributed quotas, production mutual "
        "Transport Layer Security, multi-node high availability, storm/full-disk "
        "physical-host drills, and independent assessment remain separate gates.",
    ))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | Cycle 13 consolidated 2026-08-01 | Page "
                )
                break
    _remove_duplicate_page_breaks(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle13.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
