"""Append the Local SOC Flow Dashboard operator addendum to the master manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def _insert_after(document: Document, paragraph, text: str, style: str) -> object:
    created = document.add_paragraph(text, style=style)
    paragraph._p.addnext(created._p)
    return created


def _insert_before(paragraph, text: str, style: str) -> object:
    created = paragraph.insert_paragraph_before(text)
    created.style = style
    return created


def _bullets(document: Document, values: tuple[str, ...]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def update(source: Path, destination: Path) -> None:
    document = Document(source)
    if any(
        paragraph.text.strip() == "21. Flow Dashboard and Local SOC Operations"
        for paragraph in document.paragraphs
    ):
        raise ValueError("Flow Dashboard addendum is already present")

    # Keep the static contents and the current operator guide authoritative.
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text == "19. Cycle 19 Setup, Supply Assurance, and Runtime Hardening":
            cursor = _insert_after(
                document,
                paragraph,
                "20. Device Security Lab and Scan Center Addendum",
                "List Bullet",
            )
            _insert_after(
                document,
                cursor,
                "21. Flow Dashboard and Local SOC Operations",
                "List Bullet",
            )
            break

    paragraphs = list(document.paragraphs)
    operator_heading = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 2"
        and paragraph.text.strip() == "Dashboard and operator experience"
    )
    index = next(
        offset for offset, paragraph in enumerate(paragraphs)
        if paragraph._p is operator_heading._p
    )
    next_heading = next(
        paragraph for paragraph in paragraphs[index + 1:]
        if paragraph.style.name.startswith("Heading")
    )
    for item in (
        "Settings and Full Setup offer a persistent Classic or Flow startup "
        "dashboard choice. The cyan Local SOC header action opens Flow directly, "
        "and Classic remains available without migration or data duplication.",
        "Flow uses five interactive radial infographics for case flow, normalized "
        "evidence, audit integrity, asset inventory, and trusted detection content. "
        "Hovering expands plain-language context; clicking opens the related tab.",
        "The Flow workspace provides Overview, Cases, Hunt, Assets, Detection "
        "Content, and Audit tabs plus direct jumps to scanning, forensics, the Red "
        "Team console, and the Classic dashboard.",
        "Hunts and inventory collection execute outside the Qt interface thread. "
        "Case evidence counts use one aggregate query, and decorative flow motion "
        "stops when the workspace is closed or reduced motion is active.",
    ):
        _insert_before(next_heading, item, "List Bullet")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("21. Flow Dashboard and Local SOC Operations", level=1)
    document.add_paragraph(
        "Implementation date: 21 August 2026. This addendum defines Angerona's "
        "second operator workspace: a local investigation and detection-content "
        "console layered over the existing bounded evidence, case, inventory, "
        "audit, and detection-package services. It improves workflow integration "
        "without adding cloud dependence or a general remote administration path."
    )

    document.add_heading("Dashboard selection and navigation", level=2)
    _bullets(document, (
        "The operator selects Classic dashboard or Flow Dashboard - Local SOC in "
        "Settings under Appearance or in Full Setup. An invalid persisted value "
        "fails safely to Classic.",
        "Classic remains the real-time defensive cockpit for module status, live "
        "alerts, SOAR, scanning, ARIA, resource telemetry, and response controls.",
        "Flow is the investigation cockpit. Five animated rings summarize cases, "
        "normalized evidence, tamper-evident audit, local assets, and active trusted "
        "detection content. Each ring is keyboard/tooltip described and opens its "
        "corresponding workflow.",
        "Selecting Flow as the startup preference opens it after the main window is "
        "ready. The Classic dashboard stays available behind it and can be restored "
        "from Flow with one action.",
    ))

    document.add_heading("Case and evidence workflow", level=2)
    _bullets(document, (
        "Cases support bounded title, assignee, tags, investigation state, legal "
        "hold, attributed comments, optimistic versioning, retention, and sanitized "
        "export.",
        "The case queue reads all evidence counts with one aggregate database query "
        "instead of issuing one query per visible row.",
        "A structured hunt accepts only supported fields and operators. Time range "
        "is bounded to one year, candidates remain bounded by the evidence store, "
        "and no SQL, script, or arbitrary expression is accepted.",
        "Attaching a hunt result stores only a normalized evidence reference, size, "
        "digest, schema provenance, collection time, and privacy class. Raw event "
        "content is not copied into the case database.",
        "Evidence custody is HMAC-authenticated and hash chained. The detail view "
        "shows custody verification separately from investigation status.",
    ))

    document.add_heading("Assets, detection content, and administrator audit", level=2)
    _bullets(document, (
        "Local inventory reports the operating-system family, release, architecture, "
        "Angerona version/module counts, and Python runtime component names and "
        "versions. It excludes hostnames, usernames, home paths, process command "
        "lines, and network identity.",
        "Detection packages are immutable and content addressed. Invalid packages "
        "enter quarantine; activation requires a currently trusted Ed25519 publisher "
        "signature and revalidates trust. Atomic rollback retains the previous digest.",
        "Case, evidence, hunt, inventory, detection, and export operations write to a "
        "local append-only administrator ledger protected by chained HMACs and "
        "database triggers that reject update and delete.",
        "Case-custody and administrator-audit keys are independently domain derived "
        "from a 256-bit Local SOC master key held in the operating-system protected "
        "credential store.",
    ))

    document.add_heading("Safety, privacy, and performance boundaries", level=2)
    _bullets(document, (
        "Flow does not require cloud access, add a remote shell, enable arbitrary "
        "queries, execute model output, activate unsigned detections, or copy raw "
        "evidence into case storage.",
        "Hunts and inventory collection use short-lived background workers with a "
        "single-flight task guard. Qt widgets are updated only through a signal on "
        "the interface thread.",
        "The radial connector animation is low frequency, pauses while hidden, and "
        "honors both the Angerona motion preference and operating-system reduced "
        "motion controls.",
        "This local operations layer closes a workflow-integration gap. It does not "
        "claim production multi-tenant SSO, high availability, external audit "
        "retention, deep packet-capture infrastructure, or signed kernel enforcement.",
    ))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | v1.10 Local SOC consolidated 2026-08-21 | Page "
                )
                break
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_flow_dashboard.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
