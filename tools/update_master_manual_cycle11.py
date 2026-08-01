"""Append the verified Cycle 11 atomic fleet-delivery record."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def bullets(document: Document, values: tuple[str, ...]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _insert_after(document: Document, paragraph, text: str):
    created = document.add_paragraph(text, style=paragraph.style)
    paragraph._p.addnext(created._p)
    return created


def _remove_duplicate_page_breaks(document: Document) -> None:
    previous_was_break = False
    for paragraph in list(document.paragraphs):
        xml = paragraph._p.xml
        is_break_only = not paragraph.text.strip() and 'w:type="page"' in xml
        if is_break_only and previous_was_break:
            paragraph._element.getparent().remove(paragraph._element)
            continue
        previous_was_break = is_break_only


def update(source: Path, destination: Path) -> None:
    document = Document(source)
    if any(
        paragraph.text.strip().startswith("Cycle 11. Atomic Fleet Delivery")
        for paragraph in document.paragraphs
    ):
        raise ValueError("Cycle 11 is already present in the source manual")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Cycle 10 current-development edition"):
            paragraph.text = "Cycle 11 current-development edition  |  1 August 2026"
        elif text == "15. Cycle 10 Fleet Integrity and Access Governance":
            _insert_after(
                document,
                paragraph,
                "16. Cycle 11 Atomic Fleet Delivery",
            )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("Cycle 11. Atomic Fleet Delivery", level=1)
    document.add_paragraph(
        "Verification date: 1 August 2026. This section records a bounded, "
        "all-or-nothing fleet-ingestion path intended to reduce local database "
        "overhead while preserving tenant, device, evidence, privacy, and replay "
        "boundaries. It is a tested loopback Fleet Preview capability, not a "
        "claim of production remote ingestion."
    )

    document.add_heading("Atomic bounded batch contract", level=2)
    bullets(document, (
        "The canonical fleet write path accepts one to 256 event envelopes. "
        "Single-event ingestion calls the same path with one item, preventing "
        "validation and retry semantics from diverging between interfaces.",
        "Each event retains a 256 KiB normalized body limit. A batch has a 4 MiB "
        "aggregate normalized-payload limit. JavaScript Object Notation depth, "
        "node, container, number, key, and UTF-8 limits remain enforced before "
        "a database write lock is acquired.",
        "Envelope fields are exact: device identifier, event identifier, body, "
        "and optional observed time. Missing fields, unknown fields, invalid "
        "identifiers, invalid time, custom objects, and oversized input fail "
        "without truncation.",
        "Every named endpoint must be enrolled and active. Missing, revoked, "
        "quarantined, or retired state rolls back the complete batch, including "
        "event rows, device liveness, and health counters.",
    ))

    document.add_heading("Performance and retry behavior", level=2)
    bullets(document, (
        "A batch uses one immediate SQLite transaction, looks up state only for "
        "the distinct endpoints it names, updates each touched endpoint once, "
        "and writes health counters in one aggregate statement.",
        "Exact at-least-once retries return signed duplicate receipts with the "
        "original observed time, server-received time, clock classification, "
        "skew, endpoint binding, and event digest.",
        "An event identifier reused by another endpoint or with changed evidence "
        "fails closed. If that conflict occurs after valid items in the batch, "
        "the earlier provisional writes are rolled back as well.",
        "Batch membership does not assert total ordering between endpoints or "
        "sensors. Signed receipt time and clock-quality evidence remain separate "
        "from endpoint-observed chronology.",
    ))

    document.add_heading("Application programming interface and health", level=2)
    bullets(document, (
        "The deterministic OpenAPI 3.1 contract is version 1.1.0 and documents "
        "the authenticated event-batch route, exact event schema, and maximum "
        "item count.",
        "The route retains loopback-only binding, complete path/query/body "
        "Hash-based Message Authentication Code using Secure Hash Algorithm 256, "
        "fresh timestamps, durable one-time nonces, JavaScript Object Notation "
        "media type, and the existing request-size boundary.",
        "Durable low-cardinality health adds accepted batches, attempted events, "
        "and largest accepted batch. It does not retain event payloads, hostnames, "
        "user names, paths, endpoint identifiers, or keys.",
    ))

    document.add_heading("Verification and remaining gates", level=2)
    bullets(document, (
        "The authoritative serial Windows repository suite passes 580 tests with "
        "2 intentional platform-dependent skips and 0 failures.",
        "Ruff correctness checks, Python bytecode compilation, whitespace "
        "validation, and a focused Bandit medium/high scan pass.",
        "Regression coverage proves one-transaction behavior, complete rollback, "
        "signed receipts, exact retry, count/field/byte rejection, authenticated "
        "HTTP use, contract fidelity, and migration-compatible health.",
        "Compression and adaptive flow control, production mutual Transport Layer "
        "Security, remote rate limits, high availability, long physical-host "
        "soaks, and independent penetration testing remain separate gates.",
    ))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | Cycle 11 consolidated 2026-08-01 | Page "
                )
                break
    _remove_duplicate_page_breaks(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle11.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
