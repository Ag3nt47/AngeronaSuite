"""Build the current-state Angerona capabilities bragsheet.

The document intentionally contains no release history.  Its module catalogue
is read directly from BaseModule class metadata so the public capability count
and descriptions stay tied to the current source tree.
"""
from __future__ import annotations

import ast
import shutil
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Angerona_Capabilities_Bragsheet.docx"
ANALYSIS_OUTPUT = ROOT / "analysis" / "Angerona_Capabilities_Bragsheet.docx"
MODULES = ROOT / "src" / "angerona" / "modules"

NAVY = "081B2C"
BLUE = "0B6E99"
CYAN = "20B8CD"
PALE = "EAF5F8"
LIGHT = "F3F6F8"
MID = "D4E3E9"
INK = "17242E"
MUTED = "4D626E"
WHITE = "FFFFFF"
GREEN = "18794E"
AMBER = "9A5B00"


def _literal(node: ast.AST) -> str:
    try:
        value = ast.literal_eval(node)
    except (ValueError, TypeError, SyntaxError, MemoryError, RecursionError):
        return ""
    return value.strip() if isinstance(value, str) else ""


def module_catalogue() -> list[tuple[str, str, str]]:
    """Return one (name, category, description) row per BaseModule class."""
    rows: list[tuple[str, str, str]] = []
    for path in sorted(MODULES.glob("*.py")):
        tree = ast.parse(path.read_text(encoding="utf-8-sig"), filename=str(path))
        for node in tree.body:
            if not isinstance(node, ast.ClassDef):
                continue
            bases = {
                base.id if isinstance(base, ast.Name) else base.attr
                for base in node.bases
                if isinstance(base, (ast.Name, ast.Attribute))
            }
            if "BaseModule" not in bases:
                continue
            fields: dict[str, str] = {}
            for item in node.body:
                if not isinstance(item, (ast.Assign, ast.AnnAssign)):
                    continue
                targets = item.targets if isinstance(item, ast.Assign) else [item.target]
                value = item.value
                for target in targets:
                    if isinstance(target, ast.Name) and target.id in {
                        "name", "category", "description"
                    }:
                        fields[target.id] = _literal(value)
            name = fields.get("name") or node.name
            category = fields.get("category") or "Other"
            description = fields.get("description") or "Modular Angerona security capability."
            rows.append((name, category, description))
    rows.sort(key=lambda row: (row[1].casefold(), row[0].casefold()))
    if len(rows) != 67:
        raise RuntimeError(f"expected 67 BaseModule capabilities, found {len(rows)}")
    if len({name for name, _, _ in rows}) != len(rows):
        raise RuntimeError("duplicate module names in capability catalogue")
    return rows


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, top=70, start=95, bottom=70, end=95) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _set_cell_text(cell, text: str, *, bold=False, color=INK, size=8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    _repeat_header(header)
    _no_split(header)
    for index, label in enumerate(headers):
        _shade(header.cells[index], NAVY)
        _set_cell_text(header.cells[index], label, bold=True, color=WHITE, size=8.5)
        if widths:
            header.cells[index].width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        _no_split(row)
        fill = WHITE if row_index % 2 == 0 else LIGHT
        for index, value in enumerate(values):
            _shade(row.cells[index], fill)
            _set_cell_text(row.cells[index], str(value), bold=index == 0, size=8.2)
            if widths:
                row.cells[index].width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ANGERONA  •  v1.10.1  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.04

    for level, size, color in ((1, 18, NAVY), (2, 12.5, BLUE), (3, 10.5, CYAN)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    if "Eyebrow" not in doc.styles:
        style = doc.styles.add_style("Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(9)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(CYAN)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "ANGERONA  /  CURRENT CAPABILITIES"
    header.style = doc.styles["Eyebrow"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_number(section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "Angerona Capabilities Bragsheet"
    props.subject = "Current capabilities, operating model, use cases, and verification status"
    props.author = "Angerona Project"
    props.last_modified_by = "Angerona Project"
    props.keywords = "Angerona, EDR, NDR, SOAR, cybersecurity, defensive security"
    props.comments = "Current-state document; no release history."


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Eyebrow")
    p.add_run("LOCAL-FIRST DEFENSIVE SECURITY PLATFORM")
    p.paragraph_format.space_before = Pt(44)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("ANGERONA")
    run.font.name = "Aptos Display"
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run("Capabilities Bragsheet")
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(18)
    run = lead.add_run(
        "A modular, local-first endpoint defense and security-operations workbench "
        "for real-time monitoring, correlation, autonomous containment, deception, "
        "forensics, AI-assisted analysis, and repeatable adversary validation."
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    add_table(
        doc,
        ["CURRENT EDITION", "DISCOVERED CAPABILITIES", "AUTOMATED TESTS", "RED-TEAM PROOF"],
        [["v1.10.1", "67 modules / 0 errors", "1,083 pass / 3 skip / 0 fail", "100% detection + response"]],
        [1.3, 1.8, 2.0, 1.7],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.add_run("Current status: ").bold = True
    p.add_run(
        "Windows is the primary defended endpoint. Linux and macOS provide "
        "observe-mode sensor paths. Angerona runs elevated in user mode and uses "
        "supported OS telemetry interfaces; no production custom kernel driver is required."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.add_run("Edition date: 24 August 2026")
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    modules = module_catalogue()
    categories = Counter(category for _, category, _ in modules)
    doc = Document()
    configure(doc)
    add_cover(doc)

    add_heading(doc, "1. What Angerona Is")
    doc.add_paragraph(
        "Angerona combines endpoint detection and response (EDR), network detection "
        "and response (NDR), security orchestration and automated response (SOAR), "
        "deception, forensics, resilience, and local AI in one operator-controlled desktop suite."
    )
    add_bullets(doc, [
        "Local-first by design: event storage, investigations, most analysis, and the default AI path stay on the host.",
        "Real-time defensive action: authenticated detector evidence can trigger immediate, exact-target containment without an incident prompt.",
        "Modular architecture: 67 BaseModule capabilities are auto-discovered and supervised through a shared signed EventBus.",
        "Evidence-led operations: detections, response actions, verification results, and reversible receipts are recorded for review.",
        "Practice-to-proof loop: Shark Attack runs benign simulations, correlates detector and response evidence, and produces after-action reports.",
        "Operator-adjustable posture: response intensity, severity threshold, process action, isolation, quarantine, blocking, and honeypots are configurable.",
    ])

    add_heading(doc, "2. Best-Fit Use Cases")
    add_table(doc, ["Use case", "What Angerona provides", "Best fit"], [
        ["Security home lab", "A unified target for detection engineering, telemetry study, attack simulation, response automation, and evidence review.", "Advanced learners, builders, defenders"],
        ["Blue-team workstation", "Local monitoring, correlation, threat hunting, containment, deception, forensics, and daily security briefings.", "Analysts and detection engineers"],
        ["Purple-team range", "Benign ATT&CK-aligned drills with exact detector/action correlation, closure evidence, and repeatable validation.", "Purple teams and training labs"],
        ["Small environment defense", "Single-host protection plus optional SIEM, mobile, and secure remote-telemetry bridges.", "Labs and small controlled deployments"],
        ["Security engineering portfolio", "A substantial Python/Qt system demonstrating secure architecture, endpoint telemetry, automation, testing, packaging, and documentation.", "Employment and technical review"],
        ["Research platform", "Extensible modules, normalized events, local AI, causal evidence graphs, and a deterministic test harness.", "Prototype and defensive R&D"],
    ], [1.35, 3.7, 1.55])

    add_heading(doc, "3. Platform Capability Map")
    add_table(doc, ["Capability area", "Current implementation"], [
        ["Endpoint telemetry", "Processes, files, memory, persistence, Defender, AMSI, Sysmon, ETW, WFP posture, recovery tamper, USB, and optional kernel/eBPF paths."],
        ["Network defense", "Connection monitoring, PID/port attribution, packet inspection, DNS anomaly analysis, C2 cadence detection, ARP and WLAN monitoring, exact-IOC blocking, and isolation."],
        ["Autonomous response", "Standing-authority Adversary Combat plus SOAR playbooks for process containment, file quarantine, network blocking, host isolation, and deception activation."],
        ["Detection engineering", "YARA, deterministic IOC fast paths, evidence fusion, ATT&CK mapping, learned baselines, Purple Remediation Guard, and evolution from verified misses."],
        ["Investigation and proof", "Tamper-evident flight recorder, causal incident graph, process provenance, forensics capture, signed reports, action receipts, and after-action reports."],
        ["Local AI and ARIA", "Ollama-based triage, assistant tools, runbook retrieval, briefings, voice options, research connectors, model integrity, and constrained patch staging."],
        ["Resilience", "Module supervision, watchdogs, anti-suspension heartbeat, telemetry canaries, chaos probes, self-integrity checks, bounded caches, and adaptive resource governors."],
        ["Operations", "Desktop dashboard, command console, ATT&CK coverage, threat intelligence, settings, readiness assessment, diagnostics, self-tests, installers, and release verification."],
        ["Integration", "CEF/Syslog SIEM forwarding, encrypted remote telemetry, mobile response, local MCP surface, reports, status exports, and signed extension manifests."],
        ["Privacy and trust", "Protected local storage, secret redaction, opt-in egress, encrypted bridges, signed events, exact-source extension verification, and bounded cloud payloads."],
    ], [1.65, 5.0])

    add_heading(doc, "4. Adversary Combat: Standing-Authority Response")
    doc.add_paragraph(
        "Adversary Combat is the unattended response tier. Once armed in Settings, it "
        "consumes authenticated detector evidence and acts immediately without asking "
        "for permission on each incident. Maximum mode explicitly accepts availability risk."
    )
    add_table(doc, ["Control", "Current choices / behavior"], [
        ["Response mode", "Contain, Aggressive, or Maximum"],
        ["Minimum severity", "LOW, MEDIUM, HIGH, or CRITICAL; current default is LOW"],
        ["Process response", "Suspend or terminate the exact attributed process; current default is terminate"],
        ["Network response", "Block exact remote IP indicators and exact program targets"],
        ["File response", "Quarantine an exact evidence-bound file with restoration metadata"],
        ["Host isolation", "Optional firewall isolation after a configurable event threshold; current threshold is 3"],
        ["Deception", "Optionally starts Angerona honeypots/canaries during combat operations"],
        ["Receipts and undo", "Durable JSONL action receipts; reversible actions can be undone individually or as a group"],
        ["Scope integrity", "Actions stay bound to the PID, file, program, or remote address carried by the triggering detector evidence"],
    ], [1.65, 5.0])

    add_heading(doc, "5. Architecture and Operating Model")
    add_table(doc, ["Layer", "Responsibility"], [
        ["Sensors and modules", "Collect telemetry, detect threats, correlate evidence, respond, and publish health."],
        ["Signed EventBus", "Provides thread-safe local pub/sub, bounded recent history, priority handling, and event-authenticity checks."],
        ["Module Manager", "Auto-discovers, starts, stops, supervises, and reports platform availability for built-in and verified external modules."],
        ["Flight recorder", "Persists the local event ledger for timelines, hunting, reports, and incident evidence."],
        ["Causal incident graph", "Links process generations, files, network indicators, parentage, response, and proof without conflating reused PIDs."],
        ["Response plane", "Runs exact-target playbooks, records outcomes, verifies postconditions, and exposes reversal where technically possible."],
        ["Local AI plane", "Uses Ollama for triage and assistance behind prompt, size, privacy, and output controls; deterministic fallbacks remain available."],
        ["Desktop operations", "Provides dashboards, settings, investigations, testing, response controls, readiness status, and guided setup."],
    ], [1.7, 4.95])

    add_heading(doc, "6. Assurance, Evidence, and Current Status")
    add_table(doc, ["Gate", "Current result"], [
        ["Repository test suite", "1,083 passed; 3 intentional platform skips; 0 failures"],
        ["Static quality", "Ruff clean on the changed implementation and validation paths"],
        ["Module discovery", "67 capabilities discovered; 0 import errors; 0 duplicate codes"],
        ["Adversary Combat unit/integration tests", "25 focused tests passed"],
        ["Extreme Shark Attack campaign", "58 total steps; 52/52 eligible detections; 52/52 automatic responses"],
        ["Closure verification", "13/13 contracts and 13/13 closures verified"],
        ["Measured drill latency", "Average detection 0.44 s; average mitigation 0.26 s"],
        ["Manual quality gate", "Master Manual rendered to 118 pages and visually inspected page by page"],
    ], [1.9, 4.75])
    doc.add_paragraph(
        "The campaign uses controlled, benign test artifacts and processes. Its 100% result "
        "proves coverage and action for the tested contracts; it is not a claim that any "
        "security product can detect every unknown real-world attack."
    )

    add_heading(doc, "7. Complete Capability Catalogue")
    doc.add_paragraph(
        "Every current auto-discovered BaseModule capability is listed below. The catalogue "
        "contains exactly 67 unique capabilities and is generated from current source metadata."
    )
    category_rows = [[category, str(categories[category])] for category in sorted(categories)]
    category_rows.append(["TOTAL", str(sum(categories.values()))])
    add_table(doc, ["Category", "Count"], category_rows, [4.8, 1.85])

    current_category = None
    for name, category, description in modules:
        if category != current_category:
            add_heading(doc, f"{category} ({categories[category]})", level=2)
            current_category = category
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        row = table.rows[0]
        _no_split(row)
        cell = row.cells[0]
        _shade(cell, PALE)
        cell.width = Inches(6.65)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(name)
        run.bold = True
        run.font.name = "Aptos"
        run.font.size = Pt(9.4)
        run.font.color.rgb = RGBColor.from_string(BLUE)
        detail = cell.add_paragraph(description)
        detail.paragraph_format.space_after = Pt(0)
        detail.runs[0].font.size = Pt(8.4)
        detail.runs[0].font.color.rgb = RGBColor.from_string(INK)
        _cell_margins(cell, top=75, bottom=75, start=120, end=120)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_heading(doc, "8. Interfaces and Operator Workflows")
    add_bullets(doc, [
        "Dashboard and Local SOC workspace for posture, modules, live alerts, investigations, and operations.",
        "Settings for response posture, privacy, AI, integrations, resource use, appearance, voice, trusted processes, and startup behavior.",
        "Interactive console for process and connection inspection, containment, module control, incident review, MITRE coverage, remediation logs, SQL hunting, testing, and local AI questions.",
        "Shark Attack console for selectable scenarios, difficulty, custom techniques, live offense monitoring, evidence correlation, and after-action reporting.",
        "Threat Intelligence and CVE analysis surfaces for CISA KEV correlation, local exposure context, and staged remediation guidance.",
        "ARIA assistant, HUD, runbook retrieval, routines, research, inbox triage, voice, and optional outbound notifications.",
        "Installer, startup repair, diagnostics, backup validation, self-checks, release checksums, SBOM/provenance workflows, and packaged Windows deployment paths.",
    ])

    add_heading(doc, "9. Current Boundaries")
    add_table(doc, ["Boundary", "Current position"], [
        ["Product maturity", "Substantial working security-engineering platform and lab suite; not independently certified as an enterprise EDR replacement."],
        ["Deployment scale", "Strongest as a local workstation/lab platform. Central fleet administration, multi-node high availability, and enterprise identity governance remain separate deployment concerns."],
        ["Kernel enforcement", "Primary product is elevated user mode using supported Windows interfaces. Optional kernel bridge code exists, but no unsigned production driver should be treated as a trust anchor."],
        ["Cross-platform", "Windows is the primary response platform; Linux and macOS currently emphasize privacy-minimized observation and optional native sensor paths."],
        ["AI", "Local AI improves explanation and workflow assistance; deterministic controls and detector evidence remain the authority for response."],
        ["Validation meaning", "Automated tests and drills establish repeatability for implemented contracts, not universal detection of every future technique."],
        ["Autonomous action", "Maximum mode is intentionally disruptive. Administrators choose the policy and retain durable evidence plus undo for reversible actions."],
    ], [1.65, 5.0])

    doc.add_page_break()
    add_heading(doc, "10. Practical Value")
    doc.add_paragraph(
        "Angerona is a credible advanced home-lab suite, defensive R&D platform, purple-team "
        "exercise system, and security-engineering portfolio project. Its strongest evidence is "
        "the integrated breadth: sensors, correlation, response, reversal, testing, UI, packaging, "
        "privacy controls, and documentation operate as one system rather than isolated demos."
    )
    add_bullets(doc, [
        "For defenders: a hands-on platform for building and proving detections and containment playbooks.",
        "For developers: a large modular Python/Qt codebase with concurrency, persistence, security boundaries, packaging, and automated verification.",
        "For employers: concrete evidence of initiative and cross-domain security engineering when presented with a live demo, architecture explanation, and honest scope boundaries.",
        "For researchers: an extensible place to prototype local-first defensive ideas and measure them against repeatable simulations.",
    ])
    add_heading(doc, "Evidence a reviewer can inspect", level=2)
    add_table(doc, ["Evidence", "What it demonstrates"], [
        ["Working desktop application", "Product integration, operator experience, lifecycle handling, and real-time state management"],
        ["67-module catalogue", "Breadth across telemetry, detection, response, AI, resilience, forensics, deception, and integration"],
        ["1,083-test green gate", "Regression discipline, edge-case coverage, and maintainable interfaces"],
        ["Shark Attack AAR", "Measurable detection and response behavior rather than feature-list claims alone"],
        ["Receipts, reversal, and evidence binding", "Security reasoning around authorization, attribution, auditability, and recovery"],
        ["Master Manual and current bragsheet", "Ability to communicate system design, operations, limits, and proof to technical and non-technical reviewers"],
    ], [2.05, 4.6])

    # Avoid an empty trailing paragraph carrying an accidental blank page.
    for section in doc.sections[1:]:
        if section.start_type == WD_SECTION.NEW_PAGE:
            section.start_type = WD_SECTION.CONTINUOUS
    doc.save(OUTPUT)
    shutil.copyfile(OUTPUT, ANALYSIS_OUTPUT)
    print(f"wrote {OUTPUT}")
    print(f"synchronized {ANALYSIS_OUTPUT}")
    print(f"catalogued {len(modules)} modules across {len(categories)} categories")


if __name__ == "__main__":
    build()
