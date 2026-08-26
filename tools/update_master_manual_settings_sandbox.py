"""Append the tab-aware Settings source-sandbox addition to the master manual."""
from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "Angerona_Master_Manual.docx"
HEADING = "17.3 Tab-aware Settings code sandbox (2026-08-25)"


def _insert_before(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def main() -> None:
    document = Document(MANUAL)
    if any(paragraph.text.strip() == HEADING for paragraph in document.paragraphs):
        print("Master Manual already contains the Settings sandbox addition.")
        return

    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Appendix A. Command reference"
    )
    _insert_before(document, anchor, HEADING, "Heading 2")
    _insert_before(
        document,
        anchor,
        "Every functional tab in Settings now exposes a tab-aware code-sandbox "
        "button. The button follows the selected tab, resolves its registered "
        "implementation files, opens isolated editable working copies, and places "
        "the editor at the corresponding tab-builder code.",
        "Normal",
    )
    bullets = (
        "Covered tabs: Overview, Information, General, System, Adversary Combat, "
        "Enterprise, ARIA, Trusted Processes, Mobile Integration, and API Keys.",
        "Each sandbox includes the tab's related implementation files in a picker. "
        "For example, ARIA opens the Settings UI at _tab_aria and also exposes its "
        "assistant and configuration implementation files.",
        "Check Syntax, Save Sandbox Copy, Reset Current Copy, and Reset All Sandbox "
        "Changes operate only on the private working copy. Installed Angerona code "
        "and live sensors are never rewritten or hot-loaded by this editor.",
        "The Adversary Combat Settings topic now has a complete source mapping for "
        "its policy, action, receipt, history, and Undo implementation.",
        "Verification: 49 focused Settings/UI regression tests passed; the full "
        "suite passed 1,258 tests with three intentional platform skips; Ruff and "
        "documentation-drift checks passed.",
    )
    for bullet in bullets:
        _insert_before(document, anchor, bullet, "List Bullet")
    document.save(MANUAL)
    print(f"Updated {MANUAL}")


if __name__ == "__main__":
    main()
