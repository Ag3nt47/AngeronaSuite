"""Append the verified Cycle 7 enterprise foundation record to the manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def update(source: Path, destination: Path) -> None:
    doc = Document(source)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Cycle 7. Enterprise Control-Plane Foundations", level=1)
    doc.add_paragraph(
        "Verification date: 29 July 2026. This section consolidates locally "
        "testable enterprise recommendations from the Angerona analysis set. "
        "It distinguishes implemented foundations from external production gates."
    )

    doc.add_heading("Implemented local foundations", level=2)
    add_bullets(doc, (
        "A bounded asynchronous authoritative flight recorder and normalized "
        "evidence worker use batching, overflow/dead-letter metrics, and orderly "
        "shutdown draining. A 10,000-event recorder benchmark sustained about "
        "19,600 publishes per second without authoritative loss.",
        "Signed detection packages revalidate package bytes, detached Ed25519 "
        "signature, current trust key, expiry, and activation policy. Tampering "
        "or trust revocation fails closed and quarantines the active package.",
        "Endpoint identity, one-time enrollment, revocation, replay protection, "
        "persisted connection sequencing, targeted fleet jobs, signed results, "
        "signed policy rollout/rollback, and asset inventory now have bounded "
        "offline-first contracts.",
        "Role-Based Access Control (RBAC) supports human and expiring service "
        "principals, explicit-deny precedence, canonical fleet/group/host scopes, "
        "bounded permission wildcards, and request-bound authenticated decisions.",
        "Case management provides optimistic concurrency, bounded timelines, "
        "reference-only evidence, legal holds, privacy-minimized export, "
        "authenticated custody heads, and signed retention receipts.",
        "The Artificial Intelligence (AI) broker separates explain, recommend, "
        "plan, and execute modes. Generated PowerShell and Python are staged for "
        "review and cannot execute or hot-load as model authority.",
        "Release assurance provides deterministic manifests, CycloneDX Software "
        "Bill of Materials (SBOM), signed offline bundles, compatibility and "
        "rollback preflight, Vulnerability Exploitability eXchange (VEX) data, "
        "and a Supply-chain Levels for Software Artifacts (SLSA)-compatible "
        "provenance skeleton.",
    ))

    doc.add_heading("Security and privacy remediation closure", level=2)
    add_bullets(doc, (
        "Case exports redact common credentials, bearer tokens, email identities, "
        "and user-profile paths. Restricted evidence references and free-form "
        "comment bodies are excluded by default and the artifact includes a "
        "privacy manifest.",
        "Evidence custody verification authenticates both each row and an "
        "independent final hash/count head. Deleting the newest row, changing "
        "the head, or altering a signed retention receipt now fails verification.",
        "Authorization receipts bind principal, permission, canonical scope, "
        "resource, policy, request identity, and request digest. Conflicting "
        "same-ID operations fail; exact retries are idempotent in the retained "
        "10,000-entry process-local window.",
        "GitHub Actions are pinned to immutable 40-character commits, use "
        "least-privilege permissions, and are checked by a deterministic workflow "
        "policy test. Secret scanning, CodeQL, dependency review, Scorecard, and "
        "multi-version Continuous Integration (CI) are configured.",
    ))

    doc.add_heading("Verification evidence", level=2)
    add_bullets(doc, (
        "Repository test suite: 334 passed, 2 intentional platform skips, 0 failed.",
        "Headless integration self-check: 26 of 26 passed.",
        "Python compilation, dependency integrity, and workflow-policy gates passed.",
        "Final focused security remediation suite: 16 of 16 passed.",
        "Independent final adversarial recheck: no remaining High or Critical "
        "finding in the new enterprise foundations.",
    ))

    doc.add_heading("Honest enterprise boundary", level=2)
    doc.add_paragraph(
        "These controls materially raise assurance and close locally testable "
        "roadmap foundations; they do not by themselves create a deployed "
        "enterprise control plane. Production mutual Transport Layer Security "
        "(mTLS), tenant-isolated central services, durable cross-process replay "
        "storage, a separately privileged containment broker, Authenticode "
        "certificate custody, GitHub branch rules, long-duration physical-host "
        "soak tests, compatibility labs, external penetration testing, production "
        "driver qualification, and Apple entitlement/notarization remain external "
        "or release-level gates."
    )

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | Cycle 7 consolidated 2026-07-29 | Page "
                )
                break
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle7.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
