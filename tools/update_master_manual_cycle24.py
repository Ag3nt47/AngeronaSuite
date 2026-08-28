"""Apply the minimal Cycle 24 addendum to the canonical Angerona manual.

The first run preserves the complete Cycle 23 manual under ``.tmp``.  Every
later QA iteration rebuilds from that snapshot, so the release addendum cannot
be duplicated and the established manual layout remains otherwise unchanged.
"""
from __future__ import annotations

import os
from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Angerona_Master_Manual.docx"
BUILD_ROOT = ROOT / ".tmp" / "docx_cycle24"
PRISTINE = BUILD_ROOT / "Angerona_Master_Manual_pre_cycle24.docx"
STAGED = BUILD_ROOT / "Angerona_Master_Manual_cycle24.docx"
MARKER = "17.6 v1.11.0 independent-trust hardening (2026-08-27)"


def _set_text(paragraph, text: str) -> None:
    """Replace visible text while retaining the paragraph and first-run style."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _find_exact(document: Document, text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _insert_before(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _update_control_and_validation(document: Document) -> None:
    if len(document.tables) < 2:
        raise ValueError("manual is missing its document-control tables")

    _set_text(
        document.tables[0].cell(0, 0).paragraphs[0],
        "Current release: v1.11.0 - independent-trust hardening for "
        "authenticated telemetry continuity, hostile network paths, "
        "operator-gated response, governed AI memory, recovery assurance, "
        "and signed release boundaries.",
    )

    control_rows = {
        row.cells[0].text.strip(): row.cells[1]
        for row in document.tables[1].rows
        if len(row.cells) >= 2
    }
    required_control = {"Version", "Release state", "Source of truth"}
    if not required_control.issubset(control_rows):
        raise ValueError("manual document-control fields changed unexpectedly")
    _set_text(control_rows["Version"].paragraphs[0], "1.11.0")
    _set_text(
        control_rows["Release state"].paragraphs[0],
        "Cycle 24 defensive hardening and final serial validation complete; "
        "signed deployment artifacts remain externally governed outputs",
    )
    _set_text(
        control_rows["Source of truth"].paragraphs[0],
        "Current repository code and analysis/loop/cycle24 evidence as of "
        "27 August 2026",
    )

    validation_tables = [
        table
        for table in document.tables
        if table.rows
        and len(table.rows[0].cells) >= 2
        and table.rows[0].cells[0].text.strip() == "Gate"
        and table.rows[0].cells[1].text.strip() == "Final result"
    ]
    if len(validation_tables) != 1:
        raise ValueError("manual validation summary table changed unexpectedly")
    validation_rows = {
        row.cells[0].text.strip(): row.cells[1]
        for row in validation_tables[0].rows[1:]
        if len(row.cells) >= 2
    }
    results = {
        "Full pytest": (
            "1,675 collected across 229 files; 1,670 passed; 5 expected "
            "host-capability skips; 0 failed"
        ),
        "Compile": "611/611 Python files; 345/345 product files",
        "Static/runtime quality": (
            "Ruff; 80 Windows / 14 Linux / 13 macOS discovery; zero discovery "
            "errors; documentation-drift and patch-integrity gates passed"
        ),
        "Module self-tests": (
            "60 genuine pass; 0 genuine fail; 21 expected inactive/platform "
            "skips; EventBus pipeline passed"
        ),
        "Application selfcheck": "26/26 passed",
    }
    if not set(results).issubset(validation_rows):
        raise ValueError("manual validation summary rows changed unexpectedly")
    for field, value in results.items():
        _set_text(validation_rows[field].paragraphs[0], value)

    verdict_cells = [
        cell
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.startswith("Final red-team verdict:")
    ]
    if len(verdict_cells) != 1:
        raise ValueError("manual final-verdict callout changed unexpectedly")
    _set_text(
        verdict_cells[0].paragraphs[0],
        "Final red-team verdict: No open Critical, High, or Medium Cycle 24 "
        "code blocker remains in the current v1.11.0 tree. Publisher/root "
        "custody, clean-machine package validation, hardware rollback "
        "resistance, and the external gates below remain outstanding "
        "acceptance work.",
    )

    footer_matches = [
        paragraph
        for paragraph in document.sections[0].first_page_footer.paragraphs
        if "Angerona Suite" in paragraph.text
    ]
    if len(footer_matches) != 1:
        raise ValueError("manual first-page footer changed unexpectedly")
    _set_text(
        footer_matches[0],
        "Angerona Suite  |  v1.11.0  |  27 August 2026",
    )

    header_updates = 0
    for section in document.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            for paragraph in header.paragraphs:
                if "ANGERONA  |  MASTER MANUAL  |" in paragraph.text:
                    _set_text(
                        paragraph,
                        "ANGERONA  |  MASTER MANUAL  |  v1.11.0",
                    )
                    header_updates += 1
    if not header_updates:
        raise ValueError("manual running header changed unexpectedly")


def _update_existing_text(document: Document) -> None:
    cover_dates = [p for p in document.paragraphs if p.text.strip() == "26 August 2026"]
    if len(cover_dates) != 1:
        raise ValueError(f"expected one cover date, found {len(cover_dates)}")
    _set_text(cover_dates[0], "27 August 2026")

    overview_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Angerona unifies Windows endpoint")
        and "discovered modules share a typed EventBus" in paragraph.text
    ]
    if len(overview_matches) != 1:
        raise ValueError("manual overview paragraph changed unexpectedly")
    _set_text(
        overview_matches[0],
        "Angerona unifies Windows endpoint and network visibility, local "
        "evidence retention, case work, MITRE ATT&CK mapping, governed "
        "containment, recovery, local AI assistance, and non-destructive "
        "purple-team validation. Its current discovery contract covers 80 "
        "Windows, 14 Linux, and 13 macOS defensive modules sharing a typed "
        "EventBus and explicit platform contracts.",
    )

    install_updates = {
        "Download the Setup executable and its adjacent SHA-256 file from the GitHub Releases page.": (
            "Download the signed x64 full-trust MSIX and its exact SHA-256 and "
            "provenance records from the tagged release."
        ),
        "Verify the digest and GitHub build attestation. The current publication path is not backed by an Authenticode publisher certificate, so Windows may show Unknown Publisher.": (
            "Confirm the exact digest and expected Windows package publisher; "
            "refuse unsigned, unexpected-publisher, or unprovisioned builds."
        ),
        "Run Setup, approve UAC, and complete Full Setup. The protected data root prefers D:\\AngeronaData and uses protected %ProgramData%\\Angerona only when D: is unavailable.": (
            "Install through the Windows package UI. Classic Setup is non-public "
            "and prior-install migration-only; ZIPs are verifier-gated upgrades, "
            "never first-install paths."
        ),
    }
    for old, new in install_updates.items():
        _set_text(_find_exact(document, old), new)

    memory_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(
            "ARIA now indexes a built-in defensive reference with 10 entries"
        )
    ]
    if len(memory_matches) != 1:
        raise ValueError("manual Defense Memory summary changed unexpectedly")
    _set_text(
        memory_matches[0],
        memory_matches[0].text.replace("with 10 entries", "with 18 entries", 1),
    )
    cloud_memory_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(
            "Local retrieval may use the complete bounded synthesized reference."
        )
    ]
    if len(cloud_memory_matches) != 1:
        raise ValueError("manual Defense Memory cloud boundary changed unexpectedly")
    _set_text(
        cloud_memory_matches[0],
        cloud_memory_matches[0].text.replace(
            "can receive only selected redacted excerpts",
            "can receive at most one ranked canonical redacted excerpt",
            1,
        ),
    )

    evidence_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(
            "Repository evidence: analysis/loop/LOOP_LOG.md; "
            "analysis/loop/cycle23 research"
        )
    ]
    if len(evidence_matches) != 1:
        raise ValueError("manual repository-evidence summary changed unexpectedly")
    _set_text(
        evidence_matches[0],
        evidence_matches[0].text.replace(
            "analysis/loop/cycle23 research",
            "analysis/loop/cycle23 and cycle24 research",
            1,
        ),
    )


def _append_release_addendum(document: Document) -> None:
    if any(p.text.strip() == MARKER for p in document.paragraphs):
        raise ValueError("Cycle 24 release addendum already exists")
    anchor = _find_exact(document, "Appendix A. Command reference")
    blocks = (
        (
            "Heading 2",
            MARKER,
        ),
        (
            "Normal",
            "Cycle 24 adds independent-trust and continuity controls while "
            "preserving Angerona's observe, review, approve, execute, and "
            "verify boundaries. It does not add credential theft, intrusion, "
            "log deletion, attribution claims, or autonomous destructive "
            "remediation.",
        ),
        (
            "List Bullet",
            "Authenticated producer identity, schema admission, sequence/loss "
            "accounting, SSH provenance, audit-log continuity, and recovery "
            "evidence expose gaps without claiming to reconstruct erased records.",
        ),
        (
            "List Bullet",
            "Wi-Fi and Ethernet remain hostile-by-default. The Personal "
            "Sentinel reference authority attests an explicitly enrolled "
            "intermediate gateway path; it does not configure routers or grant "
            "endpoint, application, user, or response trust.",
        ),
        (
            "List Bullet",
            "New defensive coverage includes identity/session and temporal "
            "correlation, measured boot, peripheral/DMA and driver provenance, "
            "process-bound egress, RAG provenance, trusted time, recovery "
            "assurance, update authority, and release transparency.",
        ),
        (
            "List Bullet",
            "The dashboard's Live Defense Activity card shows only bounded, "
            "sanitized public EventBus records - never source code, secrets, "
            "private model reasoning, or chain of thought. ARIA Defense Memory "
            "v1.1.0 contains 18 pinned defensive references; an optional cloud "
            "fallback receives at most one ranked canonical redacted excerpt.",
        ),
        (
            "List Bullet",
            "Public first installation is the signed x64 full-trust MSIX. The "
            "classic migration wrapper is non-public and prior-install-only; "
            "portable ZIPs are upgrade-only and require an installed verifier. "
            "Signing identity, protected keys, pinned compiler images, and "
            "enterprise distribution remain external deployment controls.",
        ),
        (
            "List Bullet",
            "Final automated evidence: 1,675 collected across 229 test files; "
            "1,670 passed, five expected host-capability skips, and zero failed. "
            "All 611 Python files compiled; Ruff, 26 application selfchecks, "
            "the 60-pass module harness, and discovery for 80 Windows, 14 Linux, "
            "and 13 macOS modules passed with zero discovery errors.",
        ),
        (
            "Normal",
            "These results are project evidence, not independent certification "
            "or proof against every administrator, kernel, firmware, hypervisor, "
            "supply-chain, identity, physical-access, or whole-host rollback attack.",
        ),
    )
    for style, text in blocks:
        paragraph = _insert_before(document, anchor, text, style)
        if style == "Heading 2":
            paragraph.paragraph_format.keep_with_next = True


def update() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    BUILD_ROOT.mkdir(parents=True, exist_ok=True)
    if not PRISTINE.exists():
        probe = Document(SOURCE)
        if any(p.text.strip() == MARKER for p in probe.paragraphs):
            raise ValueError("cannot snapshot a manual that already has Cycle 24")
        shutil.copy2(SOURCE, PRISTINE)

    document = Document(PRISTINE)
    _update_existing_text(document)
    _update_control_and_validation(document)
    _append_release_addendum(document)

    document.core_properties.title = "Angerona Master Manual"
    document.core_properties.subject = (
        "v1.11.0 operator reference with Cycle 24 independent-trust hardening"
    )
    document.core_properties.comments = (
        "Canonical manual minimally updated 27 August 2026; actor-neutral "
        "defensive evidence."
    )
    document.save(STAGED)

    reopened = Document(STAGED)
    visible = "\n".join(p.text for p in reopened.paragraphs)
    required = (
        MARKER,
        "27 August 2026",
        "1,670 passed",
        "80 Windows, 14 Linux, and 13 macOS",
        "Live Defense Activity",
        "ARIA Defense Memory v1.1.0",
        "signed x64 full-trust MSIX",
    )
    missing = [item for item in required if item not in visible]
    if missing:
        raise ValueError(f"updated manual is missing required content: {missing}")
    if sum(p.text.strip() == MARKER for p in reopened.paragraphs) != 1:
        raise ValueError("Cycle 24 release addendum was not unique after reopen")
    if reopened.paragraphs[11].text.strip() != "27 August 2026":
        raise ValueError("cover date failed reopen validation")
    os.replace(STAGED, SOURCE)
    print(f"updated {SOURCE}")


if __name__ == "__main__":
    update()
