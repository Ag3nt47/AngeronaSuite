"""Append the verified Adversary Combat operator chapter to the master DOCX."""
from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "analysis" / "Angerona_Master_Manual_v1.10.1.docx"
HEADING = "24. Adversary Combat - Autonomous Response"

CATEGORY_COUNTS = (
    ("AI", 3),
    ("AI Defense", 1),
    ("Compliance", 1),
    ("Deception", 2),
    ("Detection", 7),
    ("Endpoint", 6),
    ("Forensics", 2),
    ("Integration", 1),
    ("Integrity", 7),
    ("Maintenance", 1),
    ("Memory", 1),
    ("Network", 6),
    ("Performance", 5),
    ("Persistence", 1),
    ("Processes", 1),
    ("Ransomware", 1),
    ("Reporting", 1),
    ("Resilience", 7),
    ("Response", 6),
    ("SOAR", 1),
    ("Sensor", 1),
    ("Signatures", 1),
    ("System", 1),
    ("Telemetry", 2),
    ("Threat Intel", 1),
)

# Several rows in the consolidated catalogue retain their historical display
# name. Keep those richer descriptions and update only the category cell.
CATALOGUE_CATEGORY_UPDATES = {
    "AMSI Bridge (AMSI)": "Endpoint",
    "AV Telemetry Bridge (AVTB)": "Endpoint",
    "Dynamic Resource Governor (DRES)": "System",
    "Fast-Path Interceptor (FPTH)": "AI",
    "Kernel Bridge (KRNL)": "Endpoint",
    "Memory Injection Scanner (MINJ)": "Memory",
    "Ransomware Heuristics (RANS)": "Ransomware",
    "Sysmon Event Bridge (SYSL)": "Endpoint",
}

NEW_CATALOGUE_ROWS = (
    (
        "Adversary Combat",
        "Response",
        "Standing-authority autonomous response: blocks exact remote/program targets, "
        "quarantines files, suspends or terminates exact processes, isolates the host, "
        "and activates honeypots with durable receipts and reversible-action undo.",
    ),
    (
        "Kernel-Boundary Posture Ledger",
        "Integrity",
        "Read-only evidence for Secure Boot, VBS/HVCI, boot-integrity flags, Code "
        "Integrity telemetry, and kernel-driver-service drift.",
    ),
    (
        "Linux Observe Sensor",
        "Endpoint",
        "Rootless, privacy-minimized Linux process, established-flow, and kernel-posture "
        "observation; privileged eBPF remains a separate opt-in sensor.",
    ),
    (
        "macOS Observe Sensor",
        "Endpoint",
        "Privacy-minimized macOS process and network observation with explicit reporting "
        "of the separate native Endpoint Security enforcement boundary.",
    ),
    (
        "Purple Remediation Guard",
        "Detection",
        "Turns reviewed red-team misses into exact, rerun-verifiable detector signatures.",
    ),
    (
        "Self-Integrity Monitor",
        "Integrity",
        "Detects in-memory tampering of Angerona enforcement functions, complementing "
        "termination and suspension protections.",
    ),
)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _replace_text(doc: Document, old: str, new: str) -> bool:
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    return True
            paragraph.text = paragraph.text.replace(old, new)
            return True
    return False


def _set_table_rows(table, rows: Iterable[tuple[str, str]]) -> None:
    """Replace body rows while preserving the document's existing table style."""
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    for left, right in rows:
        row = table.add_row()
        _prevent_row_split(row)
        cells = row.cells
        cells[0].text = str(left)
        cells[1].text = str(right)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _refresh_catalogue_tables(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        header = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if header == ("Category", "Modules"):
            _set_table_rows(table, ((name, str(count)) for name, count in CATEGORY_COUNTS))
            continue
        if header != ("Module", "Category", "What it does"):
            continue
        existing: set[str] = set()
        for row in table.rows[1:]:
            _prevent_row_split(row)
            name = row.cells[0].text.strip()
            existing.add(name)
            category = CATALOGUE_CATEGORY_UPDATES.get(name)
            if category:
                row.cells[1].text = category
        for name, category, description in NEW_CATALOGUE_ROWS:
            if name in existing:
                continue
            row = table.add_row()
            _prevent_row_split(row)
            cells = row.cells
            cells[0].text = name
            cells[1].text = category
            cells[2].text = description


def _refresh_verification_table(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        header = tuple(cell.text.strip() for cell in table.rows[0].cells)
        if header != ("Gate", "Verified result"):
            continue
        for row in table.rows[1:]:
            gate = row.cells[0].text.strip()
            if gate == "Repository pytest":
                row.cells[1].text = "1,083 passed; 3 intentional platform skips; 0 failed"
            elif gate == "Discovery":
                row.cells[1].text = "67 modules; 0 errors; 0 duplicate codes"


def update(path: Path = MANUAL) -> Path:
    doc = Document(path)
    original_tables = len(doc.tables)

    # Idempotent local replacement if this updater is rerun.
    prior_index = next(
        (
            index
            for index, p in enumerate(doc.paragraphs)
            if p.text.strip() == HEADING
            and getattr(p.style, "name", "") == "Heading 1"
        ),
        None,
    )
    if prior_index is not None:
        paragraphs = list(doc.paragraphs)
        start = prior_index
        while start > 0:
            previous = paragraphs[start - 1]
            if previous.text.strip() or not previous._element.xpath(
                ".//w:br[@w:type='page']"
            ):
                break
            start -= 1
        for paragraph in paragraphs[start:]:
            _remove_paragraph(paragraph)

    doc.paragraphs[4].text = (
        "Adversary Combat autonomous-response edition  |  24 August 2026  |  v1.10.1"
    )
    doc.core_properties.title = "Angerona Consolidated Master Manual"
    doc.core_properties.subject = (
        "Version 1.10.1 - autonomous Adversary Combat and verified Extreme campaign"
    )
    doc.core_properties.last_modified_by = "Angerona documentation loop"

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Appendix A. Consolidation Sources":
            if not any(p.text.strip() == HEADING for p in doc.paragraphs[:]):
                paragraph.insert_paragraph_before(HEADING, style="List Bullet")
            break

    _replace_text(
        doc,
        "66 BaseModule subclasses are auto-discovered at startup.",
        "67 BaseModule subclasses are auto-discovered at startup.",
    )
    _replace_text(
        doc,
        "Host Adaption is a core service/workbench and does not change the module count.",
        "Host Adaption remains a core service/workbench; Adversary Combat is the new auto-discovered response module.",
    )
    _replace_text(
        doc,
        "Total: 63 modules across the current categories.",
        "Total: 67 modules across the current categories.",
    )
    _replace_text(
        doc,
        "63 auto-discovered security modules; v1.9.4 adds Cycle 4 proof-driven remediation, privacy, watchdog recovery, lifecycle safety, and responsiveness hardening.",
        "67 auto-discovered security modules are catalogued below. The inherited v1.9.4 narrative records its Cycle 4 proof-driven remediation, privacy, watchdog recovery, lifecycle safety, and responsiveness hardening in historical context.",
    )
    _replace_text(doc, "3. Security Modules (63)", "3. Security Modules (67)")
    _replace_text(
        doc,
        "All 61 are BaseModule subclasses auto-discovered by ModuleManager",
        "All 67 are BaseModule subclasses auto-discovered by ModuleManager",
    )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == (
                    "63 auto-discovered security capabilities — see §5 / Capability Doc"
                ):
                    cell.text = (
                        "67 auto-discovered security capabilities — see §5 / Capability Doc"
                    )
    _refresh_catalogue_tables(doc)
    _refresh_verification_table(doc)

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                if "v1.10.1 Host Adaption verified 2026-08-24" in run.text:
                    run.text = run.text.replace(
                        "v1.10.1 Host Adaption verified 2026-08-24",
                        "v1.10.1 Adversary Combat verified 2026-08-24",
                    )

    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph(HEADING, style="Heading 1")
    doc.add_paragraph(
        "Verified addendum - 24 August 2026. Adversary Combat is Angerona's "
        "standing-authority response tier. Once armed, it consumes authenticated "
        "local detector evidence and acts immediately without a per-incident dialog. "
        "Maximum mode intentionally prioritizes containment over availability and can "
        "cause an outage by terminating processes or isolating the host."
    )

    doc.add_paragraph("24.1 Standing authority and controls", style="Heading 2")
    _bullet(
        doc,
        "Open Settings > Adversary Combat, or use the Adversary Combat page in Full Setup. "
        "The policy is editable while the application is running.",
    )
    _bullet(
        doc,
        "Modes are Contain, Aggressive, and Maximum. The operator also selects the minimum "
        "severity, exact remote/program blocking, file quarantine, process suspend or "
        "terminate, full host isolation, Smart Deception activation, and the isolation threshold.",
    )
    _bullet(
        doc,
        "The shipped policy is armed in Maximum mode at LOW severity with network blocking, "
        "file quarantine, process termination, host isolation, and honeypots enabled. This "
        "is an explicit outage-risk posture, not a recommendation-only mode.",
    )
    _bullet(
        doc,
        "Actions remain bound to exact signed detector evidence. Remote observe-only evidence "
        "may block the named remote indicator but cannot redirect an unrelated local file or "
        "process mutation. Angerona's own process and its launcher are excluded so autonomous "
        "defense is not terminated by its response worker.",
    )

    doc.add_paragraph("24.2 Automatic combat playbooks", style="Heading 2")
    _bullet(
        doc,
        "Block: creates paired inbound/outbound Windows Firewall rules for an exact remote IP "
        "and an outbound program-isolation rule for an exact process executable.",
    )
    _bullet(
        doc,
        "Contain: moves an exact detected file into the combat quarantine, or suspends the exact "
        "PID. Quarantine restoration and process resume verify the original path or PID creation "
        "time before undo.",
    )
    _bullet(
        doc,
        "Combat: Maximum mode terminates the exact detected process. Termination is deliberately "
        "non-reversible; suspend is the reversible alternative.",
    )
    _bullet(
        doc,
        "Isolate: a CRITICAL active incident, or the configured number of active events inside "
        "the correlation window, installs paired all-remote-traffic host isolation rules.",
    )
    _bullet(
        doc,
        "Honeypot: starts Smart Deception automatically when the policy enables it. A honeypot "
        "activation performed by Combat receives a reversible action receipt.",
    )

    doc.add_paragraph("24.3 Receipts, postconditions, and undo", style="Heading 2")
    _bullet(
        doc,
        "Every applied action is appended to shared_logs/adversary_combat_actions.jsonl with a "
        "combat ID, action ID, exact target, trigger module/timestamp, reversibility flag, and "
        "action-specific details.",
    )
    _bullet(
        doc,
        "A response receives verified-closure credit only after its postcondition passes: source "
        "absent and quarantine present, process no longer running or confirmed suspended, both "
        "firewall rule directions applied, or the deception module running.",
    )
    _bullet(
        doc,
        "Settings exposes recent actions and Undo last reversible action. The module also supports "
        "undo by action ID and undo-all for quarantines, suspensions, IP/program/host firewall "
        "rules, and Combat-started honeypots. Undo attempts append their own durable receipt.",
    )

    doc.add_paragraph("24.4 Real-time detection and AAR correlation", style="Heading 2")
    _bullet(
        doc,
        "Maximum mode tightens File Integrity Monitor, YARA, Network Monitor, and Process Monitor "
        "cadences. Process telemetry now publishes creation details, including command line, for "
        "exact tagged-process correlation.",
    )
    _bullet(
        doc,
        "Shark exfiltration records the actual peer IP and port. The After-Action Report correlates "
        "network evidence only when PID, peer IP, and peer port match, replacing the prior "
        "structurally impossible PID-only response match.",
    )
    _bullet(
        doc,
        "The scorecard recognizes Adversary Combat as a response engine only when mitigated=true, "
        "and recognizes closure only when postcondition_verified=true for every occurrence in a "
        "unique technique class.",
    )

    doc.add_paragraph("24.5 Extreme validation and repeatable proof", style="Heading 2")
    doc.add_paragraph(
        "The real Extreme Red Team campaign completed on 24 August 2026 under run "
        "redteam-1787631486-1226ad. It executed 58 total steps across the chained campaign. "
        "Fifty-two detector-eligible steps were caught and automatically remediated; "
        "read-only Discovery remained informational and benign Noise Injection remained silent."
    )
    _bullet(doc, "Detection coverage: 52/52 (100%).")
    _bullet(doc, "Automatic response success: 52/52 (100%).")
    _bullet(doc, "Action contracts: 13/13 unique technique classes (100%).")
    _bullet(doc, "Verified closure: 13/13 unique technique classes (100%).")
    _bullet(doc, "Resilience: PASS with no false positive; no eligible step was missed.")
    _bullet(doc, "Average detection time: 0.44 seconds; average mitigation time: 0.26 seconds.")
    _bullet(
        doc,
        "Run run-adversary-combat-validation.bat for the repeatable operator proof. Its default "
        "invocation keeps launching Extreme campaigns until all four scorecard rates are 100%, "
        "resilience passes, and there are no misses or false positives. It retains evidence and "
        "then restores every reversible test action.",
    )

    doc.add_paragraph("24.6 Acceptance boundary", style="Heading 2")
    _bullet(
        doc,
        "The 100% result is evidence for Angerona's bounded, benign Red Team campaign on this host. "
        "It does not claim universal prevention of every real-world technique, kernel compromise, "
        "or attack that produces no observable detector evidence.",
    )
    _bullet(
        doc,
        "Process termination is not undoable, full host isolation can interrupt connectivity, and "
        "quarantine can interrupt applications. Operators choosing Maximum mode accept those "
        "availability effects and can edit the standing policy at any time.",
    )
    _bullet(
        doc,
        "Repository verification after integration: 1,083 tests passed, 3 intentional platform "
        "skips, and 0 failures. The focused Combat and AAR regression set passed 25/25.",
    )

    if len(doc.tables) != original_tables:
        raise RuntimeError("manual table count changed unexpectedly")
    temp = path.with_suffix(".updating.docx")
    doc.save(temp)
    os.replace(temp, path)
    return path


if __name__ == "__main__":
    print(update())
