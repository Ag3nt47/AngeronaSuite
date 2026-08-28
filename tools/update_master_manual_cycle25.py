"""Apply the minimal Cycle 25 / v1.12.0 addendum to the master manual.

The first run preserves the complete Cycle 24 manual under ``.tmp``. Every
later QA iteration rebuilds from that snapshot, so the v1.12 material remains
unique while the established operator-manual design stays intact.
"""
from __future__ import annotations

from copy import deepcopy
import os
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Angerona_Master_Manual.docx"
BUILD_ROOT = ROOT / ".tmp" / "docx_cycle25"
PRISTINE = BUILD_ROOT / "Angerona_Master_Manual_pre_cycle25.docx"
STAGED = BUILD_ROOT / "Angerona_Master_Manual_cycle25.docx"
MARKER = "17.7 v1.12.0 guided-defense and capability expansion (2026-08-28)"


def _set_text(paragraph, text: str) -> None:
    """Replace visible text while retaining the paragraph and first-run style."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _find_exact(document: Document, text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _find_startswith(document: Document, prefix: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph starting {prefix!r}, found {len(matches)}")
    return matches[0]


def _find_exact_style(document: Document, text: str, style: str):
    matches = [
        p
        for p in document.paragraphs
        if p.text.strip() == text and p.style.name == style
    ]
    if len(matches) != 1:
        raise ValueError(
            f"expected one {style!r} paragraph {text!r}, found {len(matches)}"
        )
    return matches[0]


def _insert_before(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _insert_after(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addnext(paragraph._p)
    return paragraph


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is None:
        raise ValueError("paragraph is already detached")
    parent.remove(paragraph._p)
    paragraph._p = None
    paragraph._element = None


def _keep_table_row_together(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _update_control_and_validation(document: Document) -> None:
    if len(document.tables) < 21:
        raise ValueError("manual is missing expected control and history tables")

    _set_text(
        document.tables[0].cell(0, 0).paragraphs[0],
        "Current release: v1.12.0 - guided, proposal-only host adaptation; an "
        "immutable Windows Firewall recovery baseline; clickable operational "
        "detail; universal capability contracts; durable integrations; and "
        "standards-truth hardening.",
    )

    control_rows = {
        row.cells[0].text.strip(): row.cells[1]
        for row in document.tables[1].rows
        if len(row.cells) >= 2
    }
    required_control = {"Version", "Release state", "Source of truth"}
    if not required_control.issubset(control_rows):
        raise ValueError("manual document-control fields changed unexpectedly")
    _set_text(control_rows["Version"].paragraphs[0], "1.12.0")
    _set_text(
        control_rows["Release state"].paragraphs[0],
        "Cycle 25 three-round defensive expansion and final serial validation "
        "complete; host mutation remains exact, previewed, approved, journaled, "
        "and verified",
    )
    _set_text(
        control_rows["Source of truth"].paragraphs[0],
        "Current repository code, the 80-capability machine-readable inventory, "
        "and analysis/loop/cycle25 evidence as of 28 August 2026",
    )

    validation_tables = [
        table
        for table in document.tables
        if table.rows
        and len(table.rows[0].cells) >= 2
        and table.rows[0].cells[0].text.strip() == "Gate"
        and table.rows[0].cells[1].text.strip() == "Final result"
    ]
    if len(validation_tables) != 1:
        raise ValueError("manual validation summary table changed unexpectedly")
    validation_rows = {
        row.cells[0].text.strip(): row.cells[1]
        for row in validation_tables[0].rows[1:]
        if len(row.cells) >= 2
    }
    results = {
        "Full pytest": (
            "1,811 passed; 6 expected host/platform capability skips; 0 failed "
            "in the authoritative serial release run"
        ),
        "Compile": "346/346 product Python files compiled",
        "Static/runtime quality": (
            "Ruff clean; 82/82 module files imported; 80 capabilities discovered "
            "with zero errors or duplicates; documentation-drift and patch-"
            "integrity gates passed"
        ),
        "Module self-tests": (
            "92 pass; 12 expected inactive/platform skips; EventBus pipeline passed"
        ),
        "Application selfcheck": "26/26 passed in both direct and batch entry points",
    }
    if not set(results).issubset(validation_rows):
        raise ValueError("manual validation summary rows changed unexpectedly")
    for field, value in results.items():
        _set_text(validation_rows[field].paragraphs[0], value)

    # The legacy table mixed direct run sizes across rows. Normalize the gate
    # evidence so every result has the same visual weight and remains readable.
    for row in validation_tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    verdict_cells = [
        cell
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.startswith("Final red-team verdict:")
    ]
    if len(verdict_cells) != 1:
        raise ValueError("manual final-verdict callout changed unexpectedly")
    _set_text(
        verdict_cells[0].paragraphs[0],
        "Final red-team verdict: No open Critical, High, or Medium Cycle 25 "
        "code blocker remains in the current v1.12.0 tree. Clean-machine "
        "deployment, publisher custody, hardware-backed rollback resistance, "
        "fleet-scale trials, and independent efficacy review remain external "
        "acceptance work.",
    )

    footer_matches = [
        paragraph
        for paragraph in document.sections[0].first_page_footer.paragraphs
        if "Angerona Suite" in paragraph.text
    ]
    if len(footer_matches) != 1:
        raise ValueError("manual first-page footer changed unexpectedly")
    _set_text(footer_matches[0], "Angerona Suite  |  v1.12.0  |  28 August 2026")

    header_updates = 0
    for section in document.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            for paragraph in header.paragraphs:
                if "ANGERONA  |  MASTER MANUAL  |" in paragraph.text:
                    _set_text(paragraph, "ANGERONA  |  MASTER MANUAL  |  v1.12.0")
                    header_updates += 1
    if not header_updates:
        raise ValueError("manual running header changed unexpectedly")

    version_table = document.tables[20]
    if version_table.rows[0].cells[0].text.strip() != "Version":
        raise ValueError("manual version-history table changed unexpectedly")
    existing_versions = {row.cells[0].text.strip() for row in version_table.rows[1:]}
    if "1.11.0" in existing_versions or "1.12.0" in existing_versions:
        raise ValueError("Cycle 25 version rows already exist in pristine manual")
    rows = (
        (
            "1.11.0",
            "Independent-trust, authenticated continuity, hostile-path, governed "
            "memory, recovery-assurance, and signed-release boundary hardening.",
        ),
        (
            "1.12.0",
            "Guided Auto Adapt, immutable firewall recovery baseline, safe all-"
            "profile checkup, clickable evidence detail, 80 universal capability "
            "contracts, durable integrations, and pinned standards truth.",
        ),
    )
    template_row = version_table.rows[-1]._tr
    for version, change in rows:
        new_row = deepcopy(template_row)
        version_table._tbl.append(new_row)
        added = version_table.rows[-1]
        _set_text(added.cells[0].paragraphs[0], version)
        _set_text(added.cells[1].paragraphs[0], change)
        _keep_table_row_together(added)


def _update_existing_text(document: Document) -> None:
    _set_text(_find_exact(document, "27 August 2026"), "28 August 2026")

    _set_text(
        _find_startswith(document, "Run a fresh technique-bound campaign."),
        "Run a fresh technique-bound campaign. Only fresh evidence closes a "
        "gap; a later miss reopens it.",
    )

    emergency_cells = [
        cell
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.startswith("Do not improvise:")
    ]
    if len(emergency_cells) != 1 or len(emergency_cells[0].paragraphs[0].runs) < 2:
        raise ValueError("manual emergency-recovery callout changed unexpectedly")
    emergency_paragraph = emergency_cells[0].paragraphs[0]
    emergency_paragraph.runs[0].text = "Do not improvise: "
    emergency_paragraph.runs[1].text = (
        "Do not delete the journal, bulk-delete firewall policy, kill processes "
        "by name, or move protected evidence. Preserve the signed record and "
        "recover only the exact target."
    )
    for run in emergency_paragraph.runs[2:]:
        run.text = ""

    _set_text(
        _find_exact(document, "15.4 Cycle 23 assurance boundary"),
        "15.4 Inherited Cycle 23 assurance boundary",
    )

    section_two = _find_exact_style(
        document, "2. Product boundary and platform support", "Heading 1"
    )
    previous = section_two._p.getprevious()
    if previous is None or previous.tag != qn("w:p"):
        raise ValueError("section 2 no longer follows its expected spacer paragraph")
    previous_paragraph = next(
        (p for p in document.paragraphs if p._p is previous),
        None,
    )
    if previous_paragraph is None or previous_paragraph.text.strip():
        raise ValueError("section 2 spacer paragraph changed unexpectedly")
    _remove_paragraph(previous_paragraph)

    overview = _find_startswith(document, "Angerona unifies Windows endpoint")
    _set_text(
        overview,
        "Angerona unifies Windows endpoint and network visibility, local evidence "
        "retention, case work, MITRE ATT&CK mapping, governed containment, "
        "recovery, local AI assistance, and non-destructive purple-team "
        "validation. The v1.12 inventory discovers 80 defensive capabilities. "
        "Every capability receives a validated operational contract and health "
        "snapshot; five declare native contracts and 75 use explicit compatibility "
        "adapters. Product version and individual module versions remain separate.",
    )

    live_summary = _find_startswith(
        document, "The dashboard now includes a bounded Live Defense Activity card"
    )
    _set_text(
        live_summary,
        "The dashboard includes bounded Live Defense Activity. Its compact rows "
        "remain sanitized, but each eligible row is now clickable and opens a "
        "governed read-only evidence view with available event, subject, source, "
        "lineage, command, authenticity, fingerprint, and file-path fields. "
        "Missing or restricted fields stay explicit rather than being invented.",
    )
    live_boundary = _find_startswith(
        document, "Use this surface to confirm coarse module and EventBus activity"
    )
    _set_text(
        live_boundary,
        "Use Live Defense Activity to move from a coarse signal into bounded "
        "evidence and explanation. It never exposes source code, secrets, hidden "
        "model reasoning, or chain of thought; 'thinking' means a documented "
        "event interpretation and provenance, not private model internals.",
    )

    adaptation = _find_startswith(document, "Host Adaptation inventories hardware")
    _set_text(
        adaptation,
        "Host Adaptation inventories hardware, services, listeners, network context, "
        "and Windows Firewall state. Auto Adapt prompts for Balanced, Public, or "
        "Emergency Lockdown posture, audits completeness, builds an immutable plan, "
        "and simulates without writes. Apply still requires a second exact-command "
        "approval; automation remains proposal-only.",
    )
    anchor = adaptation
    for text in (
        "Enroll Recovery Baseline stores one exclusive, host-bound, authenticated, "
        "digest-verified complete Windows Firewall policy; it cannot silently replace "
        "the anchor.",
        "Restore snapshots pre-state, journals the action, verifies exact firewall "
        "state, and compensates or declares recovery-required.",
        "Safe automatic checkup audits once and simulates every profile without "
        "writes; it never remediates unattended.",
        "The baseline restores only firewall policy; hardware, services, ports, and "
        "network context remain observational.",
    ):
        anchor = _insert_after(document, anchor, text, "List Bullet")

    approved_pack = _find_startswith(document, "The current approved pack is")
    approved_pack.paragraph_format.page_break_before = True

    # Keep dense legacy sections from crowding the running footer. These are
    # deliberate editorial breaks, not content changes, and make each logical
    # block independently scannable in both Word and the canonical PDF render.
    for heading in (
        "6.5 Audit Log Integrity Guard",
        "6.6 Zero-Trust Network Path Monitor",
        "7.2 Action lifecycle",
        "15.3 Current research backlog",
        "17.2 Automatic Red Team validation correction (2026-08-25)",
        "17.5 v1.10.3 state-grade defensive hardening (2026-08-26)",
        "17.6 v1.11.0 independent-trust hardening (2026-08-27)",
        "Appendix B. Evidence sources",
    ):
        _find_exact(document, heading).paragraph_format.page_break_before = True

    _set_text(
        _find_startswith(document, "Candid rating: rare and ambitious"),
        "Candid rating: a credible advanced local security lab, not a commercial-"
        "parity claim. Its value is explainable exact authority, truthful receipts, "
        "verified recovery, red-team closure, measured performance, and explicit limits.",
    )

    resilience_updates = (
        (
            "Watchdog and supervisor track liveness",
            "Watchdog and supervisor use bounded restart/backoff and safe mode, "
            "never uncontrolled loops.",
        ),
        (
            "Event and storage lanes are bounded",
            "Event and storage lanes are bounded, batch safely, and fall back "
            "durably when queues saturate.",
        ),
        (
            "Source Sandbox is inert",
            "Source Sandbox is inert and path-confined; it rejects installers and "
            "never hot-reloads production code.",
        ),
        (
            "Shutdown helpers verify exact Angerona launch grammar",
            "Shutdown helpers verify exact launches and process birth, excluding "
            "test and substring lookalikes.",
        ),
    )
    for prefix, replacement in resilience_updates:
        _set_text(_find_startswith(document, prefix), replacement)

    soar_tables = [
        table
        for table in document.tables
        if table.rows
        and len(table.rows[0].cells) >= 2
        and table.rows[0].cells[0].text.strip() == "Workflow"
        and table.rows[0].cells[1].text.strip() == "Boundary"
    ]
    if len(soar_tables) != 1:
        raise ValueError("manual SOAR workflow table changed unexpectedly")
    soar_rows = {row.cells[0].text.strip(): row.cells[1] for row in soar_tables[0].rows[1:]}
    soar_updates = {
        "Manual suspend": (
            "Exact `suspend_process` contract delegated to Combat; never direct psutil."
        ),
        "Mobile directives": (
            "Observe/delegate only within exact response bounds; no cached rollback "
            "or issuance claim."
        ),
        "Posture guidance": (
            "AI advice is inert; only vetted deterministic remediation enters an "
            "action path."
        ),
    }
    if not set(soar_updates).issubset(soar_rows):
        raise ValueError("manual SOAR workflow rows changed unexpectedly")
    for field, value in soar_updates.items():
        _set_text(soar_rows[field].paragraphs[0], value)

    bug_heading = _find_exact(document, "13.1 Final bug closures")
    old_validation_heading = _find_exact(document, "13.2 Cycle 23 final validation")
    cursor = bug_heading._p.getnext()
    old_bug_paragraphs = []
    while cursor is not None and cursor is not old_validation_heading._p:
        paragraph = next((p for p in document.paragraphs if p._p is cursor), None)
        if paragraph is None:
            raise ValueError("unexpected non-paragraph content in bug-closure section")
        old_bug_paragraphs.append(paragraph)
        cursor = cursor.getnext()
    if len(old_bug_paragraphs) != 10:
        raise ValueError(
            f"expected 10 legacy bug-closure paragraphs, found {len(old_bug_paragraphs)}"
        )
    for paragraph in old_bug_paragraphs:
        _remove_paragraph(paragraph)
    anchor = bug_heading
    for text in (
        "Configuration and Settings changes are atomic or compensated; alert identity "
        "is stable; analysis work is bounded to two active and six queued jobs.",
        "Durable integration outboxes, atomic Intel Sync replacement, typed persistence "
        "completeness, and explicit cancellation preserve failure truth across restart.",
        "Allow/Undo is scoped and expiring, SOAR archive is recoverable, and remediation "
        "binds exact process birth, driver prior state, and firewall identity.",
        "The six pytest skips are declared host/platform capability gates; structured "
        "classification cannot waive an unknown failure.",
    ):
        anchor = _insert_after(document, anchor, text, "List Bullet")

    _set_text(
        _find_exact(document, "13.2 Cycle 23 final validation"),
        "13.2 Cycle 25 final validation",
    )
    validation_summary = _find_startswith(
        document, "Three sequential defensive review rounds closed 15 findings"
    )
    _set_text(
        validation_summary,
        "Three sequential Cycle 25 loops combined adversarial review, visionary "
        "design, current upstream project comparison, implementation, remediation, "
        "bug testing, and performance measurement. All 80 discovered capabilities "
        "were inspected through one enforceable v1.12 contract and inventory path.",
    )
    validation_lines = (
        "The authoritative serial suite passed 1,811 tests with six expected "
        "host/platform capability skips and zero failures.",
        "All 346 product Python files compiled. Ruff passed; 82 module files "
        "imported; 64 register hooks validated; 80 capabilities discovered with "
        "zero errors or duplicates; the inventory contains 61 unique capability codes.",
        "The module/self-test gate recorded 92 passes and 12 expected inactive or "
        "platform skips plus a passing EventBus pipeline. Direct and batch "
        "application selfcheck both passed 26 of 26.",
        "These gates are strong project evidence, not independent certification, "
        "commercial EDR parity, attribution, or proof of complete attack coverage.",
    )
    old_validation_prefixes = (
        "The authoritative one-process suite collected 1,465 tests",
        "All 321 product Python files compiled.",
        "The module harness produced 50 passes",
        "These gates are strong project evidence",
    )
    for prefix, replacement in zip(old_validation_prefixes, validation_lines, strict=True):
        _set_text(_find_startswith(document, prefix), replacement)

    _set_text(
        _find_exact(document, "14.2 Cycle 23 measured hot-path decisions"),
        "14.2 Cycle 25 measured hot-path decisions",
    )
    _set_text(
        _find_startswith(document, "Round 1 avoided redundant unchanged audit-state writes"),
        "The final gate moved the recorder's normal handoff to an exact-capacity "
        "C-backed queue, improving the measured path from 22.306 to 15.925 "
        "microseconds per event (28.6%) while preserving bounded loss accounting.",
    )
    _set_text(
        _find_startswith(document, "Round 3 retained direct pending-path security logic"),
        "Immutable capability summaries improved from 43.324 to 1.508 "
        "microseconds per call (96.5%). Revision-gated, fingerprinted Module "
        "Inspector refresh reduced unchanged ticks from 13.458 to 0.474 ms "
        "(96.5%) without weakening freshness, evidence, or click-through detail.",
    )
    deferred = (
        (
            "Signed Combat journal segmentation remains proposed",
            "Durable exporter batch commits remain proposed until crash, ordering, "
            "partial-acknowledgement, and replay equivalence are proven.",
        ),
        (
            "Ollama listener ownership is not cached",
            "Immutable compiled Sigma evaluation plans remain proposed until atomic "
            "admission, rule-generation, cache invalidation, and refusal tests are complete.",
        ),
        (
            "First-run Sysmon seek and asynchronous initial ARIA manager/index construction",
            "A global per-CVE detail-worker cap remains proposed so bursts across "
            "many distinct CVEs cannot exceed the already bounded per-item UI work.",
        ),
    )
    for prefix, replacement in deferred:
        _set_text(_find_startswith(document, prefix), replacement)

    _set_text(
        _find_startswith(document, "The ranked next work is a digest-pinned ATT&CK"),
        "ATT&CK 19.2, Navigator 5.3.2/layer 4.5, constrained OCSF 1.8, and atomic "
        "Sigma admission are now pinned and regression-tested. Ranked next work "
        "is the three measured proposals in section 14.1 plus physical Windows "
        "soak, clean-machine packaging, fleet-scale throughput, and independent "
        "false-positive and efficacy evaluation.",
    )
    _set_text(
        _find_startswith(document, "Repository evidence: analysis/loop/LOOP_LOG.md;"),
        "Repository evidence: analysis/loop/LOOP_LOG.md; analysis/loop/cycle23, "
        "cycle24, and cycle25 research and round1-round3 findings, dispositions, "
        "remediation, bug-test, performance, comparison, and final summaries; "
        "analysis/capability_inventory_v12.json; and the current source/tests.",
    )


def _append_clickable_workflow(document: Document) -> None:
    anchor = _find_exact_style(
        document, "6. Detection, evidence, and interoperability", "Heading 1"
    )
    blocks = (
        (
            "Heading 2",
            "5.3 Clickable capability and evidence organization",
        ),
        (
            "Normal",
            "Capability Center and the compact Modules view provide search, typed "
            "filters, header sorting, and click-through read-only details. Module "
            "details include contract origin, dependencies, source location, "
            "operational freshness/loss state, lifecycle, and supported actions.",
        ),
        (
            "List Bullet",
            "Host Adaptation tables sort by typed severity and numeric risk. "
            "Double-click opens the complete bounded JSON record without changing state.",
        ),
        (
            "List Bullet",
            "Live Alerts and event detail show available governed paths and evidence "
            "fields. A displayed path is context, not automatic permission to open, "
            "execute, delete, quarantine, or remediate it.",
        ),
        (
            "List Bullet",
            "Context Info meaning/path tables are sortable and double-clickable, so "
            "menu labels, implementation locations, and security boundaries can be "
            "inspected without hunting through multiple buttons.",
        ),
    )
    for style, text in blocks:
        paragraph = _insert_before(document, anchor, text, style)
        if style == "Heading 2":
            paragraph.paragraph_format.keep_with_next = True


def _append_release_addendum(document: Document) -> None:
    if any(p.text.strip() == MARKER for p in document.paragraphs):
        raise ValueError("Cycle 25 release addendum already exists")
    anchor = _find_exact(document, "Appendix A. Command reference")
    blocks = (
        ("Heading 2", MARKER),
        (
            "Normal",
            "Cycle 25 combines three adversarial, visionary, and current-upstream "
            "comparison loops. It challenged authority, crashes, spoofing, races, "
            "failure truth, and usability without claiming parity.",
        ),
        (
            "List Bullet",
            "Guided Auto Adapt consolidates audit, profile choice, baseline, plan, "
            "and simulation. Safe automatic checkup simulates every profile without "
            "writes; apply remains separately approved and proposal-only.",
        ),
        (
            "List Bullet",
            "Capability, module, adaptation, context, live activity, alert, and path "
            "surfaces provide searchable, sortable, clickable read-only detail.",
        ),
        (
            "List Bullet",
            "All 80 capabilities receive operational contracts and health snapshots: "
            "five native contracts and 75 explicit compatibility adapters.",
        ),
        (
            "List Bullet",
            "Durable outboxes, atomic Intel Sync, typed persistence completeness, exact "
            "tuner approvals, protected IPC secrets, and identity-bound remediation "
            "improve restart and failure safety.",
        ),
        (
            "List Bullet",
            "Standards truth is pinned to ATT&CK 19.2, Navigator 5.3.2/layer 4.5, "
            "typed OCSF 1.8, and atomic Sigma admission/refusal.",
        ),
        (
            "List Bullet",
            "Final evidence: 1,811 tests passed, six expected capability skips, zero "
            "failed; 346 product files compiled; Ruff, 80-capability discovery, "
            "module/self-test gates, and both 26-of-26 selfchecks passed.",
        ),
        (
            "Normal",
            "Comparison does not certify parity with Velociraptor, Wazuh, Fleet, "
            "osquery, Elastic, or commercial EDR/XDR. The immutable baseline restores "
            "Angerona-managed firewall policy, not every host setting.",
        ),
    )
    for style, text in blocks:
        paragraph = _insert_before(document, anchor, text, style)
        if style == "Heading 2":
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.page_break_before = True


def _append_sources(document: Document) -> None:
    anchor = _find_exact(document, "Appendix C. Glossary")
    blocks = (
        ("Heading 2", "B.2 Cycle 25 standards and upstream comparison sources"),
        (
            "Normal",
            "Velociraptor client monitoring and offline buffer: "
            "https://docs.velociraptor.app/docs/clients/monitoring/",
        ),
        (
            "Normal",
            "Velociraptor Artifact Exchange trust warning: "
            "https://docs.velociraptor.app/docs/artifacts/exchange_reference/",
        ),
        (
            "Normal",
            "Wazuh Active Response: https://documentation.wazuh.com/current/"
            "user-manual/capabilities/active-response/index.html",
        ),
        (
            "Normal",
            "Fleet GitOps policy/configuration: https://fleetdm.com/docs/configuration/yaml-files",
        ),
        (
            "Normal",
            "osquery configuration packs: https://osquery.readthedocs.io/en/5.12.1/"
            "deployment/configuration/",
        ),
        (
            "Normal",
            "Elastic detection-rules repository: https://github.com/elastic/detection-rules",
        ),
        (
            "Normal",
            "MITRE ATT&CK: https://attack.mitre.org/ | ATT&CK Navigator: "
            "https://github.com/mitre-attack/attack-navigator | OCSF schema: "
            "https://github.com/ocsf/ocsf-schema | Sigma specification: "
            "https://sigmahq.io/sigma-specification/",
        ),
        (
            "Normal",
            "These sources describe their own projects and standards. Cycle 25 "
            "comparison conclusions are Angerona-specific engineering inferences.",
        ),
    )
    for style, text in blocks:
        paragraph = _insert_before(document, anchor, text, style)
        if style == "Heading 2":
            paragraph.paragraph_format.keep_with_next = True


def update() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    BUILD_ROOT.mkdir(parents=True, exist_ok=True)
    if not PRISTINE.exists():
        probe = Document(SOURCE)
        if any(p.text.strip() == MARKER for p in probe.paragraphs):
            raise ValueError("cannot snapshot a manual that already has Cycle 25")
        shutil.copy2(SOURCE, PRISTINE)

    document = Document(PRISTINE)
    _update_existing_text(document)
    _append_clickable_workflow(document)
    _update_control_and_validation(document)
    _append_release_addendum(document)
    _append_sources(document)

    document.core_properties.title = "Angerona Master Manual"
    document.core_properties.subject = (
        "v1.12.0 operator reference with Cycle 25 guided-defense expansion"
    )
    document.core_properties.comments = (
        "Canonical manual minimally updated 28 August 2026; actor-neutral "
        "defensive evidence and explicit authority boundaries."
    )
    document.save(STAGED)

    reopened = Document(STAGED)
    visible = "\n".join(p.text for p in reopened.paragraphs)
    table_visible = "\n".join(
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    )
    required = (
        MARKER,
        "28 August 2026",
        "1,811 passed",
        "Auto Adapt prompts",
        "Enroll Recovery Baseline",
        "Safe automatic checkup",
        "five native contracts and 75 explicit compatibility adapters",
        "B.2 Cycle 25 standards and upstream comparison sources",
    )
    missing = [item for item in required if item not in visible and item not in table_visible]
    if missing:
        raise ValueError(f"updated manual is missing required content: {missing}")
    if sum(p.text.strip() == MARKER for p in reopened.paragraphs) != 1:
        raise ValueError("Cycle 25 release addendum was not unique after reopen")
    if reopened.paragraphs[11].text.strip() != "28 August 2026":
        raise ValueError("cover date failed reopen validation")
    version_rows = {
        row.cells[0].text.strip() for row in reopened.tables[20].rows[1:]
    }
    if {"1.11.0", "1.12.0"} - version_rows:
        raise ValueError("version-history rows failed reopen validation")
    os.replace(STAGED, SOURCE)
    print(f"updated {SOURCE}")


if __name__ == "__main__":
    update()
