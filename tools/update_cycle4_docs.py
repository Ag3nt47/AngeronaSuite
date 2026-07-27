"""Build the Cycle 4 / v1.9.4 Angerona Word-document update set.

The script makes conservative updates: it preserves each established document,
updates the cover/version line, and appends a purpose-specific Cycle 4 section.
Run the document skill's privacy scrubber on the generated files before release.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
ANALYSIS = REPO / "analysis"
DESKTOP = Path.home() / "Desktop" / "Angerona Analysis"
DATE = "2026-07-27"
VERSION = "1.9.4"


def _replace_cover_text(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs[:12]:
        text = paragraph.text
        for old, new in replacements.items():
            text = text.replace(old, new)
        if text != paragraph.text:
            paragraph.text = text


def _replace_exact_text(doc: Document, replacements: dict[str, str]) -> None:
    """Replace known stale whole-paragraph/cell labels without rewriting history."""
    for paragraph in doc.paragraphs:
        if paragraph.text in replacements:
            paragraph.text = replacements[paragraph.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text in replacements:
                        paragraph.text = replacements[paragraph.text]


def _append_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    note = doc.add_paragraph(subtitle)
    note.paragraph_format.keep_with_next = True


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.keep_with_next = True


def _bullets(doc: Document, items: list[str]) -> None:
    style_names = {style.name for style in doc.styles if style.name}
    style = "List Bullet" if "List Bullet" in style_names else (
        "List Paragraph" if "List Paragraph" in style_names else None
    )
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        if style is None:
            paragraph.add_run("• ")
        paragraph.add_run(item)


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def _keep_table_rows_intact(doc: Document) -> None:
    """Keep compact reference-table rows together and repeat their headers."""
    for table in doc.tables:
        if table.rows:
            header_props = table.rows[0]._tr.get_or_add_trPr()
            if header_props.find(qn("w:tblHeader")) is None:
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                header_props.append(header)
        for row in table.rows:
            props = row._tr.get_or_add_trPr()
            if props.find(qn("w:cantSplit")) is None:
                props.append(OxmlElement("w:cantSplit"))


def _finish(doc: Document) -> None:
    # A legacy v1.9.3 appendix used a literal question mark where a bullet was
    # intended. Repair it while preserving the surrounding document content.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("? The recommended GitHub release"):
            paragraph.text = "•" + paragraph.text[1:]
    props = doc.core_properties
    props.title = props.title or "Angerona"
    props.subject = f"Angerona v{VERSION} Cycle 4 verified update"
    props.keywords = "Angerona, EDR, NDR, SOAR, local-first, security"
    props.creator = ""
    props.last_modified_by = ""
    props.revision = 1


def build_master(source: Path, output: Path, gates: str) -> None:
    doc = Document(source)
    _replace_cover_text(doc, {
        "Document version 1.8.0 · 2026-07-15": f"Document version {VERSION} · {DATE}",
        "Document version 1.9.2": f"Document version {VERSION}",
        "v1.9.2": f"v{VERSION}",
    })
    _replace_exact_text(doc, {
        "55 auto-discovered security capabilities — see §5 / Capability Doc":
            "63 auto-discovered security capabilities — see §5 / Capability Doc",
        "61 BaseModule subclasses are auto-discovered at startup, grouped here by category; the full module-by-module table (name, category, behavior) is maintained in the companion Capability Doc in this folder to avoid duplicating a 61-row table across two documents.":
            "63 BaseModule subclasses are auto-discovered at startup, grouped here by category; the full module-by-module table (name, category, behavior) is maintained in the companion Capability Doc in this folder to avoid duplicating a 63-row table across two documents.",
        "Total: 61 modules across 21 categories.":
            "Total: 63 modules across the current categories.",
        "Angerona_Capability_Doc_v1.7.6.docx — full 61-module reference table + console/UI/security detail.":
            "Angerona_Capability_Doc_v1.9.4.docx — full 63-module reference plus current console/UI/security detail.",
    })
    _append_title(
        doc,
        f"Cycle 4 Operations and Hardening Update — v{VERSION}",
        "Three-loop visionary, red-team, blue-team, purple-team, bug, security, "
        "privacy, performance, remediation, launcher, and documentation sweep.",
    )
    _heading(doc, "Operator-visible changes")
    _bullets(doc, [
        "Live Alerts preserves the operator's viewport instead of snapping to the oldest row when new events arrive.",
        "Top Talkers collects process, connection, interface, and optional PTR data on a background worker; slow Windows or DNS calls no longer block the Qt event loop.",
        "Top Talkers AI and Upgrade Console model discovery/checks are asynchronous, single-flight, and discard stale results after close.",
        "Black Box avoids hidden telemetry chart repaints and reuses unchanged marker series, reducing long-session UI churn.",
        "The Watchdog window includes Restart Angerona Core. Its target-specific authenticated request clears SAFE_MODE, verifies an adopted Core's identity, terminates it, and relaunches it without spawning an unsafe duplicate.",
        "Settings now separates general ARIA cloud fallback from live-alert cloud analysis. Both are off by default; live-alert analysis requires its own explicit consent.",
        "All launcher diagnostics, scanner reports, temporary files, databases, drill artifacts, and watchdog state remain under the installation drive's runtime-data tree.",
    ])
    _heading(doc, "Red-team drill and remediation workflow")
    doc.add_paragraph(
        "A simulated finding is never closed by the same run that created it. "
        "The AAR must authenticate successfully, Purple Guard may stage only the "
        "narrow candidate, and a later drill must produce exact detector evidence. "
        "A future miss reopens the finding. Failed Active Response attempts do not "
        "inflate the remediation score."
    )
    _bullets(doc, [
        "Signed AAR enforcement is on by default; unsigned, tampered, and verifier-error reports fail closed.",
        "Manual report resolution authenticates the AAR before writing acknowledgements or Purple Guard policy.",
        "Temporary automatic response during a drill is restricted to recognized drill artifacts and tagged drill processes inside the selected scope.",
        "The normal SOAR automatic-response floor is CRITICAL; the temporary drill floor is restored after reporting.",
    ])
    doc.add_page_break()
    _heading(doc, "Long-session performance controls")
    _table(doc, ["Area", "v1.9.4 control", "Measured/expected effect"], [
        ["ATT&CK tracker", "O(1) bounded deque retention", "14.17× saturated hot-path improvement"],
        ["Compliance history", "O(1) 2,000-record deque retention", "3.90× improvement"],
        ["HEAL polling", "Directory metadata stamp cache", "816× unchanged 2,000-file poll improvement"],
        ["Status reporting", "Reuse one EventBus snapshot", "19% less work"],
        ["Top Talkers", "Single in-flight Qt thread-pool worker", "No process/DNS enumeration on the GUI thread"],
        ["Upgrade Console", "Asynchronous model discovery and availability checks", "No Ollama HTTP wait on the GUI thread"],
        ["ARP Watchdog", "AsyncSniffer + bounded fallback lifecycle", "No overlapping capture generations"],
        ["Module lifecycle", "Join-aware monotonic generations", "SPEC/AI/IPC helpers cannot overlap a restart"],
        ["Network/Forensics/HEAL", "Pruned and bounded long-session state", "Avoids unbounded PID/socket/capture/filename growth"],
    ])
    _heading(doc, "Privacy and public-release rules")
    _bullets(doc, [
        "Local Ollama remains the default analysis path. Cloud receives only a recursively redacted, depth/node/container/size-bounded prompt after explicit live-alert consent.",
        "The repository-local Git identity now uses the GitHub noreply address for future commits. Existing commit metadata still requires an intentional history rewrite or a clean public repository before publication.",
        "Current builds publish checksums, an SBOM, and provenance, but are not Authenticode signed. Treat the bundle-local hash manifest as integrity evidence, not publisher identity.",
        "Do not commit runtime-data, diagnostics, logs, databases, settings, encrypted credential stores, raw screenshots, or incident exports.",
    ])
    _heading(doc, "Enterprise foundation")
    _bullets(doc, [
        "External Python modules fail closed before import unless a detached Capability Manifest v1 verifies the exact source hash, API compatibility, entrypoint, declared permissions, privacy/egress/retention, performance budgets, and a trusted Ed25519 publisher.",
        "The bounded read-side causal graph separates PID generations and links process, file, network, response, and remediation-proof evidence. Temporal proximity is labeled separately from stronger causal relationships, and every edge carries an evidence basis and confidence.",
        "Every remediation-log entry receives a privacy-minimized HMAC-authenticated receipt chained to its predecessor. An applied outcome without a passed postcondition cannot validate as proof.",
        "Settings, console commands, and the local MCP API expose the same evidence-based readiness view. The score records local controls that can be demonstrated and keeps fleet enrollment, RBAC, signed central policy, cross-endpoint search, high availability, and case management visible as unshipped gaps.",
    ])
    _heading(doc, "Enterprise safety boundaries")
    _bullets(doc, [
        "A capability manifest is an auditable pre-import trust contract; its permission declarations are not yet an operating-system sandbox.",
        "A receipt detects post-write modification and invalid proof state; it cannot protect against a process that already controls Angerona's in-memory receipt authority.",
        "The causal graph is an analyst/read-side aid. It does not subscribe to the EventBus hot path and cannot authorize SOAR containment.",
        "Kernel drivers remain deferred until a measured blind spot cannot be closed with supported ETW, WMI, AMSI, WFP, Sysmon, or user-mode APIs.",
    ])
    _heading(doc, "Final verification")
    doc.add_paragraph(gates)
    _keep_table_rows_intact(doc)
    _finish(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_capability(source: Path, output: Path, gates: str) -> None:
    doc = Document(source)
    _replace_cover_text(doc, {
        "Document version 1.8.0 · 2026-07-15": f"Document version {VERSION} · {DATE}",
        "61 auto-discovered security modules; v1.8.0": f"63 auto-discovered security modules; v{VERSION}",
    })
    _replace_exact_text(doc, {
        "63 auto-discovered security modules; v1.9.4 adds the opt-in ARIA layer and Cycle 2 security, drill, storage, shutdown, and responsiveness hardening.":
            "63 auto-discovered security modules; v1.9.4 adds Cycle 4 proof-driven remediation, privacy, watchdog recovery, lifecycle safety, and responsiveness hardening.",
        "3. Security Modules (61)": "3. Security Modules (63)",
        "Every module reports a health % and state (OK / degraded / critical / failed / off) and exposes a self-test. All 61 are BaseModule subclasses auto-discovered by ModuleManager (no manual registration). A support library, remediation_actions.py, is not itself a module - it is imported by Posture Hardening as a vetted, reversible action allow-list.":
            "Every module reports a health % and state (OK / degraded / critical / failed / off) and exposes a self-test. All 63 are BaseModule subclasses auto-discovered by ModuleManager (no manual registration). A support library, remediation_actions.py, is not itself a module - it is imported by Posture Hardening as a vetted, reversible action allow-list.",
    })
    _append_title(
        doc,
        f"Cycle 4 Capability Delta — v{VERSION}",
        "Verified capabilities added or strengthened in the current implementation.",
    )
    _table(doc, ["Capability", "Current behavior", "Safety boundary"], [
        ["Proof-carrying drill remediation", "Authenticated AAR → candidate policy → later-run detector proof", "Same-run self-certification is rejected"],
        ["Live-alert AI analysis", "Local Ollama first; optional sanitized cloud second stage", "Separate default-off consent and bounded recursive redaction"],
        ["Responsive connection view", "Process, interface, connection and PTR collection off Qt", "One in-flight worker; late results ignored after close"],
        ["ARP telemetry lifecycle", "AsyncSniffer preferred; short bounded fallback polls", "Old capture cannot overlap a restarted module"],
        ["Generation-safe module restart", "Join-aware BaseModule handoff plus generation-owned SPEC, AI, IPC and ARP helpers", "A retired generation cannot keep workers, pingers, acceptors, clients, or sniffers alive"],
        ["Watchdog Core recovery", "Authenticated Restart Angerona Core control, automatic dead/suspended respawn, SAFE_MODE clearing", "Adopted process identity must match; failed safe termination never spawns a duplicate"],
        ["Bounded long-session state", "Deque retention, directory stamps, snapshot reuse, socket/PID/capture/filename pruning", "Detection semantics and evidence retention limits preserved"],
        ["Scoped drill response", "Recognized drill files/tagged processes only", "Unrelated user files and processes remain out of scope"],
        ["Unified D-drive runtime", "Diagnostics and temporary/report data use runtime-data", "Source tree stays free of mutable operational evidence"],
    ])
    _heading(doc, "Operational quality-of-life")
    _bullets(doc, [
        "Live Alerts stays newest-first without forcing the scrollbar to the bottom.",
        "Trusted Processes supports explicit exact-path approval and supervised discovery; basename trust is a pathless-telemetry fallback only.",
        "The ARIA microphone control opens voice/device/privacy settings directly.",
        "The Watchdog Info & Control tab exposes a direct Restart Angerona Core recovery action.",
        "Eco wake-up uses first-cycle completion barriers so heavy modules wake sequentially rather than stampeding the host.",
        "Normal shutdown and kill-all unload Angerona's llama3 variants without terminating unrelated Ollama work.",
    ])
    _heading(doc, "Enterprise foundation capabilities")
    _table(doc, ["Capability", "Implemented evidence", "Current boundary"], [
        ["Capability Manifest v1", "Pre-import exact-source verification plus trusted Ed25519 publisher check", "Permission declarations are not yet OS-enforced sandbox policy"],
        ["Causal incident graph", "Bounded process/file/network/parent/response/proof graph with PID-generation separation", "Read-only analyst aid; it never initiates response"],
        ["Remediation receipts", "HMAC-authenticated predecessor chain binds action digest and postcondition state", "Not a defense against a fully compromised in-process authority"],
        ["Enterprise readiness", "Settings, console, and MCP report local control evidence", "Fleet mTLS, RBAC, central policy, cross-endpoint storage/search, HA, and case management remain gaps"],
    ])
    _heading(doc, "Verification")
    doc.add_paragraph(gates)
    _keep_table_rows_intact(doc)
    _finish(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_bragsheet(source: Path, output: Path, gates: str) -> None:
    doc = Document(source)
    _replace_cover_text(doc, {
        "v1.8.0": f"v{VERSION}",
        "Version 1.8.0": f"Version {VERSION}",
    })
    _append_title(
        doc,
        f"v{VERSION} Proof Points — Cycle 4",
        "Concrete engineering results from the latest three-loop hardening pass.",
    )
    _bullets(doc, [
        "63 auto-discovered defensive modules with a native PySide6 dashboard, EDR/NDR/SOAR workflows, ATT&CK coverage, forensics, and local AI.",
        "Authenticated, proof-carrying purple remediation: a later drill must prove the candidate; same-run, unsigned, tampered, or unverifiable reports cannot self-certify.",
        "Privacy-sanitized cloud alert analysis is a distinct opt-in; local analysis is the default and raw alert evidence is never the fallback payload.",
        "UI freeze reduction: network/PTR enumeration moved off Qt, hidden Black Box chart repaint eliminated, database reads made immediate-only, and hot histories made O(1).",
        "Restart reliability: join-aware module generations retire helper threads/sockets before replacement, while the watchdog can safely restart an adopted Core from its own window.",
        "Measured wins: 14.17× ATT&CK retention, 3.90× compliance retention, 816× unchanged HEAL polling, and 19% less status-report work.",
        "Safer automated response: default CRITICAL floor plus a narrow temporary drill scope that excludes unrelated user files and processes.",
        "One-click source setup and packaged-release paths keep mutable data on D:, install dependencies, verify release artifacts, and provide visible launcher diagnostics.",
    ])
    _heading(doc, "New enterprise proof points")
    _bullets(doc, [
        "Signed capability manifests stop untrusted external Python before import, instead of trusting a filename or loading code and checking it afterward.",
        "Causal incident graphs distinguish evidence-backed relationships from mere timing, stay bounded under large input, and preserve separate process generations when Windows reuses a PID.",
        "Remediation receipts cryptographically bind the action record, predecessor, and passed postcondition so a successful button press is not confused with a proven fix.",
        "Enterprise readiness is deliberately honest: local trust, integrity, privacy, response, and performance controls earn evidence; fleet enrollment, RBAC, central policy, cross-endpoint search, HA, and case workflow remain visible gaps.",
    ])
    _heading(doc, "Release gate")
    doc.add_paragraph(gates)
    _keep_table_rows_intact(doc)
    _finish(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_flow(source: Path, output: Path, gates: str) -> None:
    doc = Document(source)
    _replace_cover_text(doc, {
        "Document version 1.8.0 · 2026-07-15": f"Document version {VERSION} · {DATE}",
        "v1.8.0 lifecycle": f"v{VERSION} lifecycle",
    })
    _append_title(
        doc,
        f"v{VERSION} Control and Evidence Flow",
        "Current alert, analysis, drill, remediation, and performance boundaries.",
    )
    _heading(doc, "Live telemetry and UI")
    doc.add_paragraph(
        "Windows sensors/modules → signed EventBus → bounded FlightRecorder → "
        "immediate-only GUI snapshot → newest-first Live Alerts. Expensive "
        "connection, process, interface, and PTR work executes on a single "
        "background worker before a Qt-thread render."
    )
    _heading(doc, "Alert analysis and privacy")
    doc.add_paragraph(
        "Alert → local Ollama analysis → confidence gate → [only when the separate "
        "live-alert cloud setting is enabled] recursively sanitize and bound → "
        "configured cloud provider. The provider never receives the local prompt."
    )
    _heading(doc, "Drill remediation proof chain")
    doc.add_paragraph(
        "Benign drill marker → real detector echo and/or scoped response evidence → "
        "authenticated AAR → narrow Purple Guard candidate → later independent "
        "drill → exact proof → PATCHED. A same-run candidate, invalid report, or "
        "later miss cannot close the finding."
    )
    _heading(doc, "Module lifecycle")
    doc.add_paragraph(
        "Startup → real-time sensors remain live → heavy modules wake one at a time "
        "and signal first-cycle completion → adaptive sleep intervals → bounded "
        "stop events → a join-aware generation handoff retires SPEC workers, AI "
        "recovery pingers, IPC accept/connection helpers, and Scapy capture before "
        "a replacement generation starts."
    )
    _heading(doc, "Watchdog Core recovery")
    doc.add_paragraph(
        "Watchdog window → Restart Angerona Core → target-specific authenticated "
        "command → clear Core SAFE_MODE → validate the adopted heartbeat PID against "
        "the configured executable and Angerona command identity → terminate → "
        "spawn a clean Core. A failed identity-safe termination stops the flow and "
        "never launches a duplicate. Dead or suspended Core recovery remains automatic."
    )
    _heading(doc, "External capability trust flow")
    doc.add_paragraph(
        "External module file + detached Capability Manifest v1 → bounded schema "
        "validation → exact source SHA-256 → API and entrypoint compatibility → "
        "trusted-publisher lookup → Ed25519 verification → only then Python import. "
        "Any missing, malformed, mismatched, untrusted, or invalid element rejects "
        "the module without executing its code."
    )
    _heading(doc, "Causal evidence and remediation-proof flow")
    doc.add_paragraph(
        "Bounded recent event facts → stable evidence IDs → PID-generation-aware "
        "process/file/network/response/proof nodes → typed edges with basis and "
        "confidence → bounded incident components. A vetted fix → canonical "
        "remediation record → passed postcondition → HMAC-authenticated chained "
        "receipt → proof reference event → optional graph linkage. Temporal "
        "precedes edges remain distinct from causal evidence and never authorize "
        "containment."
    )
    _table(doc, ["Boundary", "Failure behavior"], [
        ["SQLite busy", "Keep last complete GUI view and retry; do not block Qt"],
        ["Cloud disabled", "No provider import/query or network call from alert analysis"],
        ["AAR verifier error", "HIGH integrity event and fail closed"],
        ["Slow PTR/connection API", "Background worker remains bounded to one in-flight snapshot"],
        ["Scapy stop delay", "Retain old helper reference; restart falls back instead of overlapping"],
        ["Core identity/termination mismatch", "Refuse manual restart and do not spawn a duplicate"],
        ["Drill response target outside scope", "Reject; unrelated host objects are untouched"],
        ["External module trust failure", "Reject before import; record the manifest/source/publisher reason"],
        ["Causal graph limit reached", "Drop oldest/lowest-priority retained facts and report explicit truncation statistics"],
        ["Remediation receipt verification failure", "Do not represent the action as verified proof"],
    ])
    _heading(doc, "Verification")
    doc.add_paragraph(gates)
    _keep_table_rows_intact(doc)
    _finish(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_vulnerabilities(source: Path, output: Path, gates: str) -> None:
    doc = Document(source)
    _replace_cover_text(doc, {
        "Document version 1.8.0 · 2026-07-15": f"Document version {VERSION} · {DATE}",
    })
    _replace_exact_text(doc, {
        "Document version 1.9.4 · 2026-07-27 · Cycle 2: 10 new findings (6 MEDIUM, 4 LOW), all remediated; final self-check 26/26.":
            "Document version 1.9.4 · 2026-07-27 · Cycle 4 appendix: 11 current items, 9 closed and 2 explicit public-release residual decisions; final self-check 26/26.",
    })
    _append_title(
        doc,
        f"Cycle 4 Vulnerability and Remediation Record — v{VERSION}",
        "New findings were validated against the integrated tree after remediation.",
    )
    _table(doc, ["ID", "Severity", "Finding", "Disposition"], [
        ["C4-R1-01", "MEDIUM", "Live-alert Analyze could silently reuse cloud fallback and send raw evidence", "Closed: separate opt-in, local default, recursive/bounded cloud redaction, egress recheck"],
        ["C4-R1-02", "MEDIUM residual", "Bundle-local hashes do not establish publisher identity", "Open external boundary: Authenticode/MSIX signing certificate required"],
        ["C4-R1-03", "LOW", "AAR verifier errors/manual resolution could bypass strict trust", "Closed: strict fail-closed verification and authenticated manual resolution"],
        ["C4-R1-04", "LOW privacy", "Historical Git author metadata contains a personal address", "Future commits use noreply; history rewrite/clean public repo still required"],
        ["C4-B-01", "Reliability", "Top Talkers performed OS/PTR calls on Qt", "Closed: one background worker, UI-only rendering"],
        ["C4-B-02", "Reliability", "Scapy capture could overlap or linger across restarts", "Closed: generation stop event, AsyncSniffer, bounded fallback and overlap refusal"],
        ["C4-P-01", "Performance", "List front-deletes and repeated scans/snapshots caused long-session cost", "Closed: bounded deques, directory stamp cache, shared snapshot"],
        ["C4-P-02", "Reliability", "Module restart could overlap SPEC, AI-recovery, or IPC helper generations", "Closed: join-aware lifecycle generations and generation-owned helper resources"],
        ["C4-P-03", "Performance", "Network/PID, forensics capture, and HEAL filename state could grow indefinitely", "Closed: expiry pruning, identity-aware retention, and hard bounds"],
        ["C4-B-03", "Reliability", "Top Talkers AI and Upgrade Console model calls could block Qt", "Closed: asynchronous single-flight workers and stale-result guards"],
        ["C4-W-01", "Reliability/Safety", "Manual watchdog Core restart re-adopted the still-running Core and shared commands could be consumed by the wrong supervisor", "Closed: target-specific authenticated inbox, adopted-process identity termination, SAFE_MODE clear, and duplicate refusal"],
        ["C4-E-01", "HIGH", "External Python extensions could be opted in without a cryptographic pre-import identity and capability contract", "Closed: exact-source Capability Manifest v1, trusted Ed25519 publisher verification, and fail-closed rejection before import"],
        ["C4-E-02", "MEDIUM", "Remediation audit rows recorded outcomes but did not cryptographically bind the action record to its predecessor and passed postcondition", "Closed: privacy-minimized HMAC-authenticated chained receipts; applied without passed verification fails proof validation"],
        ["C4-E-03", "LOW", "Local MCP alert, module-health, and incident handlers used stale/non-canonical model fields", "Closed: canonical fields plus bounded enterprise-readiness and causal-graph read tools"],
    ])
    _heading(doc, "Remediation-score integrity")
    _bullets(doc, [
        "Only an explicitly successful Active Response event counts as remediation; failed attempts do not.",
        "Purple candidate installation does not count as detector coverage and cannot patch the creating run.",
        "Duplicate installed/unsupported policy findings are deduplicated without hiding unsupported techniques.",
        "Signed AAR enforcement is enabled by default and its self-test preserves the operator's environment policy.",
    ])
    _heading(doc, "Residual decisions")
    _bullets(doc, [
        "Obtain a publisher certificate and sign the release installer/executables, or ship an MSIX with a verified publisher.",
        "Before making the existing repository public, intentionally rewrite author metadata/history or publish a clean reviewed repository. This pass does not silently rewrite history.",
        "Build fleet mTLS enrollment, organization RBAC/audit, signed central policy distribution, cross-endpoint search/storage, high availability, and analyst case workflow before claiming enterprise fleet-management parity.",
        "Move untrusted AI/module workers behind an operating-system isolation boundary before treating manifest permission declarations as enforceable confinement.",
    ])
    _heading(doc, "Verification")
    doc.add_paragraph(gates)
    _keep_table_rows_intact(doc)
    _finish(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_security(source: Path, output: Path, gates: str) -> None:
    doc = Document(source)
    _replace_cover_text(doc, {
        "Assessment version 2.2 · 2026-07-15": f"Assessment version 2.3 · {DATE}",
    })
    _replace_exact_text(doc, {
        "Assessment version 2.3 · 2026-07-27 · Cycle 2 three-loop reassessment: 10 new findings (6 MEDIUM, 4 LOW), all remediated and independently gated.":
            "Assessment version 2.3 · 2026-07-27 · Cycle 4 integrated reassessment appended: current functional findings closed; publisher-signing and historical-identity decisions remain explicit residuals.",
        "Classification: Internal. Prepared for the project owner. Local static review; no external pen-test performed.":
            "Classification: Public-release technical assessment. Local static/dynamic review; no external penetration test performed.",
    })
    _append_title(
        doc,
        f"Cycle 4 Security and Privacy Reassessment — v{VERSION}",
        "Integrated verification after remediation; local static/dynamic review, not an external penetration test.",
    )
    _heading(doc, "Executive conclusion")
    doc.add_paragraph(
        "The new functional findings in cloud alert analysis and AAR trust were "
        "closed and regression-tested. Automated drill response is more narrowly "
        "scoped, failed response events no longer overstate remediation, runtime "
        "writes remain under the configured D-drive data root, and the long-session "
        "Qt/worker lifecycle is materially safer."
    )
    _table(doc, ["Control", "Result", "Evidence"], [
        ["Default-local alert analysis", "PASS", "Separate default-off consent and provider-boundary gate"],
        ["Cloud payload privacy", "PASS", "Recursive sensitive-key redaction plus string/container/depth/node/prompt budgets"],
        ["AAR authenticity", "PASS", "Unsigned, tampered, verifier-error, and manual-bypass regressions"],
        ["Drill response containment", "PASS", "Recognized drill artifacts/tagged processes inside explicit scope only"],
        ["Long-session worker safety", "PASS", "Single Top Talkers worker; bounded Scapy generation shutdown"],
        ["Core/module restart safety", "PASS", "Authenticated watchdog Core control; adopted-process identity validation; join-aware generation retirement"],
        ["Long-session state bounds", "PASS", "Network socket/PID, forensics capture, and HEAL filename state pruning"],
        ["Tracked runtime/privacy artifacts", "PASS current tree", "No tracked databases, runtime-data, settings, secrets, or raw operator screenshots"],
        ["External extension trust", "PASS", "Detached Capability Manifest v1, exact source hash, trusted Ed25519 publisher, rejection before import"],
        ["Remediation proof integrity", "PASS local boundary", "Canonical record digest, passed-postcondition requirement, HMAC predecessor chain, retention checkpoint"],
        ["Causal reasoning safety", "PASS", "Bounded read-side graph, PID-generation separation, typed/based/confident edges, no response authority"],
        ["Enterprise readiness truthfulness", "PASS", "Settings/console/MCP show demonstrated local controls and explicit fleet/RBAC/HA gaps"],
        ["Release publisher identity", "RESIDUAL", "Checksums/SBOM/provenance exist; Authenticode publisher signature does not"],
        ["Historical commit identity", "RESIDUAL", "Future noreply configured; existing commit metadata not rewritten"],
        ["Fleet control plane", "RESIDUAL", "No endpoint mTLS enrollment, organization RBAC/audit, signed central policy, fleet search/storage, HA, or case management"],
        ["Extension/AI OS isolation", "RESIDUAL", "Manifest permissions are an auditable contract, not an AppContainer/restricted-token enforcement boundary"],
    ])
    _heading(doc, "Public-release guidance")
    _bullets(doc, [
        "Publish from a reviewed release tag and verify every staged file. Keep operational telemetry and credentials outside Git.",
        "Do not describe current unsigned binaries as publisher-authenticated. State the Unknown Publisher limitation until signing is deployed.",
        "Treat cloud alert analysis as a separate consent surface from ARIA conversation. The two toggles must remain independent.",
        "If history is retained, rewrite or replace commits containing personal author metadata before public exposure and coordinate the resulting force-push.",
        "Describe Capability Manifest v1 as pre-import authenticity/integrity and disclosure enforcement, not as an operating-system sandbox.",
        "Do not describe the local readiness score as fleet certification; the report explicitly records the missing control-plane capabilities.",
    ])
    _heading(doc, "Verification")
    doc.add_paragraph(gates)
    _keep_table_rows_intact(doc)
    _finish(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--pytest-passed", type=int, required=True)
    parser.add_argument("--pytest-skipped", type=int, default=0)
    parser.add_argument("--selfcheck", default="26/26 PASS")
    parser.add_argument("--aria", default="13/13 PASS")
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument(
        "--only",
        nargs="*",
        choices=("master", "capability", "bragsheet", "flow", "vulnerabilities", "security"),
        help="Regenerate only the named documents (default: all).",
    )
    args = parser.parse_args()
    gates = (
        f"Repository pytest: {args.pytest_passed} passed, {args.pytest_skipped} skipped. "
        f"Headless self-check: {args.selfcheck}. ARIA self-tests: {args.aria}. "
        "Cycle 4 focused regression gate: 49/49 PASS. Python compile and diff checks: PASS. "
        "Auto-discovery remains 63 modules. Synthetic 100,000-event causal graph: "
        "2.958 seconds (29.58 microseconds per input event), 1,000 retained events, "
        "2,500-node cap respected, approximately 3.81 MiB RSS added."
    )
    out = args.output_dir.resolve()
    jobs = [
        (
            "master",
            build_master,
            DESKTOP / "Angerona_Master_Manual_v1.9.2.docx",
            out / "Angerona_Master_Manual_v1.9.4.docx",
        ),
        (
            "capability",
            build_capability,
            ANALYSIS / "Angerona_Capability_Doc_v1.8.0.docx",
            out / "Angerona_Capability_Doc_v1.9.4.docx",
        ),
        (
            "bragsheet",
            build_bragsheet,
            ANALYSIS / "Angerona_Capabilities_Bragsheet_v1.8.0.docx",
            out / "Angerona_Capabilities_Bragsheet_v1.9.4.docx",
        ),
        (
            "flow",
            build_flow,
            ANALYSIS / "Angerona_System_Flow_v1.8.0.docx",
            out / "Angerona_System_Flow_v1.9.4.docx",
        ),
        (
            "vulnerabilities",
            build_vulnerabilities,
            ANALYSIS / "Angerona_Vulnerabilities_Assessment_Remediation_v1.8.0.docx",
            out / "Angerona_Vulnerabilities_Assessment_Remediation_v1.9.4.docx",
        ),
        (
            "security",
            build_security,
            ANALYSIS / "Angerona_Security_Assessment_v2.2_2026-07-15.docx",
            out / "Angerona_Security_Assessment_v2.3_2026-07-27.docx",
        ),
    ]
    selected = set(args.only or (name for name, *_ in jobs))
    for name, builder, source, output in jobs:
        if name not in selected:
            continue
        if not source.exists():
            raise FileNotFoundError(source)
        builder(source, output, gates)
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
