"""Append the verified Cycle 10 fleet integrity record to the master manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def bullets(document: Document, values: tuple[str, ...]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _insert_after(document: Document, paragraph, text: str):
    created = document.add_paragraph(text, style=paragraph.style)
    paragraph._p.addnext(created._p)
    return created


def _remove_duplicate_page_breaks(document: Document) -> None:
    previous_was_break = False
    for paragraph in list(document.paragraphs):
        xml = paragraph._p.xml
        is_break_only = not paragraph.text.strip() and 'w:type="page"' in xml
        if is_break_only and previous_was_break:
            paragraph._element.getparent().remove(paragraph._element)
            continue
        previous_was_break = is_break_only


def update(source: Path, destination: Path) -> None:
    document = Document(source)
    if any(
        paragraph.text.strip().startswith("Cycle 10. Fleet Integrity")
        for paragraph in document.paragraphs
    ):
        raise ValueError("Cycle 10 is already present in the source manual")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Cycle 9 current-development edition"):
            paragraph.text = "Cycle 10 current-development edition  |  1 August 2026"
        elif text == (
            "14. Cycle 9 Security, Reliability, Performance, and Enterprise Evidence"
        ):
            _insert_after(
                document,
                paragraph,
                "15. Cycle 10 Fleet Integrity and Access Governance",
            )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "Cycle 10. Fleet Integrity and Access Governance",
        level=1,
    )
    document.add_paragraph(
        "Verification date: 1 August 2026. This section records the next "
        "offline-first enterprise tranche: trustworthy fleet chronology, "
        "device-bound event identity, a versioned local API contract, and a "
        "canonical least-privilege role model. These are tested local preview "
        "controls, not a claim of production fleet deployment."
    )

    document.add_heading("Fleet ingestion integrity", level=2)
    bullets(document, (
        "A tenant event identifier is permanently bound to its originating "
        "device and event body. Exact retries preserve the original observed "
        "time, server-received time, clock classification, skew, device identity, "
        "and event hash. Reuse from another device or with changed evidence fails "
        "closed.",
        "Event bodies are copied into finite, bounded plain JavaScript Object "
        "Notation. Non-string or empty keys, invalid UTF-8, custom objects, "
        "cycles, non-finite numbers, integers outside the signed 64-bit range, "
        "and byte/depth/node/container overflow are rejected before hashing or "
        "database insertion.",
        "Endpoint-observed time, server-received time, and signed clock skew are "
        "separate evidence. Clock quality is synchronized, skewed, untrusted, "
        "server-assigned, or legacy. Device last-seen health uses server time, so "
        "a drifting endpoint cannot move inventory state into the future or past.",
        "Existing Fleet Preview databases migrate in place. New timing columns "
        "and health counters are added without deleting evidence; older records "
        "remain explicitly legacy and therefore uncertain.",
    ))

    document.add_heading("Versioned local API and health evidence", level=2)
    bullets(document, (
        "The loopback Fleet Preview publishes a deterministic OpenAPI 3.1 "
        "contract for its health, contract, device inventory, event, and "
        "ingestion-health routes. The contract describes only routes actually "
        "shipped.",
        "Every versioned route except the minimal health check requires a fresh, "
        "one-time, complete-path/query/body Hash-based Message Authentication "
        "Code using Secure Hash Algorithm 256 request. The service remains "
        "loopback-only and exposes no generic command route.",
        "POST bodies require application/json and a bounded content length. "
        "Responses are no-store and nosniff, and the service header does not "
        "disclose the Python runtime version.",
        "Durable per-tenant counters record stored events, duplicate retries, "
        "clock classifications, last receipt time, and device states. The public "
        "readiness surface removes tenant and device identity and retains only "
        "bounded counts, state, and the public contract digest.",
        "Enterprise Settings can copy the API contract directly. Neither the "
        "contract nor the copy action includes a service key, tenant, endpoint, "
        "hostname, username, local path, event payload, or identity key.",
    ))

    document.add_heading("Least-privilege roles and duty separation", level=2)
    bullets(document, (
        "The canonical local role catalog contains Viewer, Analyst, Hunter, "
        "Responder, Detection Engineer, Fleet Operator, Tenant Administrator, "
        "Platform Administrator, and Auditor.",
        "Authorization remains default-deny with explicit-deny precedence, "
        "canonical hierarchical scope, expiring service principals, bounded "
        "permission wildcards, idempotent decisions, and authenticated receipts.",
        "A single principal cannot combine Auditor with an operational role in "
        "an overlapping scope. Detection Engineer cannot overlap Tenant or "
        "Platform Administrator policy-activation authority. Distinct non-"
        "overlapping scopes remain valid.",
        "The action-time Response Broker separately requires an independent "
        "approver and two approvers for high or critical operations. External "
        "identity-provider lifecycle, protected release-key ceremony, and "
        "administrator-ledger retention remain deployment gates.",
    ))

    document.add_heading("Verification and honest boundary", level=2)
    bullets(document, (
        "The authoritative serial Windows repository suite passes 576 tests with "
        "2 intentional platform-dependent skips and 0 failures.",
        "Ruff correctness checks and Python bytecode compilation pass. A focused "
        "Bandit scan reports no medium- or high-severity finding in the changed "
        "fleet and readiness boundaries.",
        "Adversarial regression coverage includes database migration, invalid and "
        "non-finite timestamps, invalid/deep/cyclic JSON, device-bound replay, "
        "concurrent ingestion, authenticated API access, content type, public-"
        "safe evidence, standard role permissions, and separation of duties.",
        "Production mutual Transport Layer Security, Single Sign-On/OpenID "
        "Connect, protected publisher signing, repository rulesets, high "
        "availability, external penetration review, and long-duration physical-"
        "host soaks remain separate gates.",
    ))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | Cycle 10 consolidated 2026-08-01 | Page "
                )
                break
    _remove_duplicate_page_breaks(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle10.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
