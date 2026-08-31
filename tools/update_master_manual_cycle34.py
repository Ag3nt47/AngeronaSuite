"""Apply the minimal Cycle 34 maintenance addendum to the master manual.

The updater preserves a pristine pre-Cycle-34 snapshot under ``.tmp`` so the
manual can be rebuilt after the commit-bound release gate without duplicating
content or accumulating layout drift.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
from datetime import datetime
import os
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Angerona_Master_Manual.docx"
BUILD_ROOT = ROOT / ".tmp" / "docx_cycle34"
PRISTINE = BUILD_ROOT / "Angerona_Master_Manual_pre_cycle34.docx"
STAGED = BUILD_ROOT / "Angerona_Master_Manual_cycle34.docx"
MARKER = "17.9 v1.13.0 Cycle 34 defensive convergence maintenance (2026-08-30)"


def _set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_runs(paragraph, replacements: tuple[tuple[str, str], ...]) -> None:
    for run in paragraph.runs:
        for old, new in replacements:
            run.text = run.text.replace(old, new)


def _find_exact(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _insert_before(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _clone_row_with_values(table, values: tuple[str, ...]) -> None:
    if len(values) != len(table.columns):
        raise ValueError("table row width changed unexpectedly")
    row_xml = deepcopy(table.rows[-1]._tr)
    table._tbl.append(row_xml)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values, strict=True):
        _set_text(cell.paragraphs[0], value)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9.0)


def _table_rows(table) -> dict[str, object]:
    return {
        row.cells[0].text.strip(): row.cells[1]
        for row in table.rows[1:]
        if len(row.cells) >= 2
    }


def _set_vertical_cell_margins(cell, value: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag in ("top", "bottom"):
        element = margins.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_dynamic_page_number(paragraph) -> None:
    """Replace a literal ``Page 1`` footer with a real PAGE field."""
    runs = paragraph.runs
    prefix_properties = deepcopy(runs[0]._r.rPr) if runs and runs[0]._r.rPr is not None else None
    number_properties = deepcopy(runs[-1]._r.rPr) if runs and runs[-1]._r.rPr is not None else None
    element = paragraph._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)

    prefix = OxmlElement("w:r")
    if prefix_properties is not None:
        prefix.append(prefix_properties)
    prefix_text = OxmlElement("w:t")
    prefix_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    prefix_text.text = "Page "
    prefix.append(prefix_text)
    element.append(prefix)

    def append_field_run(child) -> None:
        run = OxmlElement("w:r")
        if number_properties is not None:
            run.append(deepcopy(number_properties))
        run.append(child)
        element.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    append_field_run(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGE "
    append_field_run(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    append_field_run(separate)

    result = OxmlElement("w:t")
    result.text = "1"
    append_field_run(result)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    append_field_run(end)


def _split_cover_section(document: Document) -> None:
    """Give the cover its own first-page section.

    LibreOffice applies a single section's ``titlePg`` page style again at
    later explicit page breaks. That drops running furniture and, on dense
    pages, the top margin. A one-page cover section keeps the intentional cover
    footer while the body section can use one ordinary page style throughout.
    """
    if len(document.sections) != 1:
        raise ValueError("master manual cover-section structure changed unexpectedly")
    control = _find_exact(document, "Document control")
    break_paragraph = control._p.getprevious()
    if break_paragraph is None or break_paragraph.tag != qn("w:p"):
        raise ValueError("master manual cover page break is missing")
    page_breaks = break_paragraph.findall(".//" + qn("w:br"))
    if len(page_breaks) != 1 or page_breaks[0].get(qn("w:type")) != "page":
        raise ValueError("master manual cover page-break structure changed unexpectedly")

    final_section = document.element.body.sectPr
    cover_section = deepcopy(final_section)
    section_type = cover_section.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        cover_section.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")

    # Do not leave ``titlePg`` anywhere in the package. LibreOffice can reuse
    # that page style at later explicit breaks even when Word scopes it to the
    # first section. Point the one-page cover's default references at the
    # original first-page parts instead, which preserves the same appearance.
    cover_title = cover_section.find(qn("w:titlePg"))
    if cover_title is not None:
        cover_section.remove(cover_title)
    for reference_name in ("headerReference", "footerReference"):
        references = list(cover_section.findall(qn(f"w:{reference_name}")))
        default_reference = next(
            (item for item in references if item.get(qn("w:type")) == "default"),
            None,
        )
        first_reference = next(
            (item for item in references if item.get(qn("w:type")) == "first"),
            None,
        )
        if default_reference is None or first_reference is None:
            raise ValueError("master manual cover furniture changed unexpectedly")
        default_reference.set(qn("r:id"), first_reference.get(qn("r:id")))
        for reference in references:
            if reference is not default_reference:
                cover_section.remove(reference)

    properties = break_paragraph.get_or_add_pPr()
    for child in list(break_paragraph):
        if child is not properties:
            break_paragraph.remove(child)
    properties.append(cover_section)

    for element_name in ("titlePg",):
        element = final_section.find(qn(f"w:{element_name}"))
        if element is not None:
            final_section.remove(element)
    for reference_name in ("headerReference", "footerReference"):
        for reference in list(final_section.findall(qn(f"w:{reference_name}"))):
            if reference.get(qn("w:type")) in {"first", "even"}:
                final_section.remove(reference)


def _normalize_running_furniture(document: Document) -> None:
    """Use one non-cover header/footer and dynamic physical page numbers.

    The source carries a cover ``titlePg`` plus empty even-page parts while
    odd/even headers are disabled. Normalizing to a cover section and a single
    ordinary body-page style gives Word and LibreOffice the same layout.
    """
    _split_cover_section(document)
    body_section = document.sections[-1]
    if not body_section.footer.tables:
        raise ValueError("master manual running footer table is missing")
    footer_table = body_section.footer.tables[0]
    if len(footer_table.rows) != 1 or len(footer_table.columns) != 2:
        raise ValueError("master manual running footer structure changed unexpectedly")
    _set_dynamic_page_number(footer_table.cell(0, 1).paragraphs[0])

    referenced_parts = {
        reference.get(qn("r:id"))
        for section in document.sections
        for reference_name in ("headerReference", "footerReference")
        for reference in section._sectPr.findall(qn(f"w:{reference_name}"))
    }
    for relationship_id, relationship in list(document.part.rels.items()):
        if (
            relationship.reltype.rsplit("/", 1)[-1] in {"header", "footer"}
            and relationship_id not in referenced_parts
        ):
            document.part.drop_rel(relationship_id)


def _update_front_matter(document: Document, *, terminal_result: str | None) -> None:
    if len(document.tables) < 23:
        raise ValueError("master manual table structure changed unexpectedly")

    _set_text(
        document.tables[0].cell(0, 0).paragraphs[0],
        "Current release: v1.13.0 with Cycle 34 post-release defensive "
        "convergence across governed detection authority, Fleet custody and "
        "lifecycle, loopback canvas isolation, AegisPath lookup bounds, and "
        "coherent publisher-key validation.",
    )

    control = _table_rows(document.tables[1])
    required = {"Version", "Release state", "Source of truth"}
    if not required.issubset(control):
        raise ValueError("manual document-control rows changed unexpectedly")
    _set_text(control["Version"].paragraphs[0], "1.13.0")
    release_state = (
        "Cycle 34 three-round convergence and targeted validation complete; "
        "commit-bound release validation and guarded GitHub publication pending"
        if terminal_result is None
        else "Cycle 34 three-round convergence and commit-bound release validation "
        "complete; GitHub publication is restricted to the guarded canonical publisher"
    )
    _set_text(control["Release state"].paragraphs[0], release_state)
    _set_text(
        control["Source of truth"].paragraphs[0],
        "Current repository code, the unchanged 84-capability inventory, and "
        "analysis/loop/cycle34 evidence as of 30 August 2026",
    )

    validation = document.tables[16]
    rows = _table_rows(validation)
    required_validation = {
        "Full pytest",
        "Compile",
        "Static/runtime quality",
        "Module self-tests",
        "Application selfcheck",
    }
    if not required_validation.issubset(rows):
        raise ValueError("manual validation rows changed unexpectedly")
    if terminal_result is None:
        full_pytest = (
            "Cycle 34 pre-documentation serial: 2,877 passed; 14 intentional "
            "host/platform skips; one Cycle 27 Windows child-start deadline miss. "
            "The exact combat/ETW lease case then passed 2/2 after bounded timeout "
            "hardening. Commit-bound terminal release gate pending."
        )
    else:
        full_pytest = terminal_result
    _set_text(rows["Full pytest"].paragraphs[0], full_pytest)
    _set_text(rows["Compile"].paragraphs[0], "368/368 package Python files compiled")
    _set_text(
        rows["Static/runtime quality"].paragraphs[0],
        "Ruff clean; 84 capabilities discovered with zero errors or duplicate "
        "codes; Cycle 34 targeted matrix 91 passed with two expected Windows "
        "host-capability skips (symlink creation and POSIX fork); adjacent "
        "Cycle 31-33 matrix 128/128",
    )
    _set_text(
        rows["Module self-tests"].paragraphs[0],
        "93 passed; 16 expected platform, disabled, or optional-prerequisite "
        "skips; 0 failed across module and standalone core gates",
    )
    _set_text(rows["Application selfcheck"].paragraphs[0], "26/26 passed")

    replacements = (
        ("v1.12.0", "v1.13.0"),
        ("28 August 2026", "30 August 2026"),
    )
    for section in document.sections:
        parts = (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        )
        for part in parts:
            for paragraph in part.paragraphs:
                _replace_runs(paragraph, replacements)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _replace_runs(paragraph, replacements)
    _normalize_running_furniture(document)

    verdict = document.tables[18].cell(0, 0)
    _set_text(
        verdict.paragraphs[0],
        "Final red-team verdict: Cycle 34 fixed five Round 1 findings, the "
        "Round 2 detection/lifecycle and Fleet retained-evidence findings, and "
        "every Round 3 replay, ownership, cache, quota, migration, quarantine, "
        "and trust-coherence bypass. Final focused re-attacks found no open "
        "Critical, High, or Medium release blocker. Full-root rollback still "
        "needs an external witness, and Fleet remains a bounded local lab.",
    )

    # Keep the first-chapter safety callout intact on page 4. Compacting only
    # the five numbered posture paragraphs preserves readable spacing while
    # preventing the two-line callout from becoming an orphaned page.
    quick_start = _find_exact(document, "1.2 Recommended starting posture")
    paragraphs = list(document.paragraphs)
    quick_start_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph._p is quick_start._p
    )
    spacer = paragraphs[quick_start_index - 1]
    if spacer.text.strip():
        raise ValueError("quick-start table spacer changed unexpectedly")
    spacer.paragraph_format.space_after = Pt(4)
    quick_start.paragraph_format.space_before = Pt(6)
    quick_start.paragraph_format.space_after = Pt(4)
    posture_items = paragraphs[quick_start_index + 1 : quick_start_index + 6]
    if len(posture_items) != 5 or any(not paragraph.text.strip() for paragraph in posture_items):
        raise ValueError("recommended-posture list changed unexpectedly")
    for paragraph in posture_items:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05
    for row in document.tables[4].rows:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            _set_vertical_cell_margins(cell, 40)

    # Four new performance rows push the final two historical measurements to
    # a continuation page. Let the next chapter use the remaining body space
    # instead of leaving almost an entire page blank.
    security_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "15. Security posture, limits, and external gates"
        and paragraph.style.name == "Heading 1"
    ]
    if len(security_matches) != 1:
        raise ValueError("security-posture chapter heading changed unexpectedly")
    security_posture = security_matches[0]
    security_posture.paragraph_format.page_break_before = False
    security_posture.paragraph_format.keep_with_next = True


def _append_performance_rows(document: Document) -> None:
    table = document.tables[17]
    existing = {row.cells[0].text.strip() for row in table.rows}
    rows = (
        (
            "AegisPath immutable selection indexes",
            "Path lookup ~1,358x; node lookup ~8,908x in bounded fixtures",
            "Exact evidence-bound graph and selection semantics retained",
        ),
        (
            "Fleet authenticated custody fast path",
            "250 retained rows ~744.6 ms to 18.8 ms; 5,000-row cached intake ~289 ms",
            "Full startup audit, 5,000-row/8 KiB caps, cache invalidation, and rollback checks retained",
        ),
        (
            "Detection Runtime decode coalescing",
            "1,920 to 30 decodes; ~4.0x evaluated-event throughput",
            "Per-rule budgets, malformed-event failures, shadow isolation, and rate gates retained",
        ),
        (
            "Coherent registry validation",
            "64 to 2 governance reads; publisher trust reads reduced from N to 2",
            "Immutable key snapshot plus exit identity/hash and stable artifact generations",
        ),
    )
    for values in rows:
        if values[0] not in existing:
            _clone_row_with_values(table, values)


def _append_version_row(document: Document) -> None:
    table = document.tables[20]
    label = "1.13.0 / Cycle 34 maintenance"
    if label not in {row.cells[0].text.strip() for row in table.rows}:
        _clone_row_with_values(
            table,
            (
                label,
                "Governed detection owner/time/quarantine convergence; coherent "
                "publisher trust; authenticated bounded Fleet custody and "
                "restart-safe transactional quotas; isolated loopback canvas; "
                "AegisPath and Detection Runtime performance hardening.",
            ),
        )


def _append_cycle34_addendum(
    document: Document,
    *,
    validated_commit: str | None,
    evidence_sha256: str | None,
) -> None:
    anchor = _find_exact(document, "Appendix A. Command reference")
    blocks = (
        (
            "Heading 2",
            MARKER,
        ),
        (
            "Normal",
            "Cycle 34 is a three-round post-v1.13 defensive convergence pass. It "
            "keeps version 1.13.0 and the 84-capability inventory unchanged; no "
            "proposal-only research item was relabeled as shipped behavior.",
        ),
        (
            "List Bullet",
            "Round 1 closed repository-wide canvas exposure, legacy detection "
            "governance bypass, detached runtime ownership, multi-package eviction, "
            "and an unbound Fleet monitor. AegisPath gained immutable path/node indexes.",
        ),
        (
            "List Bullet",
            "Round 2 bound detection state, registry, runtime, receipts, expiry, and "
            "recovery into one fail-closed authority; hardened descriptor-only canvas "
            "serving and cancellable Local SOC startup; authenticated every retained "
            "Fleet health row; and replaced 3N+1 repeated signature verification with "
            "a guarded exact-row custody projection.",
        ),
        (
            "List Bullet",
            "Round 3 added a nondecreasing promotion time floor, creator-PID-bound "
            "runtime-owner lease tied to exact registry/state/quality/policy/clock/path "
            "authority with fork-safe descriptor reset, governance anchor, journaled "
            "quarantine convergence, safe legacy migration, zero-mutation invalid/replay "
            "intake, restart-safe Fleet quotas, transactional quota reservation, and "
            "coherent publisher-key snapshots across multi-package validation.",
        ),
        (
            "Normal",
            "Performance evidence is bounded and reproducible: Detection Runtime "
            "reduced representative JSON decodes from 1,920 to 30 (~4x throughput); "
            "registry governance reads fell from 64 to 2; Fleet's 250-row cached path "
            "fell from about 744.6 ms to 18.8 ms; and AegisPath path/node lookup work "
            "improved by roughly 1,358x/8,908x in its declared fixtures.",
        ),
        (
            "Normal",
            "Final targeted evidence: Cycle 34 91 passed with two expected Windows "
            "host-capability skips (symlink creation and POSIX fork); adjacent Cycle "
            "31-33 coverage 128/128; 368/368 "
            "package files compiled; 93 self-tests passed with 16 expected skips and "
            "zero failures; application selfcheck 26/26; final migration malformed-"
            "history matrix 16/16 and owner crash/reopen/quarantine probes clean.",
        ),
        (
            "Normal",
            "Residual boundary: local HMAC/SQLite/file anchors cannot prove freshness "
            "against rollback of the entire trusted root. Truncated or unprovable "
            "legacy detection history requires explicit operator recovery. Fleet is "
            "still a local lab with no remote transport, dispatch, HA, distributed "
            "quota, or production mTLS service; its 5,000-row retention bound may "
            "conservatively delay post-prune restart intake until refill.",
        ),
        (
            "Normal",
            "Proposal-only next work includes Runtime Custody Lease Broker (best future "
            "MVP), Domain Writer Fencing Tokens, View-Bound Action Receipts, Invariant "
            "Failure Capsules, Disposable Authority Recovery Rehearsal, Authoritative "
            "Mutation Inventory Gate, Cross-Domain Commit Envelope, and Forward-"
            "Integrity Ledger Epochs. None is implemented by this maintenance cycle.",
        ),
    )
    for style, text in blocks:
        paragraph = _insert_before(document, anchor, text, style)
        if style == "Heading 2":
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.page_break_before = True

    if validated_commit is not None or evidence_sha256 is not None:
        if not validated_commit or not evidence_sha256:
            raise ValueError("terminal commit and evidence SHA-256 must be supplied together")
        _insert_before(
            document,
            anchor,
            "Commit-bound five-check release evidence passed on "
            f"{validated_commit}. Evidence-manifest SHA-256: {evidence_sha256}. "
            "GitHub publication remains governed by tools/publish_github_update.py.",
            "Normal",
        )


def update(
    *,
    terminal_result: str | None,
    validated_commit: str | None,
    evidence_sha256: str | None,
) -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    BUILD_ROOT.mkdir(parents=True, exist_ok=True)
    if not PRISTINE.exists():
        probe = Document(SOURCE)
        if any(paragraph.text.strip() == MARKER for paragraph in probe.paragraphs):
            raise ValueError("cannot snapshot a manual that already has Cycle 34")
        shutil.copy2(SOURCE, PRISTINE)

    document = Document(PRISTINE)
    _update_front_matter(document, terminal_result=terminal_result)
    _append_performance_rows(document)
    _append_version_row(document)
    _append_cycle34_addendum(
        document,
        validated_commit=validated_commit,
        evidence_sha256=evidence_sha256,
    )

    document.core_properties.title = "Angerona Master Manual"
    document.core_properties.subject = (
        "v1.13.0 operator reference with Cycle 34 defensive convergence maintenance"
    )
    document.core_properties.comments = (
        "Canonical manual minimally updated 30 August 2026; defensive-only "
        "maintenance with explicit authority and deployment boundaries."
    )
    document.core_properties.modified = datetime(2026, 8, 30)
    document.save(STAGED)

    reopened = Document(STAGED)
    visible = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    table_visible = "\n".join(
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    )
    required = (
        MARKER,
        "Cycle 34 post-release defensive convergence",
        "91 passed",
        "1,920 to 30 decodes",
        "Runtime Custody Lease Broker",
        "1.13.0 / Cycle 34 maintenance",
    )
    missing = [item for item in required if item not in visible and item not in table_visible]
    if missing:
        raise ValueError(f"updated manual is missing required content: {missing}")
    if sum(paragraph.text.strip() == MARKER for paragraph in reopened.paragraphs) != 1:
        raise ValueError("Cycle 34 addendum is not unique after reopen")

    os.replace(STAGED, SOURCE)
    print(f"updated {SOURCE}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--terminal-result")
    parser.add_argument("--validated-commit")
    parser.add_argument("--evidence-sha256")
    args = parser.parse_args()
    update(
        terminal_result=args.terminal_result,
        validated_commit=args.validated_commit,
        evidence_sha256=args.evidence_sha256,
    )


if __name__ == "__main__":
    main()
