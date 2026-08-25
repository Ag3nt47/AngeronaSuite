"""Append the verified automatic Red Team validation correction to the manual."""
from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "Angerona_Master_Manual.docx"
HEADING = "17.2 Automatic Red Team validation correction (2026-08-25)"


def _insert_before(document: Document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def main() -> None:
    document = Document(MANUAL)
    if any(paragraph.text.strip() == HEADING for paragraph in document.paragraphs):
        print("Master Manual already contains the automatic validation correction.")
        return

    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Appendix A. Command reference"
    )
    _insert_before(document, anchor, HEADING, "Heading 2")
    _insert_before(
        document,
        anchor,
        "A live Extreme GUI campaign exposed a launch-contract defect: Auto-contain "
        "armed the response tier but did not activate Purple Guard's fixed simulation "
        "detector pack. The report therefore recorded only 4/52 eligible detections. "
        "This was a real validation-path failure, not acceptable detector latency.",
        "Normal",
    )
    bullets = (
        "Corrected launch behavior: an Auto-contain Red Team run activates all 13 "
        "reviewed simulation-only technique signatures before the first marker is "
        "created. Activation preserves earlier signed candidate lineage and fails "
        "closed if the pack cannot be persisted.",
        "Scope remains exact: the pack recognizes only inert _redteam_* artifacts in "
        "the dedicated or explicitly registered drill target and nonce-tagged idle "
        "processes. It does not convert suspicious filenames elsewhere into host-wide "
        "threat authority.",
        "AAR correlation now prefers the exact Adversary Combat receipt with a verified "
        "postcondition over an earlier delegation wrapper. A real verified Combat "
        "receipt counts as both an applied action contract and a closure.",
        "The first corrected live round achieved 52/52 detection but only 51/52 response "
        "and 12/13 closure, so the validator rejected it and automatically continued.",
        "The second round, run redteam-1787697587-bf119f, achieved 52/52 detection, "
        "52/52 automatic response, 13/13 action contracts, 13/13 verified closure, and "
        "a passing no-false-alert resilience control. Average detection was 0.57 seconds "
        "and average mitigation was 1.09 seconds.",
        "Post-run cleanup verified 223 authenticated action records, zero active "
        "reversible actions, an empty journal error, zero recovery requirements, zero "
        "remaining drill markers, and zero tagged probe processes.",
        "Regression evidence: 1,257 tests passed with three intentional platform skips; "
        "308 Python files compiled; repository-wide Ruff checks passed; the headless "
        "application self-check passed 26/26; and 128 adversary-boundary negative "
        "controls passed before the live campaign.",
    )
    for bullet in bullets:
        _insert_before(document, anchor, bullet, "List Bullet")
    _insert_before(
        document,
        anchor,
        "Interpretation limit: this proves 100% end-to-end handling of the defined "
        "13-class, 52-step inert campaign. It is not a claim that any finite product "
        "detects every possible real-world threat.",
        "Normal",
    )
    document.save(MANUAL)
    print(f"Updated {MANUAL}")


if __name__ == "__main__":
    main()
