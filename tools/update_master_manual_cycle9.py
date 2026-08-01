"""Append the verified Cycle 9 five-loop record to the consolidated manual."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def bullets(document, values) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _insert_after(document, paragraph, text: str) -> object:
    created = document.add_paragraph(text, style=paragraph.style)
    paragraph._p.addnext(created._p)
    return created


def _remove_duplicate_page_breaks(document: Document) -> None:
    """Collapse adjacent break-only paragraphs left by historical merges."""
    previous_was_break = False
    for paragraph in list(document.paragraphs):
        xml = paragraph._p.xml
        is_break_only = not paragraph.text.strip() and 'w:type="page"' in xml
        if is_break_only and previous_was_break:
            paragraph._element.getparent().remove(paragraph._element)
            continue
        previous_was_break = is_break_only


def update(source: Path, destination: Path) -> None:
    document = Document(source)
    # The consolidated source intentionally preserves historical sections, but
    # its cover, static contents, and Current-State Guide are authoritative and
    # must identify this current edition.
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Cycle 6 current-development edition"):
            paragraph.text = "Cycle 9 current-development edition  |  1 August 2026"
        elif text == "11. Cycle 6 Enterprise Hardening Sweep":
            cursor = paragraph
            for item in (
                "12. Cycle 7 Enterprise Control-Plane Foundations",
                "13. Cycle 8 Local Fleet Preview and Response Safety",
                "14. Cycle 9 Security, Reliability, Performance, and Enterprise Evidence",
            ):
                cursor = _insert_after(document, cursor, item)
        elif text == (
            "Fleet enrollment, organization RBAC/audit, centrally signed policy, "
            "cross-endpoint search/storage, HA, and case management remain future work."
        ):
            paragraph.text = (
                "Local Fleet Preview identity, scoped Role-Based Access Control, "
                "signed policy, hunts, cases, and evidence are shipped foundations. "
                "Production mutual Transport Layer Security, OpenID Connect, "
                "cross-host high availability/search, publisher signing, and "
                "independent disaster recovery remain external gates."
            )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "Cycle 9. Security, Reliability, Performance, and Enterprise Evidence",
        level=1,
    )
    document.add_paragraph(
        "Verification date: 1 August 2026. This section records a five-loop "
        "engineering pass over security and privacy, crash recovery, sustained "
        "responsiveness, enterprise controls, and public-release evidence. "
        "Claims below describe shipped local behavior; production deployment "
        "dependencies remain explicit external gates."
    )

    document.add_heading("Security and privacy boundaries", level=2)
    bullets(document, (
        "Process-baseline trust now validates exact canonical executable paths, "
        "rejects redirection and reparse points, authenticates durable state, "
        "uses non-blocking lifecycle control, and keeps trust approval separate "
        "from observation.",
        "A central outbound URL policy constrains schemes, credentials, host and "
        "address classes, redirects, request size, response size, and timeouts. "
        "Local model traffic is restricted to loopback and public retrieval "
        "rejects private, link-local, loopback, and metadata-service destinations.",
        "Extensible Markup Language parsing uses hardened parsers and explicit "
        "document-size budgets. Product network callers no longer bypass the "
        "central transport boundary.",
        "Startup and enterprise diagnostics retain exception type and impact "
        "without copying raw paths, credentials, prompts, responses, or host "
        "identifiers into public-safe evidence.",
    ))

    document.add_heading("Crash visibility and recovery", level=2)
    bullets(document, (
        "Optional service startup failures publish signed Startup Health evidence "
        "instead of disappearing in daemon threads. A protection-module loader "
        "failure is classified Critical because it can stop protection coverage.",
        "Black Box text presentation, Settings save compatibility, dialog worker "
        "lifecycle, Watchdog restart state, and hidden-window timer behavior have "
        "deterministic regression coverage.",
        "Failures remain degraded and observable rather than crashing the core. "
        "No recovery path receives a generic command shell or model authority.",
    ))

    document.add_heading("Sustained responsiveness and local model governance", level=2)
    bullets(document, (
        "World View no longer performs a synthetic model generation every eight "
        "seconds. It reads aggregate timing from the last real inference, avoiding "
        "continuous model wakeups and unnecessary processor use.",
        "All Ollama generation routes through one guarded client with loopback-only "
        "transport, bounded request/response streams, redaction before display, "
        "safe errors, and aggregate-only diagnostics. Prompt and response bodies "
        "are not cached for performance reporting.",
        "The Artificial Intelligence proxy moves upstream work off its event loop, "
        "bounds request bodies, and rejects invalid JSON. Cloud code synthesis is "
        "off by default and no longer places provider credentials in a URL.",
        "Long-running histories, voice work, scanner queues, legacy defense queues, "
        "and diagnostic text widgets have fixed bounds and non-blocking overflow "
        "behavior. Hidden threat-intelligence polling stops until the window is "
        "shown again.",
        "Closed asynchronous windows reject late completion without deleting live "
        "workers prematurely. The full lifecycle suite covers these boundaries.",
    ))

    document.add_heading("Runtime enterprise identity and readiness evidence", level=2)
    bullets(document, (
        "Enabling Fleet Preview creates or loads a protected per-endpoint Ed25519 "
        "identity, registers its stable public identity under the configured "
        "tenant, and stores the hostname only as a tenant-keyed token.",
        "The authenticated preview remains loopback-only. It uses freshness, full-"
        "target request authentication, durable replay resistance, tenant-scoped "
        "inventory, deduplicated ingestion, quarantine and revocation enforcement, "
        "and authenticated receipts.",
        "Enterprise Readiness assessment version 2 credits tested local identity, "
        "Role-Based Access Control, signed policy rollout, backup/restore, release, "
        "audit, privacy, integrity, and bounded-storage foundations. It separately "
        "lists production-only dependencies rather than treating them as local bugs.",
        "Settings shows whether Fleet Preview is disabled, configured, or running, "
        "plus a low-cardinality endpoint-identity state and registered-device count. "
        "A deterministic Copy public-safe evidence action excludes hostnames, user "
        "names, paths, credentials, event payloads, and endpoint identifiers.",
    ))

    document.add_heading("Production deployment gates", level=2)
    bullets(document, (
        "Mutual Transport Layer Security (mTLS), external certificate authority "
        "custody, and preferably non-exportable hardware-backed endpoint keys.",
        "Single Sign-On (SSO), OpenID Connect (OIDC), organization lifecycle, "
        "separation of duties, and an independently retained administrator ledger.",
        "High-availability ingestion/search, independent-host backup and disaster-"
        "recovery drills, long-duration physical-host soak evidence, and external "
        "penetration testing.",
        "Authenticode publisher identity, enforced repository rulesets, signed "
        "release promotion, and platform-vendor signing/notarization where required.",
        "No local score or test result is a claim of enterprise certification, "
        "kernel tamper-proofing, or a production multi-tenant service.",
    ))

    document.add_heading("Final verification evidence", level=2)
    bullets(document, (
        "The authoritative serial Windows suite passes 556 tests with 2 intentional "
        "platform-dependent skips and 0 failures.",
        "Python bytecode compilation and Ruff correctness checks pass. The release "
        "gate records dependency audit, documentation drift, unit tests, lint, and "
        "bytecode results as bounded content-addressed evidence.",
        "The public-facing repository and documentation are scanned for tracked "
        "credentials, private keys, user-profile paths, databases, runtime state, "
        "and cache artifacts. Synthetic screenshots remain explicitly labeled.",
    ))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "Angerona" in paragraph.text and paragraph.runs:
                paragraph.runs[0].text = (
                    "Angerona | Cycle 9 consolidated 2026-08-01 | Page "
                )
                break
    _remove_duplicate_page_breaks(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: update_master_manual_cycle9.py SOURCE.docx DESTINATION.docx")
        return 2
    update(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
